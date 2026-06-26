"""Reproducible statistical report for BunsekiChat pretest-posttest research."""

from __future__ import annotations

import json
import math
import os
import tempfile
import warnings
from datetime import datetime
from io import BytesIO

MATPLOTLIB_CACHE = os.path.join(tempfile.gettempdir(), "bunsekichat-matplotlib")
os.makedirs(MATPLOTLIB_CACHE, exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", MATPLOTLIB_CACHE)

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ALPHA = 0.05
COLORS = {
    "pre": "#2563EB",
    "post": "#DB2777",
    "gain": "#0F766E",
    "accent": "#D97706",
    "muted": "#64748B",
    "light": "#E2E8F0",
}


def _safe_json(value):
    if isinstance(value, dict):
        return value
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return {}
    try:
        parsed = json.loads(str(value))
        return parsed if isinstance(parsed, dict) else {}
    except (TypeError, ValueError, json.JSONDecodeError):
        return {}


def _number(value, digits=2, suffix=""):
    try:
        value = float(value)
    except (TypeError, ValueError):
        return "N/D"
    if not np.isfinite(value):
        return "N/D"
    return f"{value:.{digits}f}{suffix}"


def _pvalue(value):
    try:
        value = float(value)
    except (TypeError, ValueError):
        return "N/D"
    if not np.isfinite(value):
        return "N/D"
    return "< 0,001" if value < 0.001 else f"{value:.3f}".replace(".", ",")


def _bool_series(series):
    mapping = {
        "true": True,
        "1": True,
        "si": True,
        "sí": True,
        "yes": True,
        "false": False,
        "0": False,
        "no": False,
    }
    return series.map(lambda x: x if isinstance(x, bool) else mapping.get(str(x).strip().lower(), np.nan))


def _latest_completed(research_df):
    df = research_df.copy()
    if df.empty:
        return df
    for col in ("user_id", "quiz_type", "score"):
        if col not in df.columns:
            df[col] = np.nan
    if "status" in df.columns:
        df = df[df["status"].astype(str).str.lower().eq("completed")]
    df = df[df["quiz_type"].astype(str).str.lower().isin(["pretest", "posttest"])]
    df["quiz_type"] = df["quiz_type"].astype(str).str.lower()
    df["score"] = pd.to_numeric(df["score"], errors="coerce")
    if "id" in df.columns:
        df["_order"] = pd.to_numeric(df["id"], errors="coerce")
    elif "completed_at" in df.columns:
        df["_order"] = pd.to_datetime(df["completed_at"], errors="coerce").astype("int64")
    else:
        df["_order"] = np.arange(len(df))
    return (
        df.sort_values(["user_id", "quiz_type", "_order"])
        .drop_duplicates(["user_id", "quiz_type"], keep="last")
        .drop(columns=["_order"], errors="ignore")
    )


def _paired_scores(latest):
    if latest.empty:
        return pd.DataFrame(columns=["pretest", "posttest", "gain"])
    paired = latest.pivot_table(index="user_id", columns="quiz_type", values="score", aggfunc="last")
    for col in ("pretest", "posttest"):
        if col not in paired.columns:
            paired[col] = np.nan
    paired = paired[["pretest", "posttest"]].dropna().copy()
    paired["gain"] = paired["posttest"] - paired["pretest"]
    return paired


def _descriptive(series):
    values = pd.to_numeric(pd.Series(series), errors="coerce")
    valid = values.dropna()
    if valid.empty:
        return {"n": 0, "missing": int(values.isna().sum())}
    return {
        "n": int(valid.size),
        "missing": int(values.isna().sum()),
        "mean": float(valid.mean()),
        "sd": float(valid.std(ddof=1)) if valid.size > 1 else np.nan,
        "median": float(valid.median()),
        "q1": float(valid.quantile(0.25)),
        "q3": float(valid.quantile(0.75)),
        "min": float(valid.min()),
        "max": float(valid.max()),
        "skew": float(stats.skew(valid, bias=False)) if valid.size > 2 and valid.nunique() > 1 else np.nan,
    }


def _shapiro(values):
    values = pd.to_numeric(pd.Series(values), errors="coerce").dropna().to_numpy(dtype=float)
    if len(values) < 3:
        return {"n": len(values), "statistic": np.nan, "p": np.nan, "status": "Requiere al menos 3 observaciones."}
    if np.ptp(values) == 0:
        return {"n": len(values), "statistic": np.nan, "p": np.nan, "status": "Variable constante; normalidad no evaluable."}
    evaluated = values[:5000]
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        result = stats.shapiro(evaluated)
    return {
        "n": len(values),
        "statistic": float(result.statistic),
        "p": float(result.pvalue),
        "status": "Compatible con normalidad" if result.pvalue >= ALPHA else "Evidencia de no normalidad",
    }


def _bootstrap_mean_ci(values, repetitions=4000):
    values = np.asarray(values, dtype=float)
    values = values[np.isfinite(values)]
    if len(values) < 2:
        return (np.nan, np.nan)
    rng = np.random.default_rng(20260626)
    samples = rng.choice(values, size=(repetitions, len(values)), replace=True).mean(axis=1)
    return tuple(np.quantile(samples, [0.025, 0.975]))


def _paired_statistics(paired):
    result = {
        "n": len(paired),
        "normality": _shapiro(paired.get("gain", pd.Series(dtype=float))),
        "tests": [],
        "primary": None,
        "effect": {},
    }
    if len(paired) < 2:
        result["note"] = "Se requieren al menos dos pares para calcular variabilidad y contrastes; para inferencia estable se recomienda una muestra mayor."
        return result

    pre = paired["pretest"].to_numpy(dtype=float)
    post = paired["posttest"].to_numpy(dtype=float)
    diff = post - pre
    mean_diff = float(np.mean(diff))
    sd_diff = float(np.std(diff, ddof=1))
    se = sd_diff / math.sqrt(len(diff)) if sd_diff > 0 else 0.0
    critical = stats.t.ppf(0.975, len(diff) - 1)
    ci = (mean_diff - critical * se, mean_diff + critical * se)
    bootstrap_ci = _bootstrap_mean_ci(diff)
    dz = mean_diff / sd_diff if sd_diff > 0 else np.nan
    correction = 1 - (3 / (4 * len(diff) - 5)) if len(diff) > 2 else np.nan
    hedges_g = dz * correction if np.isfinite(dz) and np.isfinite(correction) else np.nan

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        t_result = stats.ttest_rel(post, pre, nan_policy="omit")
    t_stat = float(t_result.statistic) if np.isfinite(t_result.statistic) else np.nan
    t_p_two = float(t_result.pvalue) if np.isfinite(t_result.pvalue) else np.nan
    if np.isfinite(t_p_two) and np.isfinite(t_stat):
        t_p_greater = t_p_two / 2 if t_stat >= 0 else 1 - (t_p_two / 2)
    else:
        t_p_greater = np.nan
    result["tests"].append({
        "name": "t de Student pareada",
        "statistic": t_stat,
        "p_two": t_p_two,
        "p_greater": t_p_greater,
        "assumption": "Normalidad aproximada de las diferencias.",
    })

    nonzero = diff[diff != 0]
    if len(nonzero):
        try:
            w_two = stats.wilcoxon(post, pre, alternative="two-sided", zero_method="wilcox", method="auto")
            w_greater = stats.wilcoxon(post, pre, alternative="greater", zero_method="wilcox", method="auto")
            w_stat, w_p_two, w_p_greater = float(w_two.statistic), float(w_two.pvalue), float(w_greater.pvalue)
        except ValueError:
            w_stat = w_p_two = w_p_greater = np.nan
        ranks = stats.rankdata(np.abs(nonzero))
        positive_ranks = float(ranks[nonzero > 0].sum())
        negative_ranks = float(ranks[nonzero < 0].sum())
        rank_biserial = (positive_ranks - negative_ranks) / (positive_ranks + negative_ranks)
        positives = int((nonzero > 0).sum())
        sign_test = stats.binomtest(positives, len(nonzero), 0.5, alternative="greater")
        sign_p = float(sign_test.pvalue)
    else:
        w_stat = 0.0
        w_p_two = w_p_greater = sign_p = 1.0
        rank_biserial = 0.0
        positives = 0
    result["tests"].append({
        "name": "Wilcoxon de rangos con signo",
        "statistic": w_stat,
        "p_two": w_p_two,
        "p_greater": w_p_greater,
        "assumption": "Contraste no paramétrico para datos pareados.",
    })
    result["tests"].append({
        "name": "Prueba de los signos exacta",
        "statistic": positives,
        "p_two": np.nan,
        "p_greater": sign_p,
        "assumption": "Evalúa la dirección del cambio sin usar su magnitud.",
    })

    normal_p = result["normality"].get("p")
    if len(paired) < 3:
        result["primary"] = None
        result["note"] = "Muestra piloto demasiado pequeña para seleccionar de forma fiable un contraste inferencial."
    elif np.isfinite(normal_p) and normal_p >= ALPHA:
        result["primary"] = "t de Student pareada"
    else:
        result["primary"] = "Wilcoxon de rangos con signo"

    q1, q3 = np.quantile(diff, [0.25, 0.75])
    iqr = q3 - q1
    outliers = int(((diff < q1 - 1.5 * iqr) | (diff > q3 + 1.5 * iqr)).sum()) if iqr > 0 else 0
    result["effect"] = {
        "mean_diff": mean_diff,
        "median_diff": float(np.median(diff)),
        "ci_low": float(ci[0]),
        "ci_high": float(ci[1]),
        "bootstrap_low": float(bootstrap_ci[0]),
        "bootstrap_high": float(bootstrap_ci[1]),
        "cohen_dz": dz,
        "hedges_gz": hedges_g,
        "rank_biserial": rank_biserial,
        "improved": int((diff > 0).sum()),
        "unchanged": int((diff == 0).sum()),
        "decreased": int((diff < 0).sum()),
        "outliers_iqr": outliers,
    }
    return result


def _effect_label(value):
    if not np.isfinite(value):
        return "no estimable"
    magnitude = abs(value)
    if magnitude < 0.20:
        return "trivial"
    if magnitude < 0.50:
        return "pequeño"
    if magnitude < 0.80:
        return "moderado"
    return "grande"


def _holm_adjust(pvalues):
    values = np.asarray(pvalues, dtype=float)
    adjusted = np.full(len(values), np.nan)
    valid_indices = np.where(np.isfinite(values))[0]
    if not len(valid_indices):
        return adjusted
    order = valid_indices[np.argsort(values[valid_indices])]
    running = 0.0
    m = len(order)
    for rank, idx in enumerate(order):
        candidate = min(1.0, (m - rank) * values[idx])
        running = max(running, candidate)
        adjusted[idx] = running
    return adjusted


def _dimension_pairs(latest):
    rows = []
    for _, record in latest.iterrows():
        for dimension, value in _safe_json(record.get("dimension_scores_json")).items():
            try:
                score = float(value)
            except (TypeError, ValueError):
                continue
            rows.append({"user_id": record.get("user_id"), "quiz_type": record.get("quiz_type"), "dimension": str(dimension), "score": score})
    long_df = pd.DataFrame(rows)
    if long_df.empty:
        return long_df, pd.DataFrame()
    wide = long_df.pivot_table(index=["user_id", "dimension"], columns="quiz_type", values="score", aggfunc="last").reset_index()
    if {"pretest", "posttest"}.issubset(wide.columns):
        wide = wide.dropna(subset=["pretest", "posttest"]).copy()
        wide["gain"] = wide["posttest"] - wide["pretest"]
    else:
        wide = pd.DataFrame()
    return long_df, wide


def _dimension_tests(dimension_wide):
    rows = []
    if dimension_wide.empty:
        return pd.DataFrame()
    for dimension, group in dimension_wide.groupby("dimension", sort=True):
        paired = group[["pretest", "posttest", "gain"]]
        stats_result = _paired_statistics(paired)
        primary_name = stats_result.get("primary")
        primary_test = next((test for test in stats_result.get("tests", []) if test["name"] == primary_name), None)
        rows.append({
            "Dimensión": dimension,
            "n": len(group),
            "Pre M": group["pretest"].mean(),
            "Post M": group["posttest"].mean(),
            "Ganancia M": group["gain"].mean(),
            "Prueba": primary_name or "Solo descriptiva",
            "p unilateral": primary_test.get("p_greater") if primary_test else np.nan,
            "d_z": stats_result.get("effect", {}).get("cohen_dz", np.nan),
        })
    result = pd.DataFrame(rows)
    if not result.empty:
        result["p Holm"] = _holm_adjust(result["p unilateral"].to_numpy())
    return result


def _cronbach_alpha(item_matrix):
    complete = item_matrix.apply(pd.to_numeric, errors="coerce").dropna()
    n, k = complete.shape
    if n < 2 or k < 2:
        return np.nan, n, k
    total_variance = complete.sum(axis=1).var(ddof=1)
    if not np.isfinite(total_variance) or total_variance <= 0:
        return np.nan, n, k
    alpha = (k / (k - 1)) * (1 - complete.var(axis=0, ddof=1).sum() / total_variance)
    return float(alpha), n, k


def _survey_analysis(survey_df):
    result = {"items": pd.DataFrame(), "dimensions": pd.DataFrame(), "alpha": np.nan, "complete_n": 0, "k": 0}
    if survey_df is None or survey_df.empty or "score" not in survey_df.columns:
        return result
    data = survey_df.copy()
    data["score"] = pd.to_numeric(data["score"], errors="coerce")
    data = data[data["score"].between(1, 5)]
    if data.empty:
        return result
    item_key = "item_no" if "item_no" in data.columns else "item_text"
    matrix = data.pivot_table(index="user_id", columns=item_key, values="score", aggfunc="last")
    alpha, complete_n, k = _cronbach_alpha(matrix)
    item_group = data.groupby([item_key, "item_text"], dropna=False)["score"].agg(["count", "mean", "std", "median"]).reset_index()
    dimension_student = data.groupby(["user_id", "dimension"], dropna=False)["score"].mean().reset_index()
    dim_rows = []
    for dimension, group in dimension_student.groupby("dimension", dropna=False):
        values = group["score"].dropna().to_numpy(dtype=float)
        if len(values) >= 2 and np.any(values != 3):
            test = stats.wilcoxon(values - 3, alternative="greater", zero_method="wilcox", method="auto")
            statistic, p = float(test.statistic), float(test.pvalue)
        else:
            statistic = p = np.nan
        dim_rows.append({"Dimensión": str(dimension), "n": len(values), "Media": np.mean(values) if len(values) else np.nan, "DE": np.std(values, ddof=1) if len(values) > 1 else np.nan, "Mediana": np.median(values) if len(values) else np.nan, "W": statistic, "p > 3": p})
    dimensions = pd.DataFrame(dim_rows)
    if not dimensions.empty:
        dimensions["p Holm"] = _holm_adjust(dimensions["p > 3"].to_numpy())
    result.update({"items": item_group, "dimensions": dimensions, "alpha": alpha, "complete_n": complete_n, "k": k, "student_scores": dimension_student})
    return result


def _exercise_analysis(exercise_df):
    result = {"summary": pd.DataFrame(), "chi2": None}
    if exercise_df is None or exercise_df.empty or "is_correct" not in exercise_df.columns:
        return result
    data = exercise_df.copy()
    data["correct"] = _bool_series(data["is_correct"])
    data = data.dropna(subset=["correct"])
    if data.empty:
        return result
    if "difficulty_level" not in data.columns:
        data["difficulty_level"] = "Sin nivel"
    summary = data.groupby("difficulty_level", dropna=False)["correct"].agg(["count", "mean", "sum"]).reset_index()
    summary["accuracy_pct"] = summary["mean"] * 100
    table = pd.crosstab(data["difficulty_level"], data["correct"])
    chi2_result = None
    if table.shape[0] >= 2 and table.shape[1] == 2:
        chi2, p, dof, expected = stats.chi2_contingency(table)
        n = table.to_numpy().sum()
        denominator = min(table.shape[0] - 1, table.shape[1] - 1)
        cramers_v = math.sqrt(chi2 / (n * denominator)) if n and denominator else np.nan
        chi2_result = {"chi2": float(chi2), "p": float(p), "dof": int(dof), "min_expected": float(expected.min()), "cramers_v": cramers_v, "valid": bool(expected.min() >= 5)}
    result.update({"summary": summary, "chi2": chi2_result, "data": data})
    return result


def _shade_cell(cell, color="6F0F49"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def _set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _add_table(doc, headers, rows, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    _set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        _shade_cell(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph()
    return table


def _add_note(doc, title, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(f"{title}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(111, 15, 73)
    paragraph.add_run(text)


def _figure_to_stream(fig):
    stream = BytesIO()
    fig.savefig(stream, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    stream.seek(0)
    return stream


def _add_figure(doc, fig, number, title, interpretation):
    stream = _figure_to_stream(fig)
    doc.add_picture(stream, width=Inches(6.25))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"Figura {number}. {title}")
    run.bold = True
    run.font.size = Pt(9)
    _add_note(doc, "Interpretación", interpretation)


def _base_figure(figsize=(8.2, 4.6)):
    fig, ax = plt.subplots(figsize=figsize)
    ax.grid(axis="y", color="#E2E8F0", linewidth=0.8, alpha=0.8)
    ax.spines[["top", "right"]].set_visible(False)
    return fig, ax


def _score_boxplot(paired):
    fig, ax = _base_figure()
    values = [paired["pretest"].to_numpy(), paired["posttest"].to_numpy()]
    plot = ax.boxplot(values, tick_labels=["Pretest", "Postest"], patch_artist=True, widths=0.52)
    for patch, color in zip(plot["boxes"], [COLORS["pre"], COLORS["post"]]):
        patch.set_facecolor(color)
        patch.set_alpha(0.75)
    rng = np.random.default_rng(42)
    for pos, series, color in zip([1, 2], values, [COLORS["pre"], COLORS["post"]]):
        ax.scatter(pos + rng.normal(0, 0.035, len(series)), series, color=color, edgecolor="white", linewidth=0.5, zorder=3)
    ax.set_ylabel("Puntaje (0-100)")
    ax.set_ylim(0, 105)
    ax.set_title("Distribución de puntajes emparejados")
    return fig


def _paired_lines(paired):
    fig, ax = _base_figure()
    for _, row in paired.iterrows():
        color = COLORS["gain"] if row["gain"] >= 0 else "#B91C1C"
        ax.plot([0, 1], [row["pretest"], row["posttest"]], color=color, alpha=0.55, marker="o", linewidth=1.2)
    ax.set_xticks([0, 1], ["Pretest", "Postest"])
    ax.set_ylabel("Puntaje (0-100)")
    ax.set_ylim(0, 105)
    ax.set_title("Trayectoria individual de cada estudiante")
    return fig


def _gain_histogram(paired):
    fig, ax = _base_figure()
    bins = min(max(5, int(math.sqrt(len(paired))) + 1), 15)
    ax.hist(paired["gain"], bins=bins, color=COLORS["gain"], edgecolor="white", alpha=0.85)
    ax.axvline(0, color="#B91C1C", linestyle="--", linewidth=1.4, label="Sin cambio")
    ax.axvline(paired["gain"].mean(), color=COLORS["accent"], linewidth=1.7, label="Media")
    ax.set_xlabel("Ganancia (postest - pretest)")
    ax.set_ylabel("Frecuencia")
    ax.set_title("Distribución de la ganancia")
    ax.legend(frameon=False)
    return fig


def _qq_plot(paired):
    fig, ax = _base_figure()
    stats.probplot(paired["gain"].to_numpy(), dist="norm", plot=ax)
    ax.get_lines()[0].set_markerfacecolor(COLORS["gain"])
    ax.get_lines()[0].set_markeredgecolor("white")
    ax.get_lines()[1].set_color(COLORS["accent"])
    ax.set_title("Gráfico Q-Q de las diferencias")
    return fig


def _prepost_scatter(paired):
    fig, ax = _base_figure()
    ax.scatter(paired["pretest"], paired["posttest"], color=COLORS["post"], s=42, alpha=0.8, edgecolor="white")
    ax.plot([0, 100], [0, 100], color=COLORS["muted"], linestyle="--", label="Sin cambio")
    ax.set_xlabel("Puntaje pretest")
    ax.set_ylabel("Puntaje postest")
    ax.set_xlim(0, 105)
    ax.set_ylim(0, 105)
    ax.set_title("Relación entre puntajes pretest y postest")
    ax.legend(frameon=False)
    return fig


def _dimension_chart(dimension_wide):
    means = dimension_wide.groupby("dimension")[["pretest", "posttest"]].mean().sort_index()
    fig, ax = _base_figure((8.5, 5.0))
    x = np.arange(len(means))
    width = 0.36
    ax.bar(x - width / 2, means["pretest"], width, label="Pretest", color=COLORS["pre"])
    ax.bar(x + width / 2, means["posttest"], width, label="Postest", color=COLORS["post"])
    ax.set_xticks(x, means.index, rotation=20, ha="right")
    ax.set_ylabel("Puntaje medio (0-100)")
    ax.set_ylim(0, 105)
    ax.set_title("Rendimiento por dimensión")
    ax.legend(frameon=False)
    return fig


def _categorical_chart(series, title, x_label="Frecuencia"):
    counts = series.fillna("Sin dato").astype(str).value_counts().sort_values()
    fig, ax = _base_figure((8.2, max(3.2, min(6.5, 0.48 * len(counts) + 1.8))))
    ax.barh(counts.index, counts.values, color=[COLORS["pre"], COLORS["post"], COLORS["gain"], COLORS["accent"]] * (len(counts) // 4 + 1))
    ax.set_xlabel(x_label)
    ax.set_title(title)
    return fig


def _survey_dimension_chart(dimensions):
    ordered = dimensions.sort_values("Media")
    fig, ax = _base_figure((8.3, max(3.8, 0.5 * len(ordered) + 1.7)))
    ax.barh(ordered["Dimensión"], ordered["Media"], color=COLORS["gain"])
    ax.axvline(3, color=COLORS["accent"], linestyle="--", label="Punto neutral")
    ax.set_xlim(1, 5)
    ax.set_xlabel("Media Likert (1-5)")
    ax.set_title("Percepción por dimensión")
    ax.legend(frameon=False)
    return fig


def _survey_item_chart(items):
    labels = [f"Ítem {int(x)}" if str(x).replace(".", "", 1).isdigit() else str(x)[:28] for x in items.iloc[:, 0]]
    fig, ax = _base_figure((8.3, max(4.5, 0.34 * len(items) + 1.8)))
    ax.barh(labels[::-1], items["mean"].to_numpy()[::-1], color=COLORS["pre"])
    ax.axvline(3, color=COLORS["accent"], linestyle="--", label="Punto neutral")
    ax.set_xlim(1, 5)
    ax.set_xlabel("Media Likert (1-5)")
    ax.set_title("Media de cada ítem de la encuesta")
    ax.legend(frameon=False)
    return fig


def _exercise_chart(summary):
    ordered = summary.copy()
    fig, ax = _base_figure()
    bars = ax.bar(ordered["difficulty_level"].astype(str), ordered["accuracy_pct"], color=[COLORS["gain"], COLORS["accent"], COLORS["post"]][: len(ordered)])
    ax.bar_label(bars, labels=[f"{value:.1f}%\n(n={int(n)})" for value, n in zip(ordered["accuracy_pct"], ordered["count"])], padding=3, fontsize=9)
    ax.set_ylim(0, 112)
    ax.set_ylabel("Aciertos (%)")
    ax.set_title("Desempeño en ejercicios guiados por dificultad")
    return fig


def _configure_document(doc, title):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in [("Title", 22, "6F0F49"), ("Heading 1", 16, "6F0F49"), ("Heading 2", 13, "0F766E"), ("Heading 3", 11, "1E3A8A")]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("BunsekiChat | Informe estadístico reproducible").font.size = Pt(8)
    doc.core_properties.title = title
    doc.core_properties.subject = "Análisis pretest-postest y evidencia complementaria"
    doc.core_properties.author = "BunsekiChat"


def _filter_text(filters):
    if not filters:
        return "Todos los registros disponibles"
    labels = {"subject": "Materia", "course": "Curso", "parallel": "Paralelo", "shift": "Jornada", "cohort": "Cohorte"}
    parts = [f"{labels.get(key, key)}: {value}" for key, value in filters.items() if value and str(value) != "Todos"]
    return " | ".join(parts) if parts else "Todos los registros disponibles"


def _primary_interpretation(stat_result):
    if stat_result.get("primary") is None:
        return stat_result.get("note", "No fue posible ejecutar inferencia estadística.")
    test = next(test for test in stat_result["tests"] if test["name"] == stat_result["primary"])
    p = test.get("p_greater")
    effect = stat_result.get("effect", {})
    direction = "aumentó" if effect.get("mean_diff", 0) > 0 else "no aumentó"
    decision = "se rechaza H0" if np.isfinite(p) and p < ALPHA else "no se rechaza H0"
    evidence = "existe evidencia estadísticamente significativa de mejora" if np.isfinite(p) and p < ALPHA and effect.get("mean_diff", 0) > 0 else "la muestra no aporta evidencia estadística suficiente de una mejora"
    caution = "" if stat_result["n"] >= 30 else " La muestra es menor de 30 pares, por lo que la estimación debe interpretarse con cautela."
    return f"Con {stat_result['primary']} (p unilateral {_pvalue(p)}), {decision}: {evidence}. El puntaje medio {direction} {_number(effect.get('mean_diff'))} puntos y el tamaño de efecto d_z fue {_number(effect.get('cohen_dz'))} ({_effect_label(effect.get('cohen_dz', np.nan))}).{caution}"


def build_research_word_report(research_df, exercise_df=None, survey_df=None, filters=None, research_title="Efectos de un Tutor Inteligente basado en IA Generativa"):
    """Build a formatted DOCX report using the selected academic cohort."""
    research_df = pd.DataFrame(research_df).copy()
    exercise_df = pd.DataFrame() if exercise_df is None else pd.DataFrame(exercise_df).copy()
    survey_df = pd.DataFrame() if survey_df is None else pd.DataFrame(survey_df).copy()
    latest = _latest_completed(research_df)
    paired = _paired_scores(latest)
    stat_result = _paired_statistics(paired)
    dimension_long, dimension_wide = _dimension_pairs(latest)
    dimension_tests = _dimension_tests(dimension_wide)
    survey = _survey_analysis(survey_df)
    exercises = _exercise_analysis(exercise_df)

    doc = Document()
    _configure_document(doc, research_title)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Informe de análisis estadístico")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(research_title).bold = True
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n{_filter_text(filters)}")

    doc.add_heading("1. Resumen ejecutivo", level=1)
    completed_users = latest["user_id"].nunique() if not latest.empty else 0
    _add_table(doc, ["Indicador", "Resultado"], [
        ["Estudiantes con alguna medición", completed_users],
        ["Pares completos pretest-postest", len(paired)],
        ["Media pretest", _number(paired["pretest"].mean() if len(paired) else np.nan, suffix="%")],
        ["Media postest", _number(paired["posttest"].mean() if len(paired) else np.nan, suffix="%")],
        ["Ganancia media", _number(paired["gain"].mean() if len(paired) else np.nan, suffix=" puntos")],
        ["Participantes con encuesta", survey_df["user_id"].nunique() if not survey_df.empty and "user_id" in survey_df else 0],
        ["Intentos de ejercicios guiados", len(exercise_df)],
    ])
    _add_note(doc, "Conclusión principal", _primary_interpretation(stat_result))
    _add_note(doc, "Alcance causal", "El diseño de un solo grupo con pretest y postest permite estimar cambio intraestudiante. Sin grupo de control y asignación aleatoria, la mejoría no puede atribuirse exclusivamente al tutor; deben considerarse historia, maduración, práctica y otros factores externos.")

    doc.add_heading("2. Diseño, hipótesis y reglas de análisis", level=1)
    doc.add_paragraph("Unidad de análisis: estudiante con pretest y postest completados. Cuando existen intentos repetidos, se conserva el último registro completado de cada instrumento. Nivel de significación: α = 0,05.")
    doc.add_paragraph("H0: la diferencia poblacional postest - pretest no es positiva (la intervención no mejora el rendimiento).")
    doc.add_paragraph("H1: la diferencia poblacional postest - pretest es positiva (el rendimiento posterior es mayor).")
    doc.add_paragraph("La prueba principal se selecciona mediante la normalidad de las diferencias: t pareada cuando Shapiro-Wilk no rechaza normalidad y Wilcoxon cuando existe evidencia de no normalidad. Se reportan además pruebas de sensibilidad, intervalos de confianza y tamaños del efecto. Las comparaciones por dimensiones y encuesta aplican corrección de Holm por multiplicidad.")

    doc.add_heading("3. Calidad y preparación de los datos", level=1)
    pre_users = latest.loc[latest["quiz_type"].eq("pretest"), "user_id"].nunique() if not latest.empty else 0
    post_users = latest.loc[latest["quiz_type"].eq("posttest"), "user_id"].nunique() if not latest.empty else 0
    eligible = research_df.copy()
    if "status" in eligible.columns:
        eligible = eligible[eligible["status"].astype(str).str.lower().eq("completed")]
    if "quiz_type" in eligible.columns:
        eligible = eligible[eligible["quiz_type"].astype(str).str.lower().isin(["pretest", "posttest"])]
    duplicate_count = max(0, len(eligible) - len(latest))
    incomplete_count = max(0, len(research_df) - len(eligible))
    completeness = (len(paired) / max(pre_users, post_users) * 100) if max(pre_users, post_users) else 0
    _add_table(doc, ["Control", "Resultado", "Interpretación"], [
        ["Pretests únicos", pre_users, "Participantes con medición inicial"],
        ["Postests únicos", post_users, "Participantes con medición final"],
        ["Pares completos", len(paired), f"Completitud de emparejamiento: {completeness:.1f}%"],
        ["Registros repetidos excluidos", duplicate_count, "Se conservó el último intento completado"],
        ["Registros pendientes o incompletos", incomplete_count, "No ingresan en los contrastes"],
        ["Atípicos en la ganancia (regla 1,5 RIC)", stat_result.get("effect", {}).get("outliers_iqr", 0), "Se mantienen; deben revisarse, no eliminarse automáticamente"],
    ])

    doc.add_heading("4. Análisis exploratorio por variable", level=1)
    doc.add_heading("4.1 Variables cuantitativas principales", level=2)
    numeric_rows = []
    for label, series in [("Puntaje pretest", paired.get("pretest", [])), ("Puntaje postest", paired.get("posttest", [])), ("Ganancia", paired.get("gain", []))]:
        desc = _descriptive(series)
        numeric_rows.append([label, desc.get("n", 0), _number(desc.get("mean")), _number(desc.get("sd")), _number(desc.get("median")), f"{_number(desc.get('q1'))} - {_number(desc.get('q3'))}", f"{_number(desc.get('min'))} - {_number(desc.get('max'))}", _number(desc.get("skew"))])
    if "total_time_seconds" in latest.columns:
        numeric_rows.append(["Tiempo total (segundos)", *_eda_row_values(_descriptive(latest["total_time_seconds"]))])
    _add_table(doc, ["Variable", "n", "Media", "DE", "Mediana", "RIC", "Mín-Máx", "Asimetría"], numeric_rows)
    _add_note(doc, "Lectura", "La media resume el centro, la DE la dispersión, el RIC el 50% central y la asimetría la forma de la distribución. Para resultados pretest-postest, la variable decisiva es la ganancia individual.")

    if len(paired):
        figure_no = 1
        _add_figure(doc, _score_boxplot(paired), figure_no, "Distribución de puntajes pretest y postest", "Permite comparar centro, dispersión y posibles valores extremos. Una posición superior del postest sugiere mejor rendimiento posterior, pero la significación se determina con el contraste pareado.")
        figure_no += 1
        _add_figure(doc, _paired_lines(paired), figure_no, "Trayectorias individuales", f"Cada línea representa un estudiante. Mejoraron {int((paired['gain'] > 0).sum())}, permanecieron iguales {int((paired['gain'] == 0).sum())} y disminuyeron {int((paired['gain'] < 0).sum())} participantes.")
        figure_no += 1
        _add_figure(doc, _gain_histogram(paired), figure_no, "Distribución de la ganancia", f"La ganancia media fue {_number(paired['gain'].mean())} puntos. Valores a la derecha de cero indican mejora y valores a la izquierda indican disminución.")
        figure_no += 1
        if len(paired) >= 3 and paired["gain"].nunique() > 1:
            _add_figure(doc, _qq_plot(paired), figure_no, "Q-Q de las diferencias", "La cercanía de los puntos a la recta apoya normalidad aproximada. La decisión formal se complementa con Shapiro-Wilk y con el tamaño muestral.")
            figure_no += 1
        _add_figure(doc, _prepost_scatter(paired), figure_no, "Relación pretest-postest", "Los puntos por encima de la diagonal representan estudiantes que mejoraron. La dispersión muestra heterogeneidad individual y posible dependencia del resultado final respecto del nivel inicial.")
        figure_no += 1
    else:
        figure_no = 1
        doc.add_paragraph("No existen pares completos para generar gráficos pretest-postest.")

    doc.add_heading("4.2 Variables categóricas y de contexto", level=2)
    categorical_columns = [
        ("profile_subject", "Materia"),
        ("profile_course_level", "Curso"),
        ("profile_parallel", "Paralelo"),
        ("profile_shift", "Jornada"),
        ("profile_cohort", "Cohorte"),
        ("version_code", "Versión del instrumento"),
        ("cognitive_profile", "Perfil cognitivo"),
    ]
    frequency_rows = []
    for column, label in categorical_columns:
        if column not in latest.columns:
            continue
        source = latest[column].fillna("Sin dato").astype(str)
        counts = source.value_counts(dropna=False)
        for category, count in counts.items():
            frequency_rows.append([label, category, int(count), f"{count / len(source) * 100:.1f}%"])
    if frequency_rows:
        _add_table(doc, ["Variable", "Categoría", "n", "%"], frequency_rows)
    if "cognitive_profile" in latest.columns and latest["cognitive_profile"].notna().any():
        _add_figure(doc, _categorical_chart(latest["cognitive_profile"], "Distribución de perfiles cognitivos"), figure_no, "Perfiles cognitivos", "Describe la composición académica registrada por el sistema. Es una variable secundaria y no sustituye el puntaje continuo para probar la hipótesis principal.")
        figure_no += 1

    doc.add_heading("5. Supuestos y contrastes de hipótesis", level=1)
    normality_rows = []
    for label, series in [("Pretest", paired.get("pretest", [])), ("Postest", paired.get("posttest", [])), ("Diferencia post-pre", paired.get("gain", []))]:
        normal = _shapiro(series)
        normality_rows.append([label, normal["n"], _number(normal.get("statistic"), 3), _pvalue(normal.get("p")), normal["status"]])
    _add_table(doc, ["Variable", "n", "W", "p", "Conclusión (α=0,05)"], normality_rows)
    _add_note(doc, "Supuesto principal", "En una t pareada la normalidad corresponde a las diferencias, no a los puntajes pretest y postest considerados por separado. Con muestras pequeñas, Shapiro-Wilk tiene poca potencia; por eso se incluye Q-Q y Wilcoxon como sensibilidad.")

    test_rows = []
    for test in stat_result.get("tests", []):
        decision = "Rechazar H0" if np.isfinite(test.get("p_greater", np.nan)) and test["p_greater"] < ALPHA else "No rechazar H0"
        test_rows.append([test["name"], _number(test.get("statistic"), 3), _pvalue(test.get("p_two")), _pvalue(test.get("p_greater")), decision])
    if test_rows:
        _add_table(doc, ["Prueba", "Estadístico", "p bilateral", "p unilateral H1: post>pre", "Decisión"], test_rows)
    else:
        doc.add_paragraph("No hay tamaño muestral suficiente para ejecutar contrastes.")
    effect = stat_result.get("effect", {})
    if effect:
        _add_table(doc, ["Estimador", "Resultado", "Interpretación"], [
            ["Diferencia media", _number(effect.get("mean_diff"), suffix=" puntos"), "Cambio promedio postest - pretest"],
            ["IC 95% paramétrico", f"{_number(effect.get('ci_low'))} a {_number(effect.get('ci_high'))}", "Incertidumbre de la diferencia media"],
            ["IC 95% bootstrap", f"{_number(effect.get('bootstrap_low'))} a {_number(effect.get('bootstrap_high'))}", "Estimación robusta por remuestreo pareado"],
            ["Cohen d_z", _number(effect.get("cohen_dz")), _effect_label(effect.get("cohen_dz", np.nan)).capitalize()],
            ["Hedges g_z", _number(effect.get("hedges_gz")), "Corrección del efecto para muestras pequeñas"],
            ["Correlación biserial de rangos", _number(effect.get("rank_biserial")), "Tamaño de efecto no paramétrico con dirección"],
        ])
    _add_note(doc, "Decisión inferencial", _primary_interpretation(stat_result))

    if len(paired) >= 3 and paired["pretest"].nunique() > 1 and paired["posttest"].nunique() > 1:
        pearson = stats.pearsonr(paired["pretest"], paired["posttest"])
        spearman = stats.spearmanr(paired["pretest"], paired["posttest"])
        _add_table(doc, ["Asociación", "Coeficiente", "p", "Uso"], [
            ["Pearson", _number(pearson.statistic, 3), _pvalue(pearson.pvalue), "Relación lineal entre mediciones"],
            ["Spearman", _number(spearman.statistic, 3), _pvalue(spearman.pvalue), "Relación monotónica robusta"],
        ])
        _add_note(doc, "Precaución", "La correlación pretest-postest informa estabilidad relativa entre estudiantes; no mide la magnitud de la mejora ni demuestra efecto causal.")

    if not latest.empty and "passed" in latest.columns:
        passed_data = latest.assign(_passed=_bool_series(latest["passed"])).pivot_table(index="user_id", columns="quiz_type", values="_passed", aggfunc="last").dropna()
        if {"pretest", "posttest"}.issubset(passed_data.columns) and len(passed_data):
            b = int(((passed_data["pretest"] == True) & (passed_data["posttest"] == False)).sum())
            c = int(((passed_data["pretest"] == False) & (passed_data["posttest"] == True)).sum())
            discordant = b + c
            mcnemar_p = stats.binomtest(c, discordant, 0.5, alternative="greater").pvalue if discordant else 1.0
            _add_table(doc, ["Cambio de aprobación", "n"], [["Aprobó pre y no post", b], ["No aprobó pre y sí post", c], ["Pares discordantes", discordant]])
            _add_note(doc, "McNemar exacta", f"La prueba direccional sobre el cambio de aprobación produjo p = {_pvalue(mcnemar_p)}. Se usa como resultado secundario; el puntaje continuo conserva más información.")

    doc.add_heading("6. Resultados por dimensión de aprendizaje", level=1)
    if dimension_wide.empty:
        doc.add_paragraph("No existen puntajes dimensionales emparejados suficientes.")
    else:
        rows = []
        for _, row in dimension_tests.iterrows():
            decision = "Significativa" if np.isfinite(row.get("p Holm", np.nan)) and row["p Holm"] < ALPHA else "No significativa"
            rows.append([row["Dimensión"], int(row["n"]), _number(row["Pre M"]), _number(row["Post M"]), _number(row["Ganancia M"]), row["Prueba"], _pvalue(row["p Holm"]), decision])
        _add_table(doc, ["Dimensión", "n", "Pre M", "Post M", "Ganancia", "Prueba", "p Holm", "Resultado"], rows, font_size=7.5)
        _add_figure(doc, _dimension_chart(dimension_wide), figure_no, "Comparación por dimensiones", "Compara comprensión conceptual, procedimientos, aplicaciones y resolución de problemas. La corrección de Holm controla el aumento del error tipo I al evaluar varias dimensiones.")
        figure_no += 1

    doc.add_heading("7. Encuesta final de percepción", level=1)
    if survey["items"].empty:
        doc.add_paragraph("No existen respuestas de encuesta para los filtros seleccionados.")
    else:
        alpha = survey["alpha"]
        if np.isfinite(alpha):
            alpha_label = "excelente" if alpha >= 0.90 else "buena" if alpha >= 0.80 else "aceptable" if alpha >= 0.70 else "cuestionable" if alpha >= 0.60 else "baja"
            _add_note(doc, "Consistencia interna", f"Alfa de Cronbach = {_number(alpha, 3)} con {survey['complete_n']} participantes completos y {survey['k']} ítems; consistencia {alpha_label}. El alfa debe interpretarse junto con la estructura multidimensional del instrumento.")
        else:
            _add_note(doc, "Consistencia interna", "No estimable: se necesitan al menos dos participantes con respuestas completas y variabilidad en el puntaje total.")
        if not survey["dimensions"].empty:
            survey_rows = []
            for _, row in survey["dimensions"].iterrows():
                result_label = "Superior al neutral" if np.isfinite(row.get("p Holm", np.nan)) and row["p Holm"] < ALPHA else "Sin evidencia suficiente"
                survey_rows.append([row["Dimensión"], int(row["n"]), _number(row["Media"]), _number(row["DE"]), _number(row["Mediana"]), _pvalue(row["p Holm"]), result_label])
            _add_table(doc, ["Dimensión", "n", "Media", "DE", "Mediana", "p Holm (>3)", "Interpretación"], survey_rows, font_size=8)
            _add_figure(doc, _survey_dimension_chart(survey["dimensions"]), figure_no, "Percepción por dimensión", "Las medias superiores a 3 reflejan valoración favorable respecto al punto neutral. La significación se contrasta por estudiante con Wilcoxon y corrección de Holm.")
            figure_no += 1
        _add_figure(doc, _survey_item_chart(survey["items"]), figure_no, "Resultados por ítem de la encuesta", "Permite identificar fortalezas y aspectos específicos por mejorar. Los ítems deben interpretarse en su dimensión y no como pruebas independientes del rendimiento académico.")
        figure_no += 1

    doc.add_heading("8. Ejercicios guiados durante la intervención", level=1)
    if exercises["summary"].empty:
        doc.add_paragraph("No existen intentos de ejercicios guiados para los filtros seleccionados.")
    else:
        exercise_rows = [[row["difficulty_level"], int(row["count"]), int(row["sum"]), f"{row['accuracy_pct']:.1f}%"] for _, row in exercises["summary"].iterrows()]
        _add_table(doc, ["Dificultad", "Intentos", "Aciertos", "Exactitud"], exercise_rows)
        chi2 = exercises.get("chi2")
        if chi2:
            validity = "supuesto cumplido" if chi2["valid"] else "frecuencias esperadas menores de 5; resultado exploratorio"
            _add_note(doc, "Asociación dificultad-acierto", f"χ²({chi2['dof']}) = {_number(chi2['chi2'], 3)}, p = {_pvalue(chi2['p'])}, V de Cramér = {_number(chi2['cramers_v'], 3)}; {validity}.")
        _add_figure(doc, _exercise_chart(exercises["summary"]), figure_no, "Exactitud en ejercicios por dificultad", "Describe el desempeño durante la intervención. Los intentos no son independientes dentro de un mismo estudiante, por lo que el contraste χ² es complementario y no sustituye el análisis pretest-postest.")
        figure_no += 1

    doc.add_heading("9. Conclusiones y criterios de reporte", level=1)
    doc.add_paragraph(_primary_interpretation(stat_result))
    doc.add_paragraph("La conclusión sustantiva debe integrar cuatro componentes: magnitud de la ganancia, intervalo de confianza, valor p y tamaño del efecto. Un valor p significativo por sí solo no demuestra relevancia educativa.")
    doc.add_paragraph("Para el informe final de tesis se recomienda conservar este documento, el CSV original y la fecha de extracción; declarar el criterio de último intento; informar pérdidas de seguimiento; y evitar afirmar causalidad fuerte mientras no exista un grupo comparador equivalente.")

    doc.add_heading("Anexo A. Diccionario de variables analíticas", level=1)
    variable_rows = [
        ["user_id", "Identificador", "Vincula de forma anónima las mediciones del estudiante"],
        ["quiz_type", "Categórica", "Pretest o postest"],
        ["score", "Cuantitativa", "Puntaje de rendimiento, escala 0-100"],
        ["gain", "Cuantitativa derivada", "Postest menos pretest"],
        ["dimension_scores_json", "Cuantitativa multidimensional", "Puntajes por competencia"],
        ["total_time_seconds", "Cuantitativa", "Tiempo total de evaluación"],
        ["version_code", "Categórica", "Versión aleatoria del instrumento"],
        ["cognitive_profile", "Ordinal/categórica", "Perfil estimado por el sistema"],
        ["is_correct", "Dicotómica", "Resultado de cada ejercicio guiado"],
        ["difficulty_level", "Ordinal", "Dificultad del ejercicio"],
        ["survey score", "Ordinal", "Respuesta Likert de 1 a 5"],
    ]
    _add_table(doc, ["Variable", "Tipo", "Uso"], variable_rows)

    doc.add_heading("Anexo B. Nota de reproducibilidad", level=1)
    doc.add_paragraph("El informe fue generado automáticamente con Python, pandas, SciPy, Matplotlib y python-docx. Las pruebas usan α=0,05; el remuestreo bootstrap utiliza una semilla fija; las comparaciones múltiples emplean Holm; y los gráficos se producen con los datos filtrados al momento de generar el archivo.")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def _eda_row_values(desc):
    return [
        desc.get("n", 0),
        _number(desc.get("mean")),
        _number(desc.get("sd")),
        _number(desc.get("median")),
        f"{_number(desc.get('q1'))} - {_number(desc.get('q3'))}",
        f"{_number(desc.get('min'))} - {_number(desc.get('max'))}",
        _number(desc.get("skew")),
    ]
