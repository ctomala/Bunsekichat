import os, re, json, hashlib, time, threading, random, unicodedata
from datetime import datetime, date, timedelta
from io import BytesIO

import pandas as pd
import plotly.express as px
import streamlit as st
import psycopg2
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv
from google import genai

# Exportaciones opcionales: instalar python-docx y reportlab para Word/PDF.
try:
    from docx import Document
except Exception:
    Document = None
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:
    A4 = None
    canvas = None

# GPS automático desde el navegador.
# Instalar si se desea captura automática real: pip install streamlit-js-eval
try:
    from streamlit_js_eval import get_geolocation
except Exception:
    get_geolocation = None

# =========================================================
# CONFIGURACIÓN GENERAL
# =========================================================
# Carga el archivo .env desde la misma carpeta donde está app.py
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ENV_PATH = os.path.join(BASE_DIR, ".env")
load_dotenv(dotenv_path=ENV_PATH)

APP_NAME = "BunsekiChat"
# Supabase/PostgreSQL: usar DATABASE_URL en Streamlit Secrets o .env
try:
    DATABASE_URL = st.secrets["DATABASE_URL"]
except Exception:
    DATABASE_URL = os.getenv("DATABASE_URL", "").strip()

DB_LOCK = threading.RLock()
SESSION_TIMEOUT_MINUTES = int(os.getenv("SESSION_TIMEOUT_MINUTES", "20"))
# Lee la API desde .env. Para pruebas locales, puedes pegar una clave temporal en GEMINI_API_KEY_FALLBACK.
# IMPORTANTE: no subas claves reales a internet ni a GitHub.
GEMINI_API_KEY_FALLBACK = ""

# Compatible con local (.env) y Streamlit Cloud (Secrets)
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
except Exception:
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip() or GEMINI_API_KEY_FALLBACK.strip()
ADMIN_USER = os.getenv("ADMIN_USER", "admin")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin123")

TOPICS = {
    "Álgebra lineal": ["Vectores", "Matrices", "Determinantes", "Espacios vectoriales", "Autovalores y autovectores"],
    "Estadística": ["Estadística descriptiva", "Probabilidad", "Distribuciones", "Inferencia", "Regresión"],
    "Cálculo diferencial": ["Funciones", "Límites", "Continuidad", "Derivadas", "Aplicaciones de la derivada"],
    "Cálculo integral": ["Antiderivadas", "Integral definida", "Técnicas de integración", "Aplicaciones de la integral", "Series básicas"],
}
LEVELS = ["Inicial", "Básico", "Intermedio", "Avanzado"]


def env_list(name, default):
    raw = os.getenv(name, "").strip()
    if not raw:
        return default
    values = [x.strip() for x in raw.split(",") if x.strip()]
    return values or default


def env_int(name, default):
    try:
        return int(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        return default


ACADEMIC_SUBJECTS = env_list("BUNSEKI_SUBJECTS", [
    "Cálculo diferencial",
    "Cálculo integral",
    "Álgebra lineal",
    "Estadística",
])
ACADEMIC_COURSE_LEVELS = env_list("BUNSEKI_COURSE_LEVELS", ["4", "1", "2", "3", "5", "6", "7", "8"])
ACADEMIC_PARALLELS = env_list("BUNSEKI_PARALLELS", ["4-A1", "4-A2", "4-B1", "4-B2"])
ACADEMIC_SHIFTS = env_list("BUNSEKI_SHIFTS", ["Matutino", "Vespertino", "Nocturno"])
RESEARCH_TITLE = os.getenv(
    "BUNSEKI_RESEARCH_TITLE",
    "Efectos de un Tutor Inteligente basado en Inteligencia Artificial Generativa sobre el Rendimiento Académico en Cálculo Diferencial en la Universidad de Guayaquil",
)
RESEARCH_QUESTION_COUNT = env_int("BUNSEKI_RESEARCH_QUESTION_COUNT", 20)
RESEARCH_TEST_TYPES = {
    "adaptive": "Evaluación adaptativa",
    "pretest": "Pretest",
    "posttest": "Posttest",
}

# Banco local de respaldo para que la investigación funcione aun si la API IA no responde.
DIFFERENTIAL_CALCULUS_QUESTION_BANK = [
    {"question": "Si f(x)=3x^2-5x+2, ¿cuál es f'(x)?", "options": ["6x-5", "3x-5", "6x+2", "x^3-5x"], "correct_answer": "6x-5", "explanation": "Se aplica la regla de la potencia término a término.", "bloom_level": "Aplicar", "topic": "Derivadas"},
    {"question": "El límite de (x^2-9)/(x-3) cuando x tiende a 3 es:", "options": ["6", "3", "0", "No existe"], "correct_answer": "6", "explanation": "Se factoriza x^2-9=(x-3)(x+3) y queda x+3.", "bloom_level": "Aplicar", "topic": "Límites"},
    {"question": "La derivada de sen(x) es:", "options": ["cos(x)", "-sen(x)", "-cos(x)", "tan(x)"], "correct_answer": "cos(x)", "explanation": "Es una derivada trigonométrica fundamental.", "bloom_level": "Recordar", "topic": "Derivadas trigonométricas"},
    {"question": "Para que una función sea continua en x=a debe cumplirse que:", "options": ["f(a) existe, el límite existe y ambos son iguales", "Solo exista f(a)", "Solo exista el límite lateral derecho", "La derivada sea cero"], "correct_answer": "f(a) existe, el límite existe y ambos son iguales", "explanation": "La continuidad exige existencia de valor, existencia de límite y coincidencia.", "bloom_level": "Comprender", "topic": "Continuidad"},
    {"question": "Si f'(a)>0, entonces cerca de a la función tiende a:", "options": ["Crecer", "Decrecer", "Ser constante siempre", "No ser continua"], "correct_answer": "Crecer", "explanation": "Una derivada positiva indica tasa de cambio positiva.", "bloom_level": "Comprender", "topic": "Aplicaciones de la derivada"},
    {"question": "La derivada de ln(x), para x>0, es:", "options": ["1/x", "x", "ln(x)", "e^x"], "correct_answer": "1/x", "explanation": "Es la regla básica de la función logaritmo natural.", "bloom_level": "Recordar", "topic": "Derivadas logarítmicas"},
    {"question": "¿Cuál es la derivada de (x^2+1)^5?", "options": ["10x(x^2+1)^4", "5(x^2+1)^4", "2x(x^2+1)^5", "10(x^2+1)^4"], "correct_answer": "10x(x^2+1)^4", "explanation": "Se usa regla de la cadena: 5(x^2+1)^4 por 2x.", "bloom_level": "Aplicar", "topic": "Regla de la cadena"},
    {"question": "La pendiente de la recta tangente a una curva en un punto se obtiene con:", "options": ["La derivada en ese punto", "El área bajo la curva", "La ordenada al origen", "El promedio de los valores de y"], "correct_answer": "La derivada en ese punto", "explanation": "La derivada representa la tasa instantánea y la pendiente tangente.", "bloom_level": "Comprender", "topic": "Interpretación de la derivada"},
    {"question": "El límite de sen(x)/x cuando x tiende a 0 es:", "options": ["1", "0", "Infinito", "No existe"], "correct_answer": "1", "explanation": "Es un límite trigonométrico notable.", "bloom_level": "Recordar", "topic": "Límites trigonométricos"},
    {"question": "Si f''(x)>0 en un intervalo, la gráfica es:", "options": ["Cóncava hacia arriba", "Cóncava hacia abajo", "Siempre decreciente", "Discontinua"], "correct_answer": "Cóncava hacia arriba", "explanation": "La segunda derivada positiva indica concavidad hacia arriba.", "bloom_level": "Comprender", "topic": "Concavidad"},
    {"question": "Un punto crítico puede aparecer cuando:", "options": ["f'(x)=0 o f'(x) no existe", "f(x)=0 siempre", "x es negativo", "La función es lineal"], "correct_answer": "f'(x)=0 o f'(x) no existe", "explanation": "Los candidatos a extremos locales se buscan donde la derivada se anula o no existe.", "bloom_level": "Comprender", "topic": "Puntos críticos"},
    {"question": "La derivada de e^(2x) es:", "options": ["2e^(2x)", "e^(2x)", "2x e^(2x)", "e^2"], "correct_answer": "2e^(2x)", "explanation": "Se aplica regla de la cadena a la función exponencial.", "bloom_level": "Aplicar", "topic": "Derivadas exponenciales"},
    {"question": "La regla del producto para derivar u(x)v(x) es:", "options": ["u'v+uv'", "u'v'", "u/v", "u'-v'"], "correct_answer": "u'v+uv'", "explanation": "La derivada del producto combina la derivada de cada factor.", "bloom_level": "Recordar", "topic": "Regla del producto"},
    {"question": "Para hallar máximos y mínimos absolutos en un intervalo cerrado se revisan:", "options": ["Puntos críticos y extremos del intervalo", "Solo puntos donde f(x)=0", "Solo valores negativos", "Solo el punto medio"], "correct_answer": "Puntos críticos y extremos del intervalo", "explanation": "El teorema del valor extremo exige comparar candidatos internos y extremos.", "bloom_level": "Aplicar", "topic": "Optimización"},
    {"question": "La linealización de f en x=a tiene la forma:", "options": ["L(x)=f(a)+f'(a)(x-a)", "L(x)=f'(a)+a(x-a)", "L(x)=f(a)(x-a)", "L(x)=f''(a)"], "correct_answer": "L(x)=f(a)+f'(a)(x-a)", "explanation": "La aproximación lineal usa el valor y la pendiente en a.", "bloom_level": "Aplicar", "topic": "Aproximación lineal"},
    {"question": "En la ecuación x^2+y^2=25, la derivada implícita dy/dx es:", "options": ["-x/y", "x/y", "-y/x", "2x+2y"], "correct_answer": "-x/y", "explanation": "Derivando: 2x+2y y'=0, por tanto y'=-x/y.", "bloom_level": "Aplicar", "topic": "Derivación implícita"},
    {"question": "La función f(x)=|x| no es derivable en x=0 porque:", "options": ["Las derivadas laterales son diferentes", "No existe f(0)", "No es continua en 0", "Su dominio es vacío"], "correct_answer": "Las derivadas laterales son diferentes", "explanation": "En x=0 la pendiente lateral izquierda y derecha no coinciden.", "bloom_level": "Analizar", "topic": "Derivabilidad"},
    {"question": "El teorema del valor medio requiere que f sea:", "options": ["Continua en [a,b] y derivable en (a,b)", "Solo derivable en a", "Discontinua en [a,b]", "Constante en todo su dominio"], "correct_answer": "Continua en [a,b] y derivable en (a,b)", "explanation": "Son las condiciones necesarias para aplicar el teorema.", "bloom_level": "Comprender", "topic": "Teorema del valor medio"},
    {"question": "La derivada de tan(x) es:", "options": ["sec^2(x)", "cos^2(x)", "-csc^2(x)", "sen(x)"], "correct_answer": "sec^2(x)", "explanation": "Es una identidad básica de derivación trigonométrica.", "bloom_level": "Recordar", "topic": "Derivadas trigonométricas"},
    {"question": "La derivada de una constante c es:", "options": ["0", "c", "1", "x"], "correct_answer": "0", "explanation": "Una constante no cambia respecto de x.", "bloom_level": "Recordar", "topic": "Reglas básicas de derivación"},
    {"question": "La regla de la potencia establece que la derivada de x^n es:", "options": ["nx^(n-1)", "x^(n+1)", "n/x", "x^n"], "correct_answer": "nx^(n-1)", "explanation": "El exponente baja como coeficiente y se reduce en una unidad.", "bloom_level": "Recordar", "topic": "Regla de la potencia"},
    {"question": "Si f'(x)=2x, entonces f crece en:", "options": ["(0, infinito)", "(-infinito, 0)", "Todo R", "Ningún intervalo"], "correct_answer": "(0, infinito)", "explanation": "La función crece donde la derivada es positiva.", "bloom_level": "Analizar", "topic": "Monotonía"},
    {"question": "Un límite bilateral existe cuando:", "options": ["Los límites laterales izquierdo y derecho son iguales", "Solo existe el límite izquierdo", "La función no está definida", "La derivada es cero"], "correct_answer": "Los límites laterales izquierdo y derecho son iguales", "explanation": "La igualdad de límites laterales garantiza el límite bilateral.", "bloom_level": "Comprender", "topic": "Límites laterales"},
    {"question": "La derivada de cos(3x) es:", "options": ["-3sen(3x)", "-sen(3x)", "3cos(3x)", "sen(3x)"], "correct_answer": "-3sen(3x)", "explanation": "Se deriva cos(u) como -sen(u) y se multiplica por u'=3.", "bloom_level": "Aplicar", "topic": "Regla de la cadena"},
    {"question": "La regla de L'Hôpital puede aplicarse inicialmente ante formas:", "options": ["0/0 o infinito/infinito", "Solo 1/0", "Solo 0+0", "Cualquier suma"], "correct_answer": "0/0 o infinito/infinito", "explanation": "La regla se usa bajo condiciones de indeterminación de cocientes.", "bloom_level": "Comprender", "topic": "Límites indeterminados"},
    {"question": "En problemas de razones relacionadas, el primer paso recomendable es:", "options": ["Definir variables y una ecuación que las relacione", "Derivar sin leer el enunciado", "Elegir una respuesta al azar", "Eliminar todas las unidades"], "correct_answer": "Definir variables y una ecuación que las relacione", "explanation": "La modelación correcta permite derivar respecto del tiempo.", "bloom_level": "Aplicar", "topic": "Razones relacionadas"},
    {"question": "Si f''(x)<0 en un intervalo, la gráfica es:", "options": ["Cóncava hacia abajo", "Cóncava hacia arriba", "Siempre creciente", "Lineal horizontal"], "correct_answer": "Cóncava hacia abajo", "explanation": "La segunda derivada negativa indica concavidad hacia abajo.", "bloom_level": "Comprender", "topic": "Concavidad"},
    {"question": "Un máximo local suele identificarse cuando f'(x) cambia de:", "options": ["Positiva a negativa", "Negativa a positiva", "Cero a cero siempre", "No definida a positiva siempre"], "correct_answer": "Positiva a negativa", "explanation": "El cambio de crecimiento a decrecimiento indica máximo local.", "bloom_level": "Analizar", "topic": "Criterio de la primera derivada"},
    {"question": "Si f(x)=x^3, entonces f'(2) es:", "options": ["12", "6", "8", "3"], "correct_answer": "12", "explanation": "f'(x)=3x^2, por tanto f'(2)=3(4)=12.", "bloom_level": "Aplicar", "topic": "Derivadas"},
]

st.set_page_config(page_title=APP_NAME, page_icon="∑", layout="wide", initial_sidebar_state="expanded")

# Control interno: evita crear varias veces el componente GPS en un mismo rerun de Streamlit.
st.session_state['_gps_widget_rendered_this_run'] = False

CSS = """
<style>
:root{
    --magenta:#d4147f;
    --magenta-2:#ff2f93;
    --magenta-dark:#6f0f49;
    --magenta-soft:#ffe4f2;
    --bg:#eef0f3;
    --card:#ffffff;
    --line:#d9dde3;
    --text:#1f2937;
    --muted:#6b7280;
    --shadow:0 22px 55px rgba(31,41,55,.13);
}

header, footer, #MainMenu {visibility:hidden; height:0 !important;}
div[data-testid="stToolbar"], div[data-testid="stDecoration"], div[data-testid="stStatusWidget"]{display:none !important;}
html, body, [class*="css"]{font-family:"Segoe UI", Tahoma, Geneva, Verdana, sans-serif;}
.stApp{background:linear-gradient(135deg,#eef0f3 0%,#e5e7eb 55%,#f6dce9 100%); color:var(--text);}
.block-container{max-width:1180px; padding-top:0 !important; padding-bottom:1.1rem !important;}
section.main > div{padding-top:0 !important;}
.element-container:empty{display:none !important;}
/* Corrección de scroll y espacio superior del login */
section.main > div.block-container{padding-top:0!important;}
section.main div[data-testid="stVerticalBlock"]{gap:0.35rem;}
.login-premium-page + div{display:none!important;}


/* ===================== LOGIN PREMIUM ===================== */
.login-premium-page{min-height:auto; display:block; padding:34px 0 14px; margin:0;}
.ad-panel{
    min-height:300px;
    border:1px dashed #cbd1da;
    border-radius:26px;
    background:rgba(255,255,255,.52);
    display:flex;
    align-items:center;
    justify-content:center;
    text-align:center;
    padding:18px;
    color:#8a94a3;
    box-shadow:inset 0 1px 0 rgba(255,255,255,.8);
}
.ad-panel strong{display:block;color:#5f6875;margin-bottom:6px;font-size:.9rem;}
.ad-panel span{font-size:.78rem;line-height:1.45;}
.login-card-premium{
    width:100%;
    max-width:305px;
    margin:26px auto 0;
    background:rgba(255,255,255,.94);
    border:1px solid rgba(217,221,227,.95);
    border-radius:28px;
    padding:30px 18px 14px;
    box-shadow:var(--shadow);
    text-align:center;
    position:relative;
}
.logo-orb{
    width:58px;height:58px;border-radius:50%;
    background:linear-gradient(135deg,var(--magenta),var(--magenta-2));
    color:white;
    display:flex;align-items:center;justify-content:center;
    font-weight:900;font-size:1.28rem;
    margin:-58px auto 8px;
    border:4px solid rgba(255,255,255,.9);
    box-shadow:0 14px 32px rgba(212,20,127,.32);
}
.login-title{font-size:1.28rem;font-weight:900;color:var(--magenta-dark);line-height:1;margin:0;}
.login-subtitle{font-size:.78rem;color:var(--muted);margin:6px 0 10px;}
.login-note{font-size:.72rem;color:var(--muted);text-align:center;margin-top:8px;}
.powered{font-size:.73rem;color:#6b7280;text-align:center;margin-top:8px;}
.powered strong{color:var(--magenta-dark);}

/* Tabs del login */
.login-card-premium [data-testid="stTabs"]{margin-top:0 !important;}
.login-card-premium [data-testid="stTabs"] button{padding:4px 8px !important;}
.login-card-premium [data-testid="stTabs"] button p{font-size:.80rem !important;font-weight:800 !important;}
.login-card-premium hr{margin:.35rem 0 .55rem !important;}
.login-card-premium div[data-testid="stVerticalBlock"]{gap:.12rem !important;}
.login-card-premium .stTextInput{margin-bottom:.05rem !important;}
.login-card-premium .stButton{margin-top:.18rem !important;}
.login-card-premium .stAlert{font-size:.74rem;padding:.35rem .55rem;}

/* Inputs compactos globales */
.stTextInput label,.stTextArea label{font-size:.76rem;color:var(--text)!important;font-weight:800;margin-bottom:0!important;}
.stTextInput input,.stTextArea textarea{
    min-height:30px!important;
    height:30px!important;
    border-radius:999px!important;
    border:1px solid #d4d8df!important;
    background:#f3f4f6!important;
    color:var(--text)!important;
    font-size:.78rem!important;
    padding:.18rem .72rem!important;
}
.stTextInput input:focus,.stTextArea textarea:focus{border-color:var(--magenta)!important;box-shadow:0 0 0 3px rgba(212,20,127,.11)!important;}
.stButton>button{
    background:linear-gradient(90deg,var(--magenta),var(--magenta-2));
    color:white;border:0;border-radius:999px;
    min-height:32px;height:32px;
    font-size:.76rem;font-weight:900;
    box-shadow:0 12px 24px rgba(212,20,127,.24);
}
.stButton>button:hover{color:white;border:0;filter:brightness(.98);}

/* Marca y resto app */
.bunseki-brand{display:flex;align-items:center;gap:12px;}
.bunseki-logo{width:48px;height:48px;border-radius:16px;background:linear-gradient(135deg,var(--magenta),var(--magenta-2));display:flex;align-items:center;justify-content:center;color:white;font-weight:900;font-size:1.25rem;box-shadow:0 12px 28px rgba(212,20,127,.25);border:1px solid rgba(255,255,255,.45);}
.bunseki-title{font-size:1.45rem;font-weight:900;color:var(--magenta-dark);line-height:1;margin:0;}
.bunseki-subtitle{font-size:.82rem;color:var(--muted);margin-top:5px;}
.card{background:var(--card);border:1px solid var(--line);border-radius:18px;padding:18px;box-shadow:0 10px 26px rgba(31,41,55,.06);margin-bottom:14px;}
.hero{background:linear-gradient(135deg,#ffffff 0%,#fff1f8 100%);border:1px solid #f3c6df;border-radius:22px;padding:20px 22px;box-shadow:0 12px 28px rgba(194,24,122,.08);margin-bottom:16px;}
.badge{display:inline-block;background:#fce7f3;color:var(--magenta-dark);padding:6px 12px;border-radius:999px;font-weight:800;border:1px solid #f5b8d5;font-size:.82rem;}
.small{color:var(--muted);font-size:.86rem;}
.user{background:#fff7fb;border-left:5px solid var(--magenta);border-radius:14px;padding:12px;margin:10px 0;}
.bot{background:#ffffff;border-left:5px solid #9ca3af;border-radius:14px;padding:12px;margin:10px 0;border:1px solid var(--line);}
.metric{background:#fff;border-radius:16px;padding:14px;border:1px solid var(--line);text-align:center;box-shadow:0 8px 20px rgba(31,41,55,.04);}
.metric b{font-size:1.25rem;color:var(--magenta-dark);}
.metric span{color:var(--muted);font-size:.84rem;}
section[data-testid="stSidebar"]{background:#ffffff;border-right:1px solid var(--line);}
section[data-testid="stSidebar"] h1,section[data-testid="stSidebar"] h2,section[data-testid="stSidebar"] h3{color:var(--magenta-dark)!important;}
[data-testid="stDataFrame"]{border-radius:14px;overflow:hidden;}

/* Responsive celular */
@media(max-width:980px){
    .login-premium-page{min-height:auto;display:block;padding:24px 0 12px;}
    .ad-panel{display:none;}
    .login-card-premium{max-width:305px;margin:22px auto 0;padding:30px 17px 14px;}
    .block-container{padding-left:.85rem!important;padding-right:.85rem!important;}
}
@media(max-width:430px){
    .login-card-premium{max-width:92vw;border-radius:24px;}
    .login-title{font-size:1.24rem;}
    .login-subtitle{font-size:.74rem;}
    .logo-orb{width:54px;height:54px;font-size:1.18rem;margin-top:-60px;}
    .stTextInput input{height:29px!important;min-height:29px!important;font-size:.76rem!important;}
    .stButton>button{height:31px;min-height:31px;}
}

/* ===================== AULA POST-LOGIN PREMIUM ===================== */
.hero{position:relative;overflow:hidden;background:linear-gradient(135deg,#ffffff 0%,#fff7fb 55%,#fde7f3 100%)!important;border:1px solid #f4b8d5!important;box-shadow:0 18px 42px rgba(111,15,73,.10)!important;}
.hero::after{content:"";position:absolute;right:-55px;top:-55px;width:180px;height:180px;background:radial-gradient(circle,rgba(212,20,127,.16),rgba(212,20,127,0) 70%);}
.hero .bunseki-brand{position:relative;z-index:2;}
section[data-testid="stSidebar"]{background:linear-gradient(180deg,#ffffff 0%,#fff7fb 58%,#f3f4f6 100%)!important;border-right:1px solid #f1b8d4!important;box-shadow:12px 0 30px rgba(111,15,73,.06)!important;}
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"]{gap:.7rem!important;}
section[data-testid="stSidebar"] .stSelectbox{background:#ffffff;border:1px solid #f1d1e3;border-radius:15px;padding:8px 10px;box-shadow:0 8px 18px rgba(31,41,55,.04);}
section[data-testid="stSidebar"] .stButton>button{width:100%;}
.metric{border:1px solid #f0c8dc!important;background:linear-gradient(180deg,#ffffff 0%,#fffafd 100%)!important;}
.card{border:1px solid #e8d5e0!important;}
.quiz-card-premium{background:linear-gradient(180deg,#ffffff 0%,#fff7fb 100%);border:1px solid #f0c8dc;border-radius:18px;padding:16px;box-shadow:0 10px 26px rgba(31,41,55,.06);margin-bottom:14px;}
.quiz-card-premium h3{margin-top:0;color:var(--magenta-dark);}
.quiz-topic-chip{display:inline-block;background:#fce7f3;border:1px solid #f5b8d5;color:var(--magenta-dark);font-weight:800;border-radius:999px;padding:4px 10px;font-size:.78rem;margin-bottom:8px;}


/* ===================== SIDEBAR PROFESIONAL TIPO APP ===================== */
section[data-testid="stSidebar"]{
    display:block !important;
    min-width:285px !important;
    background:linear-gradient(180deg,#ffffff 0%,#fff7fb 54%,#f3f4f6 100%)!important;
    border-right:1px solid #f1b8d4!important;
    box-shadow:12px 0 30px rgba(111,15,73,.08)!important;
}
section[data-testid="stSidebar"] [data-testid="stSidebarContent"]{padding:1.05rem .9rem!important;}
.sidebar-brand-pro{background:linear-gradient(135deg,#d4147f,#ff2f93);color:#fff;border-radius:20px;padding:14px 14px;margin-bottom:14px;box-shadow:0 16px 32px rgba(212,20,127,.23)}
.sidebar-brand-pro .sb-title{font-weight:950;font-size:1.15rem;line-height:1}.sidebar-brand-pro .sb-sub{font-size:.76rem;opacity:.92;margin-top:5px}
.sidebar-chip-pro{display:inline-block;background:#fff;color:#6f0f49;border:1px solid #f5b8d5;border-radius:999px;padding:4px 10px;font-size:.74rem;font-weight:900;margin-top:10px}
.sidebar-section-title{font-size:.74rem;text-transform:uppercase;letter-spacing:.08em;color:#8a476e;font-weight:900;margin:12px 0 6px}
section[data-testid="stSidebar"] .stSelectbox{background:#ffffff;border:1px solid #f1d1e3;border-radius:16px;padding:8px 10px;box-shadow:0 8px 18px rgba(31,41,55,.04);}
section[data-testid="stSidebar"] .stButton>button{width:100%;height:36px;min-height:36px;border-radius:999px;background:linear-gradient(90deg,#d4147f,#ff2f93);}
.teacher-card{background:linear-gradient(180deg,#ffffff 0%,#fffafd 100%);border:1px solid #f0c8dc;border-radius:18px;padding:16px;margin:12px 0;box-shadow:0 10px 26px rgba(31,41,55,.06)}
.teacher-card h3{margin:.1rem 0 .35rem;color:#6f0f49}.teacher-chip{display:inline-block;background:#fce7f3;border:1px solid #f5b8d5;color:#6f0f49;font-weight:800;border-radius:999px;padding:4px 10px;font-size:.76rem;margin:4px 6px 2px 0}
.export-grid-note{font-size:.82rem;color:#6b7280;margin-top:4px}
@media(max-width:760px){section[data-testid="stSidebar"]{min-width:250px!important}.block-container{padding-left:.7rem!important;padding-right:.7rem!important}}


.student-export-card{background:linear-gradient(180deg,#ffffff 0%,#fff7fb 100%);border:1px solid #f0c8dc;border-radius:18px;padding:15px;box-shadow:0 10px 26px rgba(31,41,55,.06);margin:12px 0 14px;}
.student-export-card h3{margin:.1rem 0 .35rem;color:#6f0f49}.student-export-card p{margin:.25rem 0 .65rem;color:#6b7280;font-size:.84rem}
section[data-testid="stSidebar"]{display:block!important;visibility:visible!important;}

/* ===================== SIDEBAR SAAS PRO NOTION / CHATGPT ===================== */
[data-testid="collapsedControl"]{
    display:flex !important;
    visibility:visible !important;
    position:fixed !important;
    top:14px !important;
    left:14px !important;
    z-index:999999 !important;
    background:#ffffff !important;
    border:1px solid #f1b8d4 !important;
    border-radius:12px !important;
    box-shadow:0 10px 24px rgba(111,15,73,.16) !important;
}
button[kind="header"]{display:flex !important; visibility:visible !important;}
section[data-testid="stSidebar"]{
    display:block !important;
    visibility:visible !important;
    min-width:292px !important;
    max-width:292px !important;
    transition:all .28s ease-in-out !important;
    background:linear-gradient(180deg,#ffffff 0%,#fff7fb 52%,#f5f6f8 100%)!important;
    border-right:1px solid #f1b8d4!important;
    box-shadow:14px 0 36px rgba(111,15,73,.08)!important;
}
section[data-testid="stSidebar"] [data-testid="stSidebarContent"]{
    padding:1rem .85rem 1.25rem !important;
}
.saas-sidebar-brand{
    background:linear-gradient(135deg,#d4147f 0%,#ff2f93 100%);
    color:#fff;border-radius:22px;padding:16px 15px;margin:4px 0 14px;
    box-shadow:0 18px 36px rgba(212,20,127,.26);
}
.saas-sidebar-brand .logo-row{display:flex;align-items:center;gap:10px;}
.saas-sidebar-brand .logo-mini{width:40px;height:40px;border-radius:13px;background:rgba(255,255,255,.18);display:flex;align-items:center;justify-content:center;font-weight:950;border:1px solid rgba(255,255,255,.35)}
.saas-sidebar-brand .title{font-weight:950;font-size:1.05rem;line-height:1}
.saas-sidebar-brand .subtitle{font-size:.74rem;opacity:.92;margin-top:4px}
.saas-role-chip{display:inline-block;background:#fff;color:#6f0f49;border:1px solid #f5b8d5;border-radius:999px;padding:4px 10px;font-size:.72rem;font-weight:900;margin-top:12px}
.saas-section-title{font-size:.70rem;text-transform:uppercase;letter-spacing:.10em;color:#8a476e;font-weight:950;margin:13px 0 7px}
.saas-help-card{background:#fff;border:1px solid #f1d1e3;border-radius:16px;padding:11px 12px;margin:10px 0;color:#6b7280;font-size:.78rem;box-shadow:0 8px 18px rgba(31,41,55,.04)}
section[data-testid="stSidebar"] div[role="radiogroup"] label{
    background:#fff !important;
    border:1px solid #f1d1e3 !important;
    border-radius:14px !important;
    padding:7px 10px !important;
    margin:3px 0 !important;
    box-shadow:0 7px 15px rgba(31,41,55,.035);
}
section[data-testid="stSidebar"] div[role="radiogroup"] label:hover{
    border-color:#d4147f !important;
    background:#fff7fb !important;
}
section[data-testid="stSidebar"] .stSelectbox,
section[data-testid="stSidebar"] .stNumberInput{
    background:#fff;border:1px solid #f1d1e3;border-radius:16px;padding:8px 10px;margin-bottom:8px;box-shadow:0 8px 18px rgba(31,41,55,.04);
}
section[data-testid="stSidebar"] .stButton>button{
    width:100%;height:36px;min-height:36px;border-radius:999px;background:linear-gradient(90deg,#d4147f,#ff2f93);
}
@media(max-width:760px){
    section[data-testid="stSidebar"]{min-width:265px!important;max-width:265px!important;}
    [data-testid="collapsedControl"]{top:8px!important;left:8px!important;}
    .block-container{padding-left:.7rem!important;padding-right:.7rem!important;}
}

</style>
"""
st.markdown(CSS, unsafe_allow_html=True)
st.markdown("""
<style>

/* =========================================
   FIX QUIZ (AZUL MARINO BAJO)
========================================= */
.quiz-card-premium{
    background:linear-gradient(180deg,#1e293b 0%, #0f172a 100%) !important;
    border:1px solid #334155 !important;
    color:#e5e7eb !important;
}

.quiz-card-premium h3,
.quiz-card-premium p,
.quiz-card-premium span{
    color:#e5e7eb !important;
}

.quiz-topic-chip{
    background:#334155 !important;
    border:1px solid #475569 !important;
    color:#e2e8f0 !important;
}

/* preguntas del quiz */
div[data-testid="stRadio"] label{
    color:#e5e7eb !important;
}

/* opciones */
div[role="radiogroup"] label{
    background:#0f172a !important;
    border:1px solid #334155 !important;
    color:#e5e7eb !important;
}

div[role="radiogroup"] label:hover{
    background:#1e293b !important;
    border-color:#ff2f93 !important;
}

div[role="radiogroup"] input:checked + div{
    color:#ff2f93 !important;
}


/* =========================================
   FIX CONFIGURACIÓN ACADÉMICA (PLOMO OSCURO)
========================================= */
section[data-testid="stSidebar"] .stSelectbox{
    background:#ffe4f2 !important; /* plomo oscuro */
    border:1px solid #374151 !important;
    border-radius:14px !important;
}

/* texto dentro del select */
section[data-testid="stSidebar"] .stSelectbox div,
section[data-testid="stSidebar"] .stSelectbox span{
    color:#ffe4f2 !important;
}

/* dropdown abierto */
div[data-baseweb="select"]{
    background:#1f2937 !important;
    color:#ffe4f2 !important;
}

/* inputs dentro del sidebar */
section[data-testid="stSidebar"] input{
    background:#1f2937 !important;
    color:#e5e7eb !important;
    border:1px solid #374151 !important;
}


/* =========================================
   EXTRA: mejora contraste general sidebar
========================================= */
section[data-testid="stSidebar"]{
    background:linear-gradient(180deg,#1f2937 0%, #111827 100%) !important;
}

section[data-testid="stSidebar"] *{
    color:#ffe4f2 !important;
}

</style>
""", unsafe_allow_html=True)
# =========================================================
# NIVEL PRO TOTAL: FIX FINAL DE SIDEBAR, CONTRASTE Y BOTONES
# =========================================================
st.markdown("""
<style>
/* ---------- Base visual limpia ---------- */
html, body, .stApp, [data-testid="stAppViewContainer"]{
    color:#111827 !important;
}
.block-container{
    padding-top:.75rem !important;
}

/* ---------- Botón nativo de abrir/cerrar sidebar MUY visible ---------- */
[data-testid="collapsedControl"]{
    display:flex !important;
    visibility:visible !important;
    opacity:1 !important;
    position:fixed !important;
    top:14px !important;
    left:14px !important;
    z-index:2147483647 !important;
    width:46px !important;
    height:46px !important;
    min-width:46px !important;
    min-height:46px !important;
    align-items:center !important;
    justify-content:center !important;
    background:linear-gradient(135deg,#d4147f,#ff2f93) !important;
    border:2px solid #ffffff !important;
    border-radius:16px !important;
    box-shadow:0 14px 30px rgba(212,20,127,.45) !important;
}
[data-testid="collapsedControl"] svg,
[data-testid="collapsedControl"] svg path{
    fill:#ffffff !important;
    stroke:#ffffff !important;
    color:#ffffff !important;
}
[data-testid="collapsedControl"]::after{
    content:"Menú";
    position:absolute;
    left:54px;
    top:8px;
    background:#ffffff;
    color:#6f0f49;
    border:1px solid #f1b8d4;
    border-radius:999px;
    padding:5px 10px;
    font-size:.74rem;
    font-weight:900;
    white-space:nowrap;
    box-shadow:0 8px 18px rgba(31,41,55,.12);
}
button[kind="header"]{
    display:flex !important;
    visibility:visible !important;
    opacity:1 !important;
}

/* ---------- Sidebar tipo Notion / ChatGPT: contraste fuerte ---------- */
section[data-testid="stSidebar"]{
    display:block !important;
    visibility:visible !important;
    background:linear-gradient(180deg,#ffffff 0%,#fff7fb 56%,#f4f5f7 100%) !important;
    border-right:1px solid #f1b8d4 !important;
    box-shadow:14px 0 34px rgba(111,15,73,.10) !important;
    min-width:292px !important;
    max-width:292px !important;
    transition:all .28s ease-in-out !important;
}
section[data-testid="stSidebar"] [data-testid="stSidebarContent"]{
    padding:1rem .85rem 1.25rem !important;
}
section[data-testid="stSidebar"] *,
section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] div{
    color:#1f2937 !important;
}
section[data-testid="stSidebar"] .saas-sidebar-brand,
section[data-testid="stSidebar"] .saas-sidebar-brand *,
section[data-testid="stSidebar"] .sidebar-brand-pro,
section[data-testid="stSidebar"] .sidebar-brand-pro *{
    color:#ffffff !important;
}
section[data-testid="stSidebar"] .saas-role-chip,
section[data-testid="stSidebar"] .sidebar-chip-pro{
    color:#6f0f49 !important;
    background:#ffffff !important;
}
.saas-section-title,
.sidebar-section-title{
    color:#6f0f49 !important;
    font-weight:950 !important;
}
.saas-help-card,
.saas-help-card *,
.sidebar-help-card,
.sidebar-help-card *{
    color:#374151 !important;
}

/* ---------- Menú radio: legible y con selección magenta ---------- */
section[data-testid="stSidebar"] div[role="radiogroup"] label{
    background:#ffffff !important;
    border:1px solid #f1d1e3 !important;
    border-radius:14px !important;
    padding:8px 10px !important;
    margin:4px 0 !important;
    box-shadow:0 7px 15px rgba(31,41,55,.04) !important;
}
section[data-testid="stSidebar"] div[role="radiogroup"] label *{
    color:#111827 !important;
    font-weight:800 !important;
}
section[data-testid="stSidebar"] div[role="radiogroup"] label:hover{
    background:#fff0f7 !important;
    border-color:#d4147f !important;
}
section[data-testid="stSidebar"] div[role="radiogroup"] input:checked + div,
section[data-testid="stSidebar"] div[role="radiogroup"] [aria-checked="true"]{
    color:#d4147f !important;
}

/* ---------- Selectbox y campos en sidebar ---------- */
section[data-testid="stSidebar"] .stSelectbox,
section[data-testid="stSidebar"] .stNumberInput{
    background:#ffffff !important;
    border:1px solid #f1d1e3 !important;
    border-radius:16px !important;
    padding:8px 10px !important;
    margin-bottom:8px !important;
    box-shadow:0 8px 18px rgba(31,41,55,.04) !important;
}
section[data-testid="stSidebar"] input,
section[data-testid="stSidebar"] textarea,
section[data-testid="stSidebar"] [data-baseweb="select"] *,
section[data-testid="stSidebar"] [data-baseweb="input"] *{
    color:#111827 !important;
    -webkit-text-fill-color:#111827 !important;
}

/* ---------- Botones normales y de descarga PRO ---------- */
.stButton > button,
.stDownloadButton > button,
div[data-testid="stDownloadButton"] button{
    background:linear-gradient(90deg,#d4147f,#ff2f93) !important;
    color:#ffffff !important;
    -webkit-text-fill-color:#ffffff !important;
    border:0 !important;
    border-radius:999px !important;
    min-height:38px !important;
    height:38px !important;
    font-weight:900 !important;
    box-shadow:0 12px 24px rgba(212,20,127,.22) !important;
}
.stButton > button *,
.stDownloadButton > button *,
div[data-testid="stDownloadButton"] button *{
    color:#ffffff !important;
    -webkit-text-fill-color:#ffffff !important;
}
.stButton > button:hover,
.stDownloadButton > button:hover,
div[data-testid="stDownloadButton"] button:hover{
    filter:brightness(.96) !important;
    transform:translateY(-1px) !important;
    border:0 !important;
}
.stDownloadButton > button:focus,
div[data-testid="stDownloadButton"] button:focus{
    outline:3px solid rgba(212,20,127,.18) !important;
}

/* ---------- Botones del área de descargas en tarjetas ---------- */
.student-export-card + div .stDownloadButton button,
.teacher-card + div .stDownloadButton button{
    background:#ffffff !important;
    color:#6f0f49 !important;
    -webkit-text-fill-color:#6f0f49 !important;
    border:1px solid #f1b8d4 !important;
    box-shadow:0 8px 18px rgba(111,15,73,.08) !important;
}
.student-export-card + div .stDownloadButton button *,
.teacher-card + div .stDownloadButton button *{
    color:#6f0f49 !important;
    -webkit-text-fill-color:#6f0f49 !important;
}
.student-export-card + div .stDownloadButton button:hover,
.teacher-card + div .stDownloadButton button:hover{
    background:linear-gradient(90deg,#d4147f,#ff2f93) !important;
    color:#ffffff !important;
    -webkit-text-fill-color:#ffffff !important;
}
.student-export-card + div .stDownloadButton button:hover *,
.teacher-card + div .stDownloadButton button:hover *{
    color:#ffffff !important;
    -webkit-text-fill-color:#ffffff !important;
}

/* ---------- Área de contenido: tarjetas premium ---------- */
.hero, .card, .metric, .student-export-card, .quiz-card-premium, .teacher-card{
    backdrop-filter:blur(8px) !important;
}
.metric b{color:#6f0f49 !important;}
.metric span,.small{color:#4b5563 !important;}

/* ---------- Responsive móvil ---------- */
@media(max-width:760px){
    section[data-testid="stSidebar"]{
        min-width:270px !important;
        max-width:270px !important;
    }
    [data-testid="collapsedControl"]{
        top:8px !important;
        left:8px !important;
        width:42px !important;
        height:42px !important;
    }
    [data-testid="collapsedControl"]::after{
        display:none !important;
    }
    .block-container{
        padding-left:.75rem !important;
        padding-right:.75rem !important;
    }
}
</style>
""", unsafe_allow_html=True)



st.markdown("""
<style>
/* ===================== EVALUACIÓN INTELIGENTE ADAPTATIVA ===================== */
.eval-hero{
    background:linear-gradient(135deg,#ffffff 0%,#fff1f8 60%,#fde7f3 100%);
    border:1px solid #f4b8d5;
    border-radius:24px;
    padding:20px 22px;
    box-shadow:0 16px 38px rgba(111,15,73,.10);
    margin:14px 0;
}
.eval-hero h2{margin:.1rem 0 .35rem;color:#6f0f49;font-size:1.55rem;}
.eval-hero p{margin:.25rem 0;color:#4b5563;font-size:.92rem;}
.eval-card{
    background:#ffffff;
    border:1px solid #f1c8dc;
    border-radius:20px;
    padding:16px;
    margin:12px 0;
    box-shadow:0 10px 26px rgba(31,41,55,.06);
}
.eval-card h3{margin:.1rem 0 .45rem;color:#6f0f49;}
.eval-chip{
    display:inline-block;
    background:#fce7f3;
    border:1px solid #f5b8d5;
    color:#6f0f49;
    border-radius:999px;
    padding:5px 11px;
    font-weight:850;
    font-size:.76rem;
    margin:3px 5px 3px 0;
}
.eval-question{
    background:linear-gradient(180deg,#ffffff 0%,#fff7fb 100%);
    border:1px solid #f0c8dc;
    border-radius:18px;
    padding:14px;
    margin:12px 0;
}
.eval-result-pass{
    background:#ecfdf5;border:1px solid #86efac;color:#14532d;
    border-radius:18px;padding:16px;margin:14px 0;font-weight:800;
}
.eval-result-fail{
    background:#fff7ed;border:1px solid #fdba74;color:#7c2d12;
    border-radius:18px;padding:16px;margin:14px 0;font-weight:800;
}
.plan-upload-box{
    background:linear-gradient(135deg,#ffffff 0%,#fff7fb 100%);
    border:1px dashed #d4147f;
    border-radius:20px;
    padding:16px;
    margin:12px 0;
}
</style>
""", unsafe_allow_html=True)

def bunseki_logo_html(title="BunsekiChat", subtitle="Tutor personalizado de matemáticas universitarias"):
    return (
        "<div class='bunseki-brand'>"
        "<div class='bunseki-logo'><span>BΣ</span></div>"
        "<div>"
        f"<div class='bunseki-title'>{title}</div>"
        f"<div class='bunseki-subtitle'>{subtitle}</div>"
        "</div>"
        "</div>"
    )



# ---------------- DB PostgreSQL / Supabase ----------------
def _pg_connect():
    """Conexión PostgreSQL compatible con Supabase."""
    if not DATABASE_URL:
        st.error("Falta DATABASE_URL. Configúralo en Streamlit Cloud → App → Settings → Secrets.")
        st.stop()
    kwargs = {"cursor_factory": RealDictCursor}
    if "sslmode=" not in DATABASE_URL.lower():
        kwargs["sslmode"] = "require"
    c = psycopg2.connect(DATABASE_URL, **kwargs)
    c.autocommit = True
    return c


def conn():
    """Mantiene el nombre conn() para no cambiar la lógica visual de la app."""
    last_error = None
    for attempt in range(5):
        try:
            return _pg_connect()
        except Exception as e:
            last_error = e
            time.sleep(0.4 * (attempt + 1))
    raise last_error


def fetchone(sql, params=None):
    c = conn()
    try:
        with c.cursor() as cur:
            cur.execute(sql, params or ())
            row = cur.fetchone()
            return dict(row) if row else None
    finally:
        c.close()


def fetchall(sql, params=None):
    c = conn()
    try:
        with c.cursor() as cur:
            cur.execute(sql, params or ())
            rows = cur.fetchall()
            return [dict(r) for r in rows]
    finally:
        c.close()


def execute(sql, params=None, returning=False):
    c = conn()
    try:
        with c.cursor() as cur:
            cur.execute(sql, params or ())
            if returning:
                row = cur.fetchone()
                return dict(row) if row else None
            return None
    finally:
        c.close()


def safe_password(password: str) -> bytes:
    password = password or ""
    return hashlib.sha256(password.encode("utf-8")).digest()


def hash_password(password: str) -> str:
    import bcrypt as bcrypt_lib
    return bcrypt_lib.hashpw(safe_password(password), bcrypt_lib.gensalt()).decode("utf-8")


def verify_password(password: str, password_hash: str) -> bool:
    try:
        import bcrypt as bcrypt_lib
        return bcrypt_lib.checkpw(safe_password(password), password_hash.encode("utf-8"))
    except Exception:
        return False


def init_db():
    """Crea el esquema PostgreSQL equivalente al antiguo SQLite."""
    with DB_LOCK:
        c = conn()
        try:
            with c.cursor() as cur:
                cur.execute("""
                CREATE TABLE IF NOT EXISTS users(
                    id SERIAL PRIMARY KEY,
                    username TEXT UNIQUE NOT NULL,
                    password_hash TEXT NOT NULL,
                    role TEXT DEFAULT 'student',
                    active BOOLEAN DEFAULT TRUE,
                    created_at TEXT NOT NULL,
                    last_login TEXT
                );
                CREATE TABLE IF NOT EXISTS profiles(
                    user_id INTEGER PRIMARY KEY REFERENCES users(id) ON DELETE CASCADE,
                    first_names TEXT,
                    last_names TEXT,
                    birth_date TEXT,
                    course TEXT,
                    subject TEXT,
                    course_level TEXT,
                    parallel TEXT,
                    shift TEXT,
                    cohort TEXT,
                    full_name_normalized TEXT,
                    teacher TEXT,
                    phone TEXT,
                    province TEXT,
                    city TEXT,
                    canton TEXT,
                    address TEXT,
                    gps_lat DOUBLE PRECISION,
                    gps_lon DOUBLE PRECISION,
                    gps_accuracy TEXT,
                    level TEXT DEFAULT 'Inicial',
                    profile_completed INTEGER DEFAULT 0,
                    last_gps_lat DOUBLE PRECISION,
                    last_gps_lon DOUBLE PRECISION,
                    last_gps_accuracy DOUBLE PRECISION,
                    last_gps_at TEXT
                );
                CREATE TABLE IF NOT EXISTS interactions(
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                    role TEXT NOT NULL,
                    topic TEXT,
                    subtopic TEXT,
                    level TEXT,
                    message TEXT NOT NULL,
                    clean_message TEXT,
                    model TEXT,
                    tokens_est INTEGER,
                    latency_ms INTEGER,
                    created_at TEXT NOT NULL,
                    ip_info TEXT,
                    gps_lat DOUBLE PRECISION,
                    gps_lon DOUBLE PRECISION,
                    gps_accuracy DOUBLE PRECISION,
                    gps_source TEXT
                );
                CREATE TABLE IF NOT EXISTS location_events(
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                    event_type TEXT NOT NULL,
                    page TEXT,
                    topic TEXT,
                    subtopic TEXT,
                    gps_lat DOUBLE PRECISION,
                    gps_lon DOUBLE PRECISION,
                    gps_accuracy DOUBLE PRECISION,
                    gps_source TEXT,
                    created_at TEXT NOT NULL
                );
                CREATE TABLE IF NOT EXISTS quizzes(
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                    topic TEXT,
                    level_from TEXT,
                    level_to TEXT,
                    score DOUBLE PRECISION,
                    passed INTEGER,
                    answers_json TEXT,
                    created_at TEXT NOT NULL
                );
                CREATE TABLE IF NOT EXISTS analytic_plans(
                    id SERIAL PRIMARY KEY,
                    title TEXT NOT NULL,
                    course TEXT,
                    subject TEXT,
                    course_level TEXT,
                    parallel TEXT,
                    shift TEXT,
                    teacher_id INTEGER REFERENCES users(id) ON DELETE SET NULL,
                    filename TEXT,
                    raw_text TEXT,
                    summary TEXT,
                    created_at TEXT NOT NULL
                );
                CREATE TABLE IF NOT EXISTS plan_topics(
                    id SERIAL PRIMARY KEY,
                    plan_id INTEGER REFERENCES analytic_plans(id) ON DELETE CASCADE,
                    unit_name TEXT,
                    topic TEXT,
                    subtopic TEXT,
                    learning_outcome TEXT,
                    bloom_level TEXT,
                    keywords TEXT
                );
                CREATE TABLE IF NOT EXISTS adaptive_quizzes(
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER NOT NULL REFERENCES users(id) ON DELETE CASCADE,
                    plan_id INTEGER REFERENCES analytic_plans(id) ON DELETE SET NULL,
                    title TEXT,
                    quiz_type TEXT DEFAULT 'adaptive',
                    source_topic TEXT,
                    difficulty TEXT,
                    subject TEXT,
                    course_level TEXT,
                    parallel TEXT,
                    shift TEXT,
                    cohort TEXT,
                    research_title TEXT,
                    question_count INTEGER,
                    random_seed TEXT,
                    score DOUBLE PRECISION,
                    passed BOOLEAN DEFAULT FALSE,
                    attempt_no INTEGER DEFAULT 1,
                    parent_quiz_id INTEGER,
                    status TEXT DEFAULT 'generated',
                    recommendation TEXT,
                    created_at TEXT NOT NULL,
                    completed_at TEXT
                );
                CREATE TABLE IF NOT EXISTS adaptive_questions(
                    id SERIAL PRIMARY KEY,
                    quiz_id INTEGER NOT NULL REFERENCES adaptive_quizzes(id) ON DELETE CASCADE,
                    question TEXT NOT NULL,
                    options_json TEXT,
                    correct_answer TEXT,
                    position INTEGER,
                    topic TEXT,
                    user_answer TEXT,
                    explanation TEXT,
                    bloom_level TEXT,
                    is_correct BOOLEAN
                );
                CREATE TABLE IF NOT EXISTS settings(
                    key TEXT PRIMARY KEY,
                    value TEXT
                );
                CREATE INDEX IF NOT EXISTS idx_interactions_user_id ON interactions(user_id);
                CREATE INDEX IF NOT EXISTS idx_interactions_created_at ON interactions(created_at);
                CREATE INDEX IF NOT EXISTS idx_location_events_user_id ON location_events(user_id);
                CREATE INDEX IF NOT EXISTS idx_quizzes_user_id ON quizzes(user_id);
                CREATE INDEX IF NOT EXISTS idx_plan_topics_plan_id ON plan_topics(plan_id);
                CREATE INDEX IF NOT EXISTS idx_adaptive_quizzes_user_id ON adaptive_quizzes(user_id);
                CREATE INDEX IF NOT EXISTS idx_adaptive_questions_quiz_id ON adaptive_questions(quiz_id);
                """)
                cur.execute("""
                ALTER TABLE IF EXISTS profiles ADD COLUMN IF NOT EXISTS subject TEXT;
                ALTER TABLE IF EXISTS profiles ADD COLUMN IF NOT EXISTS course_level TEXT;
                ALTER TABLE IF EXISTS profiles ADD COLUMN IF NOT EXISTS parallel TEXT;
                ALTER TABLE IF EXISTS profiles ADD COLUMN IF NOT EXISTS shift TEXT;
                ALTER TABLE IF EXISTS profiles ADD COLUMN IF NOT EXISTS cohort TEXT;
                ALTER TABLE IF EXISTS profiles ADD COLUMN IF NOT EXISTS full_name_normalized TEXT;
                ALTER TABLE IF EXISTS analytic_plans ADD COLUMN IF NOT EXISTS subject TEXT;
                ALTER TABLE IF EXISTS analytic_plans ADD COLUMN IF NOT EXISTS course_level TEXT;
                ALTER TABLE IF EXISTS analytic_plans ADD COLUMN IF NOT EXISTS parallel TEXT;
                ALTER TABLE IF EXISTS analytic_plans ADD COLUMN IF NOT EXISTS shift TEXT;
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS quiz_type TEXT DEFAULT 'adaptive';
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS subject TEXT;
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS course_level TEXT;
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS parallel TEXT;
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS shift TEXT;
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS cohort TEXT;
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS research_title TEXT;
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS question_count INTEGER;
                ALTER TABLE IF EXISTS adaptive_quizzes ADD COLUMN IF NOT EXISTS random_seed TEXT;
                ALTER TABLE IF EXISTS adaptive_questions ADD COLUMN IF NOT EXISTS position INTEGER;
                ALTER TABLE IF EXISTS adaptive_questions ADD COLUMN IF NOT EXISTS topic TEXT;
                CREATE INDEX IF NOT EXISTS idx_profiles_cohort ON profiles(cohort);
                CREATE INDEX IF NOT EXISTS idx_adaptive_quizzes_research ON adaptive_quizzes(quiz_type, subject, course_level, parallel, shift);
                """)
                cur.execute("INSERT INTO settings(key,value) VALUES('session_timeout_minutes', %s) ON CONFLICT (key) DO NOTHING", (str(SESSION_TIMEOUT_MINUTES),))
                cur.execute("SELECT id FROM users WHERE username=%s", (ADMIN_USER,))
                if not cur.fetchone():
                    cur.execute("INSERT INTO users(username,password_hash,role,active,created_at) VALUES(%s,%s,%s,%s,%s)", (ADMIN_USER, hash_password(ADMIN_PASSWORD), 'admin', True, now()))
        finally:
            c.close()


@st.cache_resource
def setup_database_once():
    init_db()
    return True


def now(): return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def normalize_spaces(value):
    return re.sub(r"\s+", " ", str(value or "").strip())


def strip_accents(value):
    normalized = unicodedata.normalize("NFKD", str(value or ""))
    return "".join(ch for ch in normalized if not unicodedata.combining(ch))


def normalize_person_name(value):
    value = normalize_spaces(value)
    value = re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ' -]", "", value)
    small_words = {"de", "del", "la", "las", "los", "y", "da", "das", "do", "dos"}
    parts = []
    for token in value.lower().split():
        if token in small_words:
            parts.append(token)
        else:
            parts.append("-".join(piece.capitalize() for piece in token.split("-") if piece))
    return normalize_spaces(" ".join(parts))


def normalize_academic_text(value):
    return normalize_spaces(value).title()


def normalize_course_level(value):
    return normalize_spaces(value).upper()


def normalize_parallel(value):
    return normalize_spaces(value).upper().replace(" ", "")


def normalize_shift(value):
    return normalize_spaces(value).title()


def unique_values(values, normalizer=normalize_academic_text):
    seen, clean = set(), []
    for value in values:
        item = normalizer(value)
        key = strip_accents(item).lower()
        if item and key not in seen:
            clean.append(item)
            seen.add(key)
    return clean


def parse_course_label(course):
    course = normalize_spaces(course)
    if not course:
        return {}
    parts = course.split()
    shift = parts[-1].title() if parts else ""
    code = " ".join(parts[:-1]) if len(parts) > 1 else course
    if "-" in code:
        level, parallel = code.split("-", 1)
        return {
            "course_level": normalize_course_level(level),
            "parallel": normalize_parallel(parallel),
            "shift": normalize_shift(shift),
        }
    return {"course_level": normalize_course_level(code), "parallel": "", "shift": normalize_shift(shift)}


def build_course_label(course_level, parallel, shift):
    code = "-".join([x for x in [normalize_course_level(course_level), normalize_parallel(parallel)] if x])
    shift = normalize_shift(shift)
    return normalize_spaces(f"{code} {shift}") if shift else code


def academic_context_from_profile(prof):
    parsed = parse_course_label(prof.get("course") if prof else "")
    subject = normalize_academic_text((prof or {}).get("subject") or ACADEMIC_SUBJECTS[0])
    course_level = normalize_course_level((prof or {}).get("course_level") or parsed.get("course_level") or ACADEMIC_COURSE_LEVELS[0])
    parallel = normalize_parallel((prof or {}).get("parallel") or parsed.get("parallel") or ACADEMIC_PARALLELS[0])
    shift = normalize_shift((prof or {}).get("shift") or parsed.get("shift") or ACADEMIC_SHIFTS[0])
    cohort = build_course_label(course_level, parallel, shift)
    return {
        "subject": subject,
        "course_level": course_level,
        "parallel": parallel,
        "shift": shift,
        "cohort": cohort,
        "course": cohort,
    }


def academic_selectbox(label, options, current, key, normalizer=normalize_academic_text, placeholder="Escribe el valor"):
    cleaned_options = unique_values(options, normalizer)
    current_clean = normalizer(current)
    choices = cleaned_options + ["Otro..."]
    index = cleaned_options.index(current_clean) if current_clean in cleaned_options else len(cleaned_options)
    choice = st.selectbox(label, choices, index=index, key=key)
    if choice == "Otro...":
        return normalizer(st.text_input(f"{label} personalizada", value=current_clean, key=f"{key}_other", placeholder=placeholder))
    return choice


def birth_date_value(raw):
    try:
        return datetime.strptime(str(raw), "%Y-%m-%d").date()
    except Exception:
        return date(2005, 1, 1)


def get_timeout():
    r = fetchone("SELECT value FROM settings WHERE key=%s", ('session_timeout_minutes',))
    return int(r['value']) if r else SESSION_TIMEOUT_MINUTES


def set_timeout(minutes:int):
    with DB_LOCK:
        execute("INSERT INTO settings(key,value) VALUES('session_timeout_minutes', %s) ON CONFLICT (key) DO UPDATE SET value=EXCLUDED.value", (str(minutes),))


def user_by_username(username):
    return fetchone("SELECT * FROM users WHERE username=%s", (username.strip().lower(),))


def create_user(username, password):
    with DB_LOCK:
        row = execute("INSERT INTO users(username,password_hash,role,active,created_at) VALUES(%s,%s,%s,%s,%s) RETURNING id", (username.strip().lower(), hash_password(password), 'student', True, now()), returning=True)
        uid = row['id']
        execute("INSERT INTO profiles(user_id) VALUES(%s) ON CONFLICT (user_id) DO NOTHING", (uid,))
        return uid


def authenticate(username, password):
    u = user_by_username(username)
    if not u or not u.get('active') or not verify_password(password, u['password_hash']):
        return None
    with DB_LOCK:
        execute("UPDATE users SET last_login=%s WHERE id=%s", (now(), u['id']))
    return u


def get_profile(uid):
    return fetchone("SELECT * FROM profiles WHERE user_id=%s", (uid,)) or {}


def update_profile(uid, data):
    if not data:
        return
    cols = ','.join([f"{k}=%s" for k in data])
    vals = list(data.values()) + [uid]
    with DB_LOCK:
        execute(f"UPDATE profiles SET {cols} WHERE user_id=%s", vals)


def normalize_gps_payload(raw):
    if not raw or not isinstance(raw, dict):
        return None
    coords = raw.get('coords') if isinstance(raw.get('coords'), dict) else raw
    lat = coords.get('latitude', coords.get('lat'))
    lon = coords.get('longitude', coords.get('lon', coords.get('lng')))
    acc = coords.get('accuracy')
    try:
        lat = float(lat); lon = float(lon); acc = float(acc) if acc is not None else None
    except (TypeError, ValueError):
        return None
    if not (-90 <= lat <= 90 and -180 <= lon <= 180):
        return None
    return {'lat': lat, 'lon': lon, 'accuracy': acc, 'source': 'browser_geolocation'}


def capture_browser_gps(uid=None, page='', topic='', subtopic='', event_type='gps_refresh', show_status=False):
    if 'current_gps' not in st.session_state:
        st.session_state.current_gps = None
    if get_geolocation is None:
        if show_status:
            st.info('GPS automático no activo. Instala: pip install streamlit-js-eval')
        return st.session_state.current_gps
    gps = None
    if not st.session_state.get('_gps_widget_rendered_this_run', False):
        st.session_state['_gps_widget_rendered_this_run'] = True
        try:
            try:
                raw = get_geolocation(key='bunsekichat_browser_gps_unique')
            except TypeError:
                raw = get_geolocation()
            gps = normalize_gps_payload(raw)
        except Exception:
            gps = None
    else:
        gps = st.session_state.get('current_gps')
    if gps:
        st.session_state.current_gps = gps
        if uid:
            session_key = f'location_open_logged_{uid}'
            if event_type != 'app_open' or not st.session_state.get(session_key):
                log_location_event(uid, event_type, page=page, topic=topic, subtopic=subtopic, gps=gps)
                if event_type == 'app_open':
                    st.session_state[session_key] = True
            update_profile_last_gps(uid, gps)
        if show_status:
            st.caption(f"📍 GPS activo: {gps['lat']:.6f}, {gps['lon']:.6f}")
    elif show_status and st.session_state.current_gps is None:
        st.caption('📍 Esperando permiso de ubicación del navegador...')
    return st.session_state.current_gps


def update_profile_last_gps(uid, gps):
    if not gps: return
    with DB_LOCK:
        execute("""UPDATE profiles SET last_gps_lat=%s, last_gps_lon=%s, last_gps_accuracy=%s, last_gps_at=%s WHERE user_id=%s""", (gps.get('lat'), gps.get('lon'), gps.get('accuracy'), now(), uid))


def log_location_event(uid, event_type, page='', topic='', subtopic='', gps=None):
    if not gps: return
    with DB_LOCK:
        execute("""INSERT INTO location_events(user_id,event_type,page,topic,subtopic,gps_lat,gps_lon,gps_accuracy,gps_source,created_at) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""", (uid, event_type, page, topic, subtopic, gps.get('lat'), gps.get('lon'), gps.get('accuracy'), gps.get('source'), now()))


def log_interaction(uid, role, message, topic='', subtopic='', level='', model='', latency=0, gps=None):
    clean = normalize_math(message)
    gps = gps or st.session_state.get('current_gps') or {}
    with DB_LOCK:
        execute("""INSERT INTO interactions(user_id,role,topic,subtopic,level,message,clean_message,model,tokens_est,latency_ms,created_at,gps_lat,gps_lon,gps_accuracy,gps_source) VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)""", (uid, role, topic, subtopic, level, message, clean, model, max(1, len(message)//4), latency, now(), gps.get('lat') if gps else None, gps.get('lon') if gps else None, gps.get('accuracy') if gps else None, gps.get('source') if gps else None))


def interactions(uid=None):
    if uid:
        return fetchall("SELECT * FROM interactions WHERE user_id=%s ORDER BY id", (uid,))
    return fetchall("SELECT i.*,u.username FROM interactions i JOIN users u ON u.id=i.user_id ORDER BY i.id DESC")


def get_location_events(uid=None):
    if uid:
        return fetchall("SELECT * FROM location_events WHERE user_id=%s ORDER BY id DESC", (uid,))
    return fetchall("SELECT le.*,u.username FROM location_events le JOIN users u ON u.id=le.user_id ORDER BY le.id DESC")

# ---------------- IA ----------------
def normalize_math(text:str)->str:
    if not text: return ""
    text = text.replace('\\(','$').replace('\\)','$').replace('\\[','$$').replace('\\]','$$')
    text = re.sub(r'```+\s*', '', text)
    text = text.replace('âˆ’','-').replace('×','\\times').replace('÷','\\div')
    # evita caracteres invisibles frecuentes
    return ''.join(ch for ch in text if ch.isprintable() or ch in '\n\t')

@st.cache_resource
def ai_client():
    if not GEMINI_API_KEY: return None
    return genai.Client(api_key=GEMINI_API_KEY)

def build_prompt(question, topic, subtopic, level, profile):
    return f"""
Eres BunsekiChat, tutor personalizado de matemáticas universitarias para estudiantes de bajos recursos.
Área: {topic}. Subtema: {subtopic}. Nivel: {level}.
Estudiante: {profile.get('first_names','')} {profile.get('last_names','')}. Materia: {profile.get('subject','')}. Curso: {profile.get('course','')}.

Reglas de respuesta:
1) Español claro, humano y docente.
2) Explica paso a paso, sin saltos algebraicos importantes.
3) Usa LaTeX limpio para fórmulas: $x^2$, $$\\int x dx$$. No uses caracteres raros ni escapes innecesarios.
4) Termina con una pregunta breve de verificación.
5) Si la consulta no es matemática, redirige amablemente al tema académico.

Pregunta del estudiante: {question}
"""

def ask_ai(question, topic, subtopic, level, profile):
    """Consulta la IA con fallback automático entre modelos Gemini.
    Prioriza Gemini 2.5 Flash y, si hay saturación/cuota/error temporal,
    intenta modelos alternativos antes de fallar.
    """
    client = ai_client()
    if not client:
        return (
            f"Configura GEMINI_API_KEY en el archivo .env ubicado en: {ENV_PATH}. "
            "Debe llamarse exactamente .env y contener GEMINI_API_KEY=tu_clave.",
            "sin_modelo",
            0,
        )

    prompt = build_prompt(question, topic, subtopic, level, profile)
    modelos = [
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite",
        "gemini-2.0-flash-lite",
        "gemini-2.0-flash",
    ]

    errores = []
    start_total = datetime.now()

    for modelo in modelos:
        start = datetime.now()
        try:
            resp = client.models.generate_content(
                model=modelo,
                contents=prompt,
            )
            ms = int((datetime.now() - start).total_seconds() * 1000)
            texto = normalize_math(resp.text or "No pude generar respuesta.")
            return texto, modelo, ms
        except Exception as e:
            err = str(e)
            errores.append(f"{modelo}: {err[:220]}")

            # Si es cuota/saturación/error temporal, probar siguiente modelo.
            err_lower = err.lower()
            if any(x in err_lower for x in [
                "429", "resource_exhausted", "quota", "rate", "503", "unavailable", "timeout"
            ]):
                time.sleep(1.5)
                continue

            # Para errores de API key o permisos, no tiene sentido probar otros modelos.
            if any(x in err_lower for x in ["api key", "permission", "unauthorized", "401", "403"]):
                break

            # Para otros errores, igual intenta el siguiente por robustez.
            time.sleep(0.8)
            continue

    total_ms = int((datetime.now() - start_total).total_seconds() * 1000)
    detalle = " | ".join(errores[-3:])
    return (
        "No se pudo consultar la IA con los modelos disponibles. "
        "Puede ser cuota agotada, saturación temporal o API key sin permisos. "
        f"Detalle resumido: {detalle}",
        "error_modelos_gemini",
        total_ms,
    )

# ---------------- UI helpers ----------------
def session_guard():
    if 'user' not in st.session_state: return
    timeout = get_timeout()
    last = st.session_state.get('last_activity', datetime.now())
    if datetime.now() - last > timedelta(minutes=timeout):
        st.session_state.clear(); st.warning("Sesión cerrada automáticamente por inactividad."); st.stop()
    st.session_state.last_activity = datetime.now()

def render_message(role, msg):
    cls='user' if role=='user' else 'bot'
    title='👨‍🎓 Estudiante' if role=='user' else 'BΣ BunsekiChat'
    st.markdown(f"<div class='{cls}'><b>{title}</b></div>", unsafe_allow_html=True)
    st.markdown(normalize_math(msg))

# ---------------- Quizzes ----------------
QUIZ_BANK = {
    "Álgebra lineal": [("Si A es 2x2, ¿qué indica det(A)=0?", ["A es invertible", "A no es invertible", "A es identidad"], 1), ("Producto punto de vectores ortogonales", ["1", "0", "-1"], 1)],
    "Estadística": [("La media es", ["El valor más frecuente", "El promedio", "La raíz cuadrada de la varianza"], 1), ("La desviación estándar mide", ["Dispersión", "Centro", "Asimetría siempre"], 0)],
    "Cálculo diferencial": [("Derivada de x^3", ["3x^2", "x^2", "3x"], 0), ("Si f'(x)>0, la función", ["crece", "decrece", "es constante"], 0)],
    "Cálculo integral": [("Integral de 2x dx", ["x^2+C", "2+C", "ln x"], 0), ("La integral definida representa", ["Solo pendiente", "Área neta acumulada", "Solo máximo"], 1)]
}


# Banco ampliado para pruebas adaptativas de 10 preguntas.
# La selección prioriza el historial real del estudiante: temas y subtemas más consultados.
ADAPTIVE_QUIZ_BANK = {
    "Álgebra lineal": [
        ("¿Qué representa un vector en R² o R³?", ["Una magnitud con dirección y sentido", "Solo un número", "Una ecuación cuadrática"], 0),
        ("Si dos vectores son ortogonales, su producto punto es:", ["1", "0", "-1"], 1),
        ("Una matriz cuadrada A es invertible si:", ["det(A) ≠ 0", "det(A)=0", "todas sus entradas son positivas"], 0),
        ("El determinante igual a cero indica que la matriz:", ["No es invertible", "Es identidad", "Tiene inversa única"], 0),
        ("Un sistema lineal homogéneo siempre tiene:", ["Al menos la solución trivial", "Ninguna solución", "Solo soluciones negativas"], 0),
        ("Un autovector de A cumple:", ["A v = λ v", "A+v=λ", "det(v)=0"], 0),
        ("El rango de una matriz mide:", ["La cantidad de filas/columnas linealmente independientes", "El número de decimales", "La suma de entradas"], 0),
        ("Dos vectores son linealmente dependientes si:", ["Uno puede escribirse como múltiplo del otro", "Siempre son perpendiculares", "Su producto punto es 1"], 0),
        ("La matriz identidad funciona como:", ["Elemento neutro de la multiplicación matricial", "Elemento nulo", "Matriz sin diagonal"], 0),
        ("El espacio columna está formado por:", ["Combinaciones lineales de las columnas", "Solo la diagonal", "Los determinantes parciales"], 0),
    ],
    "Estadística": [
        ("La media aritmética representa:", ["El promedio", "El dato más repetido", "La diferencia máxima"], 0),
        ("La mediana divide los datos ordenados en:", ["Dos partes iguales", "Tres partes", "Diez partes"], 0),
        ("La desviación estándar mide:", ["Dispersión de los datos", "Tamaño de muestra únicamente", "Moda"], 0),
        ("La probabilidad siempre está entre:", ["0 y 1", "-1 y 1", "1 y 1000"], 0),
        ("Una distribución normal se caracteriza por:", ["Forma de campana", "Solo valores enteros", "No tener media"], 0),
        ("La varianza es:", ["El promedio de desviaciones cuadráticas", "La raíz de la media", "La moda elevada"], 0),
        ("En regresión lineal simple se modela:", ["Relación entre X e Y", "Solo frecuencias", "Únicamente porcentajes"], 0),
        ("Un intervalo de confianza expresa:", ["Rango plausible para un parámetro", "Un dato exacto garantizado", "La moda"], 0),
        ("La hipótesis nula usualmente representa:", ["Ausencia de efecto o diferencia", "La conclusión obligatoria", "El promedio de la muestra"], 0),
        ("Una muestra debe procurar ser:", ["Representativa", "Siempre pequeña", "Sin variabilidad"], 0),
    ],
    "Cálculo diferencial": [
        ("La derivada representa principalmente:", ["Tasa de cambio instantánea", "Área acumulada", "Promedio de datos"], 0),
        ("Si f'(x)>0 en un intervalo, la función:", ["Crece", "Decrece", "Es discontinua siempre"], 0),
        ("La derivada de x³ es:", ["3x²", "x²", "3x"], 0),
        ("Un punto crítico ocurre cuando:", ["f'(x)=0 o no existe", "f(x)=1 siempre", "x es negativo"], 0),
        ("La segunda derivada ayuda a estudiar:", ["Concavidad", "Dominio únicamente", "Intercepto y"], 0),
        ("La regla de la cadena se usa para derivar:", ["Funciones compuestas", "Solo constantes", "Matrices"], 0),
        ("La derivada de sen(x) es:", ["cos(x)", "-cos(x)", "tan(x)"], 0),
        ("La derivada de una constante es:", ["0", "1", "La misma constante"], 0),
        ("Si f'(x)<0 en un intervalo, la función:", ["Decrece", "Crece", "Tiene máximo absoluto siempre"], 0),
        ("Una asíntota vertical suele asociarse a:", ["Valores donde la función crece sin límite", "Un punto medio", "Una raíz simple siempre"], 0),
    ],
    "Cálculo integral": [
        ("La integral indefinida representa:", ["Familia de antiderivadas", "Solo pendiente", "Solo máximo"], 0),
        ("La integral definida puede interpretarse como:", ["Área neta acumulada", "Derivada segunda", "Moda"], 0),
        ("La integral de 2x dx es:", ["x² + C", "2 + C", "ln(x) + C"], 0),
        ("La constante C aparece en:", ["Integrales indefinidas", "Solo límites", "Matrices"], 0),
        ("El teorema fundamental del cálculo relaciona:", ["Derivadas e integrales", "Matrices y vectores", "Media y moda"], 0),
        ("La sustitución se usa cuando:", ["Hay una composición y su derivada", "Solo hay números", "No hay variable"], 0),
        ("La integral de cos(x) dx es:", ["sen(x)+C", "-sen(x)+C", "tan(x)+C"], 0),
        ("La integral de 1/x dx es:", ["ln|x|+C", "x²+C", "e^x+C"], 0),
        ("En integración por partes se usa:", ["∫u dv = uv - ∫v du", "a²+b²=c²", "y=mx+b"], 0),
        ("Una integral impropia puede aparecer por:", ["Intervalo infinito o discontinuidad", "Solo por enteros", "Porque no tiene variable"], 0),
    ],
}

def adaptive_quiz_for_user(uid, current_topic):
    """Devuelve 10 preguntas vinculadas al historial del estudiante."""
    rows = interactions(uid)
    user_rows = [r for r in rows if r.get("role") == "user"]
    topic_counts = {}
    for r in user_rows:
        t = r.get("topic") or current_topic
        topic_counts[t] = topic_counts.get(t, 0) + 1
    ranked_topics = sorted(topic_counts, key=topic_counts.get, reverse=True)
    if current_topic not in ranked_topics:
        ranked_topics.append(current_topic)
    selected = []
    used = set()
    for t in ranked_topics:
        for q in ADAPTIVE_QUIZ_BANK.get(t, []):
            if q[0] not in used:
                selected.append((t, q))
                used.add(q[0])
            if len(selected) >= 10:
                return selected
    for t, bank in ADAPTIVE_QUIZ_BANK.items():
        for q in bank:
            if q[0] not in used:
                selected.append((t, q))
                used.add(q[0])
            if len(selected) >= 10:
                return selected
    return selected[:10]

def save_quiz(uid, topic, old, new, score, passed, answers):
    with DB_LOCK:
        execute("INSERT INTO quizzes(user_id,topic,level_from,level_to,score,passed,answers_json,created_at) VALUES(%s,%s,%s,%s,%s,%s,%s,%s)",
                (uid, topic, old, new, score, int(passed), json.dumps(answers, ensure_ascii=False), now()))


# ---------------- Control de lenguaje y sugerencias ----------------
BAD_WORDS = [
    "puta", "puto", "mierda", "carajo", "verga", "pendejo", "pendeja",
    "idiota", "imbecil", "imbécil", "maldito", "maldita", "cojudo", "cojuda"
]

def contains_bad_language(text: str) -> bool:
    text = (text or "").lower()
    return any(re.search(rf"\b{re.escape(w)}\b", text) for w in BAD_WORDS)

def prompt_suggestions(topic: str, subtopic: str, level: str):
    return [
        f"Explícame {subtopic} desde cero con un ejemplo universitario.",
        f"Dame 3 ejercicios resueltos de {subtopic} en nivel {level}.",
        f"Corrige mi procedimiento paso a paso sobre {subtopic}.",
        f"Crea una prueba corta de {topic} y califícame.",
        f"Muéstrame un error común al estudiar {subtopic} y cómo evitarlo.",
    ]

# ---------------- Exportación docente ----------------
def _rows_to_df(rows, columns=None):
    """Convierte listas de diccionarios en DataFrame sin depender de pandas.read_sql_query.
    Esto evita exportaciones vacías o filas extrañas cuando el driver PostgreSQL cambia el cursor.
    """
    df = pd.DataFrame(rows or [])
    if df.empty and columns:
        return pd.DataFrame(columns=columns)
    return df


def _clean_report_df(df: pd.DataFrame) -> pd.DataFrame:
    """Limpia filas inválidas y asegura columnas esperadas para exportar."""
    expected = ["id", "created_at", "role", "topic", "subtopic", "level", "model",
                "tokens_est", "latency_ms", "gps_lat", "gps_lon", "gps_accuracy",
                "gps_source", "message"]
    if df is None or df.empty:
        return pd.DataFrame(columns=expected)

    df = df.copy()
    for col in expected:
        if col not in df.columns:
            df[col] = ""

    # Corrige casos donde alguna exportación anterior registró filas como encabezados.
    mask_bad = pd.Series(False, index=df.index)
    for col in ["created_at", "role", "topic", "subtopic", "level", "message"]:
        mask_bad = mask_bad | df[col].astype(str).str.strip().eq(col)

    df = df[~mask_bad].copy()
    df = df[expected]

    # Orden y limpieza visual
    df["created_at"] = df["created_at"].fillna("").astype(str)
    df["message"] = df["message"].fillna("").astype(str)
    return df


def get_teacher_tables():
    users_rows = fetchall("""
        SELECT u.id,u.username,u.role,u.active,u.created_at,u.last_login,
               p.first_names,p.last_names,p.course,p.subject,p.course_level,p.parallel,p.shift,p.cohort,p.full_name_normalized,
               p.teacher,p.province,p.city,p.level
        FROM users u LEFT JOIN profiles p ON p.user_id=u.id
        ORDER BY u.id DESC
    """)
    logs_rows = fetchall("""
        SELECT i.*,u.username,p.first_names,p.last_names,p.course,p.subject,p.course_level,p.parallel,p.shift,p.cohort,
               p.teacher,p.level AS profile_level
        FROM interactions i
        JOIN users u ON u.id=i.user_id
        LEFT JOIN profiles p ON p.user_id=i.user_id
        ORDER BY i.id DESC
    """)
    quizzes_rows = fetchall("""
        SELECT q.*,u.username,p.first_names,p.last_names,p.course,p.subject,p.course_level,p.parallel,p.shift,p.cohort,
               p.teacher,p.level AS profile_level
        FROM quizzes q
        JOIN users u ON u.id=q.user_id
        LEFT JOIN profiles p ON p.user_id=q.user_id
        ORDER BY q.id DESC
    """)
    users = _rows_to_df(users_rows, ["id","username","role","active","created_at","last_login","first_names","last_names","course","subject","course_level","parallel","shift","cohort","full_name_normalized","teacher","province","city","level"])
    logs = _rows_to_df(logs_rows, ["id","user_id","role","topic","subtopic","level","message","clean_message","model","tokens_est","latency_ms","created_at","username","first_names","last_names","course","subject","course_level","parallel","shift","cohort","teacher","profile_level","gps_lat","gps_lon","gps_accuracy","gps_source"])
    quizzes = _rows_to_df(quizzes_rows, ["id","user_id","topic","level_from","level_to","score","passed","answers_json","created_at","username","first_names","last_names","course","subject","course_level","parallel","shift","cohort","teacher","profile_level"])
    return users, logs, quizzes


def build_student_summary(uid: int):
    profile = fetchone("""
        SELECT u.id,u.username,u.role,u.active,u.created_at,u.last_login,
               p.first_names,p.last_names,p.course,p.subject,p.course_level,p.parallel,p.shift,p.cohort,p.full_name_normalized,
               p.teacher,p.phone,p.province,p.city,p.canton,p.address,p.level
        FROM users u LEFT JOIN profiles p ON p.user_id=u.id WHERE u.id=%s
    """, (uid,)) or {"id": uid, "username": f"usuario_{uid}"}

    rows = fetchall("""
        SELECT id,created_at,role,topic,subtopic,level,model,tokens_est,latency_ms,
               gps_lat,gps_lon,gps_accuracy,gps_source,
               COALESCE(NULLIF(clean_message,''), message) AS message
        FROM interactions
        WHERE user_id=%s
        ORDER BY id ASC
    """, (uid,))
    logs = _clean_report_df(pd.DataFrame(rows or []))
    return profile, logs


def export_student_docx(uid: int) -> bytes:
    if Document is None:
        raise RuntimeError("Instala python-docx: pip install python-docx")
    profile, df = build_student_summary(uid)
    doc = Document()
    doc.add_heading("Reporte individual BunsekiChat", level=1)
    nombre = f"{profile.get('first_names') or ''} {profile.get('last_names') or ''}".strip()
    doc.add_paragraph(f"Estudiante: {nombre or profile.get('username','')}")
    doc.add_paragraph(f"Usuario: {profile.get('username','')}")
    doc.add_paragraph(f"Materia: {profile.get('subject') or 'Sin materia'}")
    doc.add_paragraph(f"Curso: {profile.get('course') or 'Sin curso'}")
    doc.add_paragraph(f"Cohorte: {profile.get('cohort') or profile.get('course') or 'Sin cohorte'}")
    doc.add_paragraph(f"Docente: {profile.get('teacher') or 'Sin docente'}")
    doc.add_paragraph(f"Nivel actual: {profile.get('level') or 'Sin nivel'}")
    doc.add_paragraph(f"Fecha de exportación: {now()}")
    doc.add_heading("Consultas e interacciones", level=2)
    if df.empty:
        doc.add_paragraph("No existen consultas registradas para este estudiante.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Fecha", "Rol", "Tema", "Subtema", "Nivel", "Mensaje"]):
            hdr[i].text = h
        for _, r in df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(r.get("created_at", ""))
            cells[1].text = str(r.get("role", ""))
            cells[2].text = str(r.get("topic", ""))
            cells[3].text = str(r.get("subtopic", ""))
            cells[4].text = str(r.get("level", ""))
            msg = str(r.get("message", ""))
            cells[5].text = msg[:1200]
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def export_student_pdf(uid: int) -> bytes:
    if canvas is None or A4 is None:
        raise RuntimeError("Instala reportlab: pip install reportlab")
    profile, df = build_student_summary(uid)
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4
    y = height - 50
    def line(text, size=10, bold=False):
        nonlocal y
        if y < 55:
            c.showPage(); y = height - 50
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        safe = str(text).replace("\n", " ")
        c.drawString(45, y, safe[:115])
        y -= size + 7
    line("Reporte individual BunsekiChat", 16, True)
    nombre = f"{profile.get('first_names') or ''} {profile.get('last_names') or ''}".strip()
    line(f"Estudiante: {nombre or profile.get('username','')}", 10)
    line(f"Usuario: {profile.get('username','')} | Materia: {profile.get('subject') or 'Sin materia'} | Curso: {profile.get('course') or 'Sin curso'}", 10)
    line(f"Cohorte: {profile.get('cohort') or profile.get('course') or 'Sin cohorte'} | Nivel: {profile.get('level') or 'Sin nivel'}", 10)
    line(f"Docente: {profile.get('teacher') or 'Sin docente'} | Exportado: {now()}", 10)
    line("Consultas e interacciones", 13, True)
    if df.empty:
        line("No existen consultas registradas para este estudiante.")
    else:
        for _, r in df.iterrows():
            line(f"[{r.get('created_at','')}] {r.get('role','')} · {r.get('topic','')} / {r.get('subtopic','')}", 9, True)
            msg = str(r.get("message", ""))
            for chunk in [msg[i:i+100] for i in range(0, min(len(msg), 700), 100)]:
                line("  " + chunk, 8)
            y -= 4
    c.save()
    bio.seek(0)
    return bio.getvalue()


# ---------------- Planes analíticos y evaluación adaptativa IA ----------------
def _strip_code_fences(raw: str) -> str:
    raw = (raw or "").strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.I)
    raw = re.sub(r"\s*```$", "", raw)
    return raw.strip()


def _safe_json_loads(raw, default):
    try:
        return json.loads(_strip_code_fences(raw))
    except Exception:
        return default


def read_uploaded_plan_text(uploaded_file) -> str:
    """Extrae texto de TXT, DOCX o PDF sin romper la app si falta una librería."""
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    uploaded_file.seek(0)
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        if Document is None:
            return ""
        try:
            doc = Document(BytesIO(data))
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            return ""
    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(data))
            return "\n".join([(page.extract_text() or "") for page in reader.pages])
        except Exception:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(BytesIO(data))
                return "\n".join([(page.extract_text() or "") for page in reader.pages])
            except Exception:
                return ""
    return data.decode("utf-8", errors="ignore")


def parse_plan_topics_with_ai(raw_text: str, title: str = "") -> dict:
    """Convierte un plan analítico en unidades/temas/resultados. Usa IA si hay API; si no, fallback robusto."""
    raw_text = (raw_text or "")[:24000]
    fallback_topics = []
    # fallback por líneas con palabras académicas frecuentes
    for line in raw_text.splitlines():
        clean = re.sub(r"\s+", " ", line).strip(" -•\t")
        if 8 <= len(clean) <= 160 and any(w in clean.lower() for w in ["unidad", "tema", "resultado", "competencia", "contenido", "aprendizaje", "objetivo"]):
            fallback_topics.append({
                "unit_name": clean[:80],
                "topic": clean[:90],
                "subtopic": "",
                "learning_outcome": "Comprender y aplicar el contenido del plan analítico.",
                "bloom_level": "Comprender/Aplicar",
                "keywords": clean
            })
        if len(fallback_topics) >= 12:
            break
    if not fallback_topics:
        fallback_topics = [{
            "unit_name": "Unidad general",
            "topic": title or "Contenido del plan analítico",
            "subtopic": "",
            "learning_outcome": "Comprender, aplicar y evaluar los contenidos del plan.",
            "bloom_level": "Comprender/Aplicar",
            "keywords": title or "plan analítico"
        }]

    client = ai_client()
    if not client:
        return {"summary": "Plan cargado con extracción básica local.", "topics": fallback_topics}

    prompt = f"""
Actúa como diseñador curricular universitario.
Analiza este plan analítico y devuelve SOLO JSON válido con esta estructura:
{{
  "summary": "resumen breve",
  "topics": [
    {{
      "unit_name": "Unidad 1",
      "topic": "tema principal",
      "subtopic": "subtema",
      "learning_outcome": "resultado de aprendizaje evaluable",
      "bloom_level": "Recordar/Comprender/Aplicar/Analizar/Evaluar/Crear",
      "keywords": "palabras clave separadas por coma"
    }}
  ]
}}
Máximo 18 topics. No uses markdown.

Título: {title}
PLAN:
{raw_text}
"""
    try:
        resp = client.models.generate_content(model="gemini-2.5-flash-lite", contents=prompt)
        parsed = _safe_json_loads(resp.text, {})
        if isinstance(parsed, dict) and isinstance(parsed.get("topics"), list) and parsed["topics"]:
            return parsed
    except Exception:
        pass
    return {"summary": "Plan cargado con extracción básica local.", "topics": fallback_topics}


def save_analytic_plan(title, course, teacher_id, filename, raw_text, parsed, academic_context=None):
    academic_context = academic_context or {}
    row = execute("""
        INSERT INTO analytic_plans(title, course, subject, course_level, parallel, shift, teacher_id, filename, raw_text, summary, created_at)
        VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id
    """, (
        title, course,
        academic_context.get("subject"), academic_context.get("course_level"),
        academic_context.get("parallel"), academic_context.get("shift"),
        teacher_id, filename, raw_text, parsed.get("summary",""), now()
    ), returning=True)
    plan_id = row["id"]
    for item in parsed.get("topics", []):
        execute("""
            INSERT INTO plan_topics(plan_id, unit_name, topic, subtopic, learning_outcome, bloom_level, keywords)
            VALUES(%s,%s,%s,%s,%s,%s,%s)
        """, (
            plan_id,
            item.get("unit_name",""),
            item.get("topic",""),
            item.get("subtopic",""),
            item.get("learning_outcome",""),
            item.get("bloom_level",""),
            item.get("keywords",""),
        ))
    return plan_id


def get_analytic_plans():
    return fetchall("""
        SELECT p.*, u.username AS teacher_username
        FROM analytic_plans p
        LEFT JOIN users u ON u.id=p.teacher_id
        ORDER BY p.id DESC
    """)


def get_plan_topics(plan_id=None):
    if plan_id:
        return fetchall("SELECT * FROM plan_topics WHERE plan_id=%s ORDER BY id", (plan_id,))
    return fetchall("""
        SELECT pt.*, p.title AS plan_title, p.course AS plan_course, p.subject AS plan_subject
        FROM plan_topics pt
        JOIN analytic_plans p ON p.id=pt.plan_id
        ORDER BY p.id DESC, pt.id
    """)


def get_student_consultation_profile(uid: int) -> dict:
    rows = interactions(uid)
    user_rows = [r for r in rows if str(r.get("role","")).lower() == "user"]
    topic_counts = {}
    subtopic_counts = {}
    recent_messages = []
    for r in user_rows:
        t = r.get("topic") or "General"
        s = r.get("subtopic") or ""
        topic_counts[t] = topic_counts.get(t, 0) + 1
        if s:
            subtopic_counts[s] = subtopic_counts.get(s, 0) + 1
        if r.get("message"):
            recent_messages.append(str(r.get("message"))[:500])
    ranked_topics = sorted(topic_counts, key=topic_counts.get, reverse=True)
    ranked_subtopics = sorted(subtopic_counts, key=subtopic_counts.get, reverse=True)
    return {
        "topic_counts": topic_counts,
        "ranked_topics": ranked_topics,
        "ranked_subtopics": ranked_subtopics,
        "recent_messages": recent_messages[-8:],
        "total_questions": len(user_rows),
    }


def is_research_quiz(quiz_type):
    return str(quiz_type or "").lower() in {"pretest", "posttest"}


def quiz_type_label(quiz_type):
    return RESEARCH_TEST_TYPES.get(str(quiz_type or "adaptive").lower(), "Evaluación adaptativa")


def local_research_question_pool(subject):
    key = strip_accents(subject).lower()
    if "calculo diferencial" in key or "calculo" in key:
        return [dict(q) for q in DIFFERENTIAL_CALCULUS_QUESTION_BANK]
    return []


def fallback_questions(n_questions, base_topics=None, assessment_context=None):
    assessment_context = assessment_context or {}
    pool = local_research_question_pool(assessment_context.get("subject", ""))
    if pool:
        return pool
    base_topics = [x for x in (base_topics or []) if x] or [assessment_context.get("subject") or "Matemáticas"]
    fallback = []
    for i in range(int(n_questions)):
        t = base_topics[i % len(base_topics)] if base_topics else "Matemáticas"
        fallback.append({
            "question": f"Pregunta {i+1}: Resuelve o explica un concepto clave relacionado con {t}.",
            "options": [
                f"Aplicar correctamente el procedimiento de {t}",
                "Elegir una respuesta sin justificar",
                "Ignorar las condiciones del problema",
                "Usar una fórmula no relacionada",
            ],
            "correct_answer": f"Aplicar correctamente el procedimiento de {t}",
            "explanation": f"La respuesta correcta debe usar el procedimiento o concepto central de {t}.",
            "bloom_level": "Aplicar",
            "topic": t,
        })
    return fallback


def make_assessment_seed(uid, quiz_type, academic_context, salt=None):
    context = academic_context or {}
    raw = "|".join([
        str(uid),
        str(quiz_type or "adaptive"),
        str(context.get("subject", "")),
        str(context.get("course_level", "")),
        str(context.get("parallel", "")),
        str(context.get("shift", "")),
        str(salt or time.time_ns()),
    ])
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def normalize_question_payload(question):
    if not isinstance(question, dict):
        return None
    opts = [normalize_spaces(o) for o in question.get("options", []) if normalize_spaces(o)]
    correct = normalize_spaces(question.get("correct_answer", ""))
    if correct and correct not in opts:
        opts.append(correct)
    opts = list(dict.fromkeys(opts))
    if len(opts) > 4:
        opts = opts[:4] if correct in opts[:4] else opts[:3] + [correct]
    if len(opts) < 2 or not correct:
        return None
    question_text = normalize_spaces(question.get("question", ""))
    if not question_text:
        return None
    return {
        "question": question_text,
        "options": opts,
        "correct_answer": correct,
        "explanation": normalize_spaces(question.get("explanation", "")),
        "bloom_level": normalize_spaces(question.get("bloom_level", "Aplicar")) or "Aplicar",
        "topic": normalize_spaces(question.get("topic", "")),
    }


def randomize_questions_for_student(questions, uid, quiz_type, n_questions, seed):
    rng = random.Random(seed)
    clean = []
    seen = set()
    for q in questions:
        item = normalize_question_payload(q)
        if not item:
            continue
        key = strip_accents(item["question"]).lower()
        if key in seen:
            continue
        seen.add(key)
        opts = item["options"][:]
        rng.shuffle(opts)
        item["options"] = opts
        clean.append(item)
    rng.shuffle(clean)
    return clean[:int(n_questions)]


def generate_questions_with_ai(uid: int, plan_id=None, difficulty="Intermedio", n_questions=10, retry_focus=None, assessment_context=None):
    assessment_context = assessment_context or {}
    quiz_type = assessment_context.get("quiz_type", "adaptive")
    profile = get_student_consultation_profile(uid)
    plan_topics = get_plan_topics(plan_id) if plan_id else get_plan_topics()
    plan_context = "\n".join([
        f"- {x.get('unit_name','')} | {x.get('topic','')} | {x.get('subtopic','')} | Resultado: {x.get('learning_outcome','')} | Bloom: {x.get('bloom_level','')}"
        for x in plan_topics[:60]
    ])
    consulted = ", ".join(profile["ranked_topics"][:5] + profile["ranked_subtopics"][:5]) or "sin historial suficiente"
    recent = "\n".join([f"- {m}" for m in profile["recent_messages"]])
    base_topics = profile["ranked_topics"] or [x.get("topic") for x in plan_topics if x.get("topic")] or list(TOPICS.keys())
    fallback = fallback_questions(n_questions, base_topics, assessment_context)

    client = ai_client()
    if not client:
        return fallback

    research_rules = ""
    if is_research_quiz(quiz_type):
        research_rules = f"""
MODO INVESTIGACIÓN:
- Tipo de instrumento: {quiz_type_label(quiz_type)}.
- Título del estudio: {RESEARCH_TITLE}.
- Materia: {assessment_context.get('subject') or 'Cálculo diferencial'}.
- Curso: {assessment_context.get('course_level') or ''}.
- Paralelo: {assessment_context.get('parallel') or ''}.
- Jornada: {assessment_context.get('shift') or ''}.
- Genera exactamente {int(n_questions)} preguntas equivalentes en dificultad para comparar pretest y posttest.
- Si es pretest, evalúa saberes previos sin depender del historial de consultas.
- Si es posttest, evalúa las mismas competencias con enunciados no idénticos al pretest.
"""

    prompt = f"""
Eres un experto en evaluación educativa adaptativa, matemática universitaria y taxonomía de Bloom.
Genera {int(n_questions)} preguntas tipo test para BunsekiChat. Deben evaluar lo que el estudiante ha consultado y alinearse al plan analítico.

CONDICIONES:
- Devuelve SOLO JSON válido, sin markdown.
- Cada pregunta debe tener 4 opciones.
- Debe existir correct_answer igual a una de las opciones.
- Nivel de dificultad: {difficulty}.
- Si el estudiante no obtiene 70%, estas preguntas se usarán para reforzar y generar nuevos intentos.
- Incluye explicación pedagógica breve.
- Usa lenguaje claro en español.
{research_rules}

JSON:
[
  {{
    "question": "texto de la pregunta",
    "options": ["A", "B", "C", "D"],
    "correct_answer": "A",
    "explanation": "explicación",
    "bloom_level": "Aplicar",
    "topic": "tema evaluado"
  }}
]

PLAN ANALÍTICO:
{plan_context}

TEMAS CONSULTADOS POR EL ESTUDIANTE:
{consulted}

CONSULTAS RECIENTES:
{recent}

FOCO DE REINTENTO SI APLICA:
{retry_focus or "Ninguno"}
"""
    try:
        resp = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
        parsed = _safe_json_loads(resp.text, [])
        if isinstance(parsed, list) and parsed:
            clean = []
            for q in parsed[:int(n_questions)]:
                opts = q.get("options") or []
                if len(opts) >= 2 and q.get("correct_answer") in opts:
                    clean.append({
                        "question": str(q.get("question","")).strip(),
                        "options": [str(o) for o in opts[:4]],
                        "correct_answer": str(q.get("correct_answer","")),
                        "explanation": str(q.get("explanation","")),
                        "bloom_level": str(q.get("bloom_level","Aplicar")),
                        "topic": str(q.get("topic","")),
                    })
            if clean:
                return clean
    except Exception:
        pass
    return fallback


def create_adaptive_quiz(uid, plan_id=None, difficulty="Intermedio", n_questions=10, parent_quiz_id=None, retry_focus=None, quiz_type="adaptive", academic_context=None):
    academic_context = academic_context or {}
    quiz_type = str(quiz_type or "adaptive").lower()
    if is_research_quiz(quiz_type):
        n_questions = RESEARCH_QUESTION_COUNT
    prior_attempts = 1
    if parent_quiz_id:
        parent = fetchone("SELECT attempt_no FROM adaptive_quizzes WHERE id=%s", (parent_quiz_id,))
        prior_attempts = int(parent.get("attempt_no", 1)) + 1 if parent else 2

    assessment_context = {**academic_context, "quiz_type": quiz_type}
    questions = generate_questions_with_ai(uid, plan_id, difficulty, n_questions, retry_focus, assessment_context)
    if len(questions) < int(n_questions):
        questions.extend(fallback_questions(n_questions, assessment_context=assessment_context))
    random_seed = make_assessment_seed(uid, quiz_type, academic_context)
    questions = randomize_questions_for_student(questions, uid, quiz_type, n_questions, random_seed)
    title = "Evaluación adaptativa IA"
    if is_research_quiz(quiz_type):
        title = f"{quiz_type_label(quiz_type)} - {academic_context.get('subject') or 'Materia'}"
    source_topic = ", ".join(sorted(set([q.get("topic","") for q in questions if q.get("topic")]))[:3])
    row = execute("""
        INSERT INTO adaptive_quizzes(
            user_id, plan_id, title, quiz_type, source_topic, difficulty,
            subject, course_level, parallel, shift, cohort, research_title,
            question_count, random_seed, attempt_no, parent_quiz_id, status, created_at
        )
        VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id
    """, (
        uid, plan_id, title, quiz_type, source_topic, difficulty,
        academic_context.get("subject"), academic_context.get("course_level"),
        academic_context.get("parallel"), academic_context.get("shift"),
        academic_context.get("cohort") or academic_context.get("course"),
        RESEARCH_TITLE if is_research_quiz(quiz_type) else "",
        int(n_questions), random_seed, prior_attempts, parent_quiz_id, "generated", now()
    ), returning=True)
    quiz_id = row["id"]
    for position, q in enumerate(questions, start=1):
        execute("""
            INSERT INTO adaptive_questions(quiz_id, question, options_json, correct_answer, position, topic, explanation, bloom_level)
            VALUES(%s,%s,%s,%s,%s,%s,%s,%s)
        """, (quiz_id, q["question"], json.dumps(q["options"], ensure_ascii=False), q["correct_answer"], position, q.get("topic",""), q.get("explanation",""), q.get("bloom_level","Aplicar")))
    return quiz_id


def load_adaptive_quiz(quiz_id):
    quiz = fetchone("SELECT * FROM adaptive_quizzes WHERE id=%s", (quiz_id,))
    questions = fetchall("SELECT * FROM adaptive_questions WHERE quiz_id=%s ORDER BY COALESCE(position, id), id", (quiz_id,))
    for q in questions:
        q["options"] = _safe_json_loads(q.get("options_json"), [])
    return quiz, questions


def get_latest_adaptive_quiz(uid, quiz_type=None):
    if quiz_type:
        return fetchone("SELECT * FROM adaptive_quizzes WHERE user_id=%s AND COALESCE(quiz_type,'adaptive')=%s ORDER BY id DESC LIMIT 1", (uid, quiz_type))
    return fetchone("SELECT * FROM adaptive_quizzes WHERE user_id=%s ORDER BY id DESC LIMIT 1", (uid,))


def grade_adaptive_quiz(quiz_id, answers: dict):
    quiz, questions = load_adaptive_quiz(quiz_id)
    total = max(1, len(questions))
    correct = 0
    weak_topics = []
    for q in questions:
        ans = answers.get(str(q["id"]), "")
        ok = str(ans).strip() == str(q.get("correct_answer","")).strip()
        if ok:
            correct += 1
        else:
            weak_topics.append(q.get("question","")[:120])
        execute("UPDATE adaptive_questions SET user_answer=%s, is_correct=%s WHERE id=%s", (ans, ok, q["id"]))
    score = round(correct / total * 100, 2)
    passed = score >= 70
    recommendation = "Aprobado. Puedes avanzar al siguiente reto." if passed else "Refuerzo automático sugerido: revisa las preguntas falladas y genera un nuevo intento."
    execute("""
        UPDATE adaptive_quizzes
        SET score=%s, passed=%s, status=%s, recommendation=%s, completed_at=%s
        WHERE id=%s
    """, (score, passed, "completed", recommendation, now(), quiz_id))
    return score, passed, weak_topics, recommendation


def render_teacher_plan_manager(user):
    st.markdown("""
    <div class='eval-hero'>
        <h2>📚 Planes analíticos inteligentes</h2>
        <p>Carga planes analíticos, sílabos o contenidos. BunsekiChat extrae temas, resultados de aprendizaje y niveles Bloom para crear evaluaciones adaptativas.</p>
    </div>
    """, unsafe_allow_html=True)

    with st.container():
        st.markdown("<div class='plan-upload-box'><b>Subir nuevo plan analítico</b></div>", unsafe_allow_html=True)
        c1, c2 = st.columns([1.2, 1])
        title = c1.text_input("Título del plan", placeholder="Ej. Cálculo diferencial - Unidad 1")
        subject = c2.selectbox("Materia", unique_values(ACADEMIC_SUBJECTS), index=0)
        p1, p2, p3 = st.columns(3)
        with p1:
            course_level = academic_selectbox("Curso / nivel", ACADEMIC_COURSE_LEVELS, ACADEMIC_COURSE_LEVELS[0], "plan_course_level", normalize_course_level)
        with p2:
            parallel = academic_selectbox("Paralelo", ACADEMIC_PARALLELS, ACADEMIC_PARALLELS[0], "plan_parallel", normalize_parallel)
        with p3:
            shift = academic_selectbox("Jornada", ACADEMIC_SHIFTS, ACADEMIC_SHIFTS[0], "plan_shift", normalize_shift)
        course = build_course_label(course_level, parallel, shift)
        uploaded = st.file_uploader("Archivo del plan analítico", type=["txt", "md", "docx", "pdf"])
        if st.button("Procesar y guardar plan", use_container_width=True):
            raw_text = read_uploaded_plan_text(uploaded)
            if not title.strip():
                st.warning("Escribe un título para el plan.")
            elif not raw_text.strip():
                st.warning("No se pudo extraer texto del archivo. Prueba con TXT/DOCX o instala pypdf para PDF.")
            else:
                parsed = parse_plan_topics_with_ai(raw_text, title)
                academic_context = {
                    "subject": normalize_academic_text(subject),
                    "course_level": normalize_course_level(course_level),
                    "parallel": normalize_parallel(parallel),
                    "shift": normalize_shift(shift),
                    "cohort": course,
                }
                plan_id = save_analytic_plan(title.strip(), course.strip(), user.get("id"), uploaded.name if uploaded else "", raw_text, parsed, academic_context)
                st.success(f"Plan guardado correctamente. ID: {plan_id}. Temas extraídos: {len(parsed.get('topics', []))}")
                st.rerun()

    plans = pd.DataFrame(get_analytic_plans())
    if plans.empty:
        st.info("Aún no hay planes analíticos cargados.")
        return

    st.markdown("<div class='eval-card'><h3>Planes cargados</h3></div>", unsafe_allow_html=True)
    plan_cols = [c for c in ["id","title","subject","course","course_level","parallel","shift","teacher_username","created_at","summary"] if c in plans.columns]
    st.dataframe(plans[plan_cols], use_container_width=True, height=280)

    plan_id = st.selectbox("Ver temas de un plan", plans["id"].tolist(), format_func=lambda x: f"{x} · {plans.loc[plans['id']==x,'title'].iloc[0]}")
    topics_df = pd.DataFrame(get_plan_topics(int(plan_id)))
    if not topics_df.empty:
        st.dataframe(topics_df[["unit_name","topic","subtopic","learning_outcome","bloom_level","keywords"]], use_container_width=True, height=360)


def render_student_adaptive_evaluation(user, prof, topic, subtopic):
    academic_context = academic_context_from_profile(prof)
    st.markdown("""
    <div class='eval-hero'>
        <h2>🧠 Evaluación e investigación</h2>
        <p>Selecciona evaluación adaptativa, pretest o posttest. Los instrumentos de investigación usan 20 preguntas y quedan vinculados a materia, curso, paralelo y jornada.</p>
    </div>
    """, unsafe_allow_html=True)

    mode_label = st.radio(
        "Tipo de evaluación",
        ["Evaluación adaptativa", "Pretest", "Posttest"],
        horizontal=True,
        key="assessment_mode",
    )
    quiz_type = {
        "Evaluación adaptativa": "adaptive",
        "Pretest": "pretest",
        "Posttest": "posttest",
    }[mode_label]
    research_mode = is_research_quiz(quiz_type)

    profile = get_student_consultation_profile(user["id"])
    a,b,c = st.columns(3)
    a.markdown(f"<div class='metric'><b>{profile['total_questions']}</b><br><span>Consultas usadas</span></div>", unsafe_allow_html=True)
    b.markdown(f"<div class='metric'><b>{academic_context.get('cohort')}</b><br><span>Cohorte</span></div>", unsafe_allow_html=True)
    last_quiz = get_latest_adaptive_quiz(user["id"], quiz_type)
    last_score = "N/D" if not last_quiz or last_quiz.get("score") is None else f"{float(last_quiz.get('score')):.0f}%"
    c.markdown(f"<div class='metric'><b>{last_score}</b><br><span>{quiz_type_label(quiz_type)}</span></div>", unsafe_allow_html=True)

    plans = pd.DataFrame(get_analytic_plans())
    if not plans.empty and research_mode and "subject" in plans.columns:
        subject_key = strip_accents(academic_context.get("subject", "")).lower()
        filtered_plans = plans[
            plans["subject"].fillna("").astype(str).map(lambda x: strip_accents(x).lower()).isin(["", subject_key])
        ].copy()
        if not filtered_plans.empty:
            plans = filtered_plans
    plan_options = [None] + (plans["id"].tolist() if not plans.empty else [])
    plan_id = st.selectbox(
        "Plan analítico base",
        plan_options,
        format_func=lambda x: "Todos los planes disponibles" if x is None else f"{x} · {plans.loc[plans['id']==x,'title'].iloc[0]}",
    )
    difficulty = st.selectbox("Dificultad", ["Inicial", "Básico", "Intermedio", "Avanzado"], index=2, disabled=research_mode)
    n_questions = RESEARCH_QUESTION_COUNT if research_mode else st.slider("Número de preguntas", 5, 20, 10)

    if research_mode:
        st.markdown(
            f"<div class='eval-card'><b>{quiz_type_label(quiz_type)} de investigación</b><br>"
            f"<span class='eval-chip'>{academic_context.get('subject')}</span>"
            f"<span class='eval-chip'>{academic_context.get('cohort')}</span>"
            f"<span class='eval-chip'>{RESEARCH_QUESTION_COUNT} preguntas aleatorias</span></div>",
            unsafe_allow_html=True,
        )

    if quiz_type == "posttest":
        completed_pretest = get_latest_adaptive_quiz(user["id"], "pretest")
        if not completed_pretest or completed_pretest.get("status") != "completed":
            st.warning("El posttest se habilita cuando el estudiante haya completado el pretest.")
            return

    current_key = f"current_adaptive_quiz_id_{quiz_type}"
    col1, col2 = st.columns(2)
    can_generate = not (research_mode and last_quiz)
    generate_label = f"Iniciar {quiz_type_label(quiz_type)}" if research_mode else "Generar nueva evaluación IA"
    if col1.button(generate_label, use_container_width=True, disabled=not can_generate):
        quiz_id = create_adaptive_quiz(user["id"], plan_id, difficulty, n_questions, quiz_type=quiz_type, academic_context=academic_context)
        st.session_state[current_key] = quiz_id
        st.rerun()
    if research_mode and not can_generate:
        if last_quiz.get("status") == "completed":
            st.info(f"{quiz_type_label(quiz_type)} ya fue completado. El resultado queda registrado para la investigación.")
        else:
            st.info(f"{quiz_type_label(quiz_type)} ya fue generado. Completa el instrumento pendiente para conservar la trazabilidad.")

    if not research_mode and last_quiz and last_quiz.get("status") == "completed" and not last_quiz.get("passed"):
        if col2.button("Reintento automático por no alcanzar 70%", use_container_width=True):
            _, questions = load_adaptive_quiz(last_quiz["id"])
            weak = [q.get("question","") for q in questions if not q.get("is_correct")]
            quiz_id = create_adaptive_quiz(
                user["id"],
                last_quiz.get("plan_id"),
                difficulty,
                n_questions,
                parent_quiz_id=last_quiz["id"],
                retry_focus="\n".join(weak[:5]),
                quiz_type=quiz_type,
                academic_context=academic_context
            )
            st.session_state[current_key] = quiz_id
            st.rerun()

    quiz_id = st.session_state.get(current_key) or (last_quiz.get("id") if last_quiz else None)
    if not quiz_id:
        st.info("Inicia una evaluación para comenzar.")
        return

    quiz, questions = load_adaptive_quiz(int(quiz_id))
    if not quiz or not questions:
        st.warning("No se pudo cargar la evaluación.")
        return

    st.markdown(
        f"<div class='eval-card'><h3>{quiz.get('title','Evaluación adaptativa')}</h3>"
        f"<span class='eval-chip'>Intento: {quiz.get('attempt_no',1)}</span>"
        f"<span class='eval-chip'>Tipo: {quiz_type_label(quiz.get('quiz_type') or 'adaptive')}</span>"
        f"<span class='eval-chip'>Dificultad: {quiz.get('difficulty','')}</span>"
        f"<span class='eval-chip'>Estado: {quiz.get('status','')}</span>"
        f"<span class='eval-chip'>Cohorte: {quiz.get('cohort') or academic_context.get('cohort')}</span>"
        f"<span class='eval-chip'>Tema: {quiz.get('source_topic') or 'Adaptativo'}</span></div>",
        unsafe_allow_html=True,
    )

    if quiz.get("status") == "completed":
        score = float(quiz.get("score") or 0)
        cls = "eval-result-pass" if quiz.get("passed") else "eval-result-fail"
        st.markdown(f"<div class='{cls}'>Resultado: {score:.0f}% · {'Aprobado' if quiz.get('passed') else 'No aprobado'} · {quiz.get('recommendation') or ''}</div>", unsafe_allow_html=True)
        for q in questions:
            icon = "✅" if q.get("is_correct") else "❌"
            st.markdown(f"<div class='eval-question'><b>{icon} {q.get('question')}</b><br><span class='small'>Tu respuesta: {q.get('user_answer') or 'Sin respuesta'} · Correcta: {q.get('correct_answer')}</span><br><br>{q.get('explanation') or ''}</div>", unsafe_allow_html=True)
        return

    with st.form(f"adaptive_quiz_form_{quiz_id}"):
        answers = {}
        for i, q in enumerate(questions, start=1):
            opts = q.get("options") or _safe_json_loads(q.get("options_json"), [])
            st.markdown(f"<div class='eval-question'><b>{i}. {q.get('question')}</b><br><span class='eval-chip'>{q.get('bloom_level') or 'Aplicar'}</span></div>", unsafe_allow_html=True)
            answers[str(q["id"])] = st.radio("Selecciona una respuesta", opts, key=f"adaptive_q_{quiz_id}_{q['id']}", label_visibility="collapsed")
        submitted = st.form_submit_button("Calificar evaluación", use_container_width=True)

    if submitted:
        score, passed, weak_topics, recommendation = grade_adaptive_quiz(int(quiz_id), answers)
        event_type = f"{quiz_type}_submitted"
        gps_actual = capture_browser_gps(user['id'], page='evaluacion_ia', topic=topic, subtopic=subtopic, event_type=event_type, show_status=False)
        log_location_event(user['id'], event_type, page='evaluacion_ia', topic=topic, subtopic=subtopic, gps=gps_actual)
        if passed:
            st.success(f"Resultado: {score:.0f}%. Aprobaste la evaluación.")
        else:
            if research_mode:
                st.warning(f"Resultado registrado: {score:.0f}%.")
            else:
                st.warning(f"Resultado: {score:.0f}%. No alcanzaste el 70%. Se habilitó el reintento automático.")
        st.rerun()


# ---------------- Sidebar SaaS PRO por roles ----------------
def sidebar_brand(role_label: str, subtitle: str):
    st.sidebar.markdown(f"""
    <div class='saas-sidebar-brand'>
        <div class='logo-row'>
            <div class='logo-mini'>BΣ</div>
            <div>
                <div class='title'>BunsekiChat</div>
                <div class='subtitle'>{subtitle}</div>
            </div>
        </div>
        <span class='saas-role-chip'>{role_label}</span>
    </div>
    """, unsafe_allow_html=True)

def render_student_sidebar(prof):
    sidebar_brand("Estudiante", "Tutor matemático universitario")
    st.sidebar.markdown("<div class='saas-section-title'>Navegación</div>", unsafe_allow_html=True)
    page = st.sidebar.radio(
        "Menú del estudiante",
        ["💬 Tutor", "🧠 Evaluación IA", "📘 Descargar aprendizaje", "🎮 Prueba de nivel", "📈 Mi progreso"],
        label_visibility="collapsed",
        key="student_nav"
    )
    st.sidebar.markdown("<div class='saas-section-title'>Configuración académica</div>", unsafe_allow_html=True)
    topic = st.sidebar.selectbox("Área", list(TOPICS.keys()), key="student_topic")
    subtopic = st.sidebar.selectbox("Subtema", TOPICS[topic], key="student_subtopic")
    level = prof.get("level") or "Inicial"
    academic_context = academic_context_from_profile(prof)
    st.sidebar.markdown(
        f"<div class='saas-help-card'><b>Nivel actual:</b> {level}<br><b>Materia:</b> {academic_context.get('subject') or 'Sin materia'}<br><b>Curso:</b> {academic_context.get('cohort') or 'Sin curso'}</div>",
        unsafe_allow_html=True
    )
    st.sidebar.markdown("<div class='saas-section-title'>Sesión</div>", unsafe_allow_html=True)
    if st.sidebar.button("Cerrar sesión", key="student_logout"):
        st.session_state.clear()
        st.rerun()
    return page, topic, subtopic, level

def render_teacher_sidebar():
    sidebar_brand("Docente / Admin", "Analítica y exportación")
    st.sidebar.markdown("<div class='saas-section-title'>Navegación</div>", unsafe_allow_html=True)
    page = st.sidebar.radio(
        "Menú docente",
        ["📊 Dashboard docente PRO", "📚 Planes analíticos", "🧪 Investigación pre/post", "📥 Exportación", "👤 Seguimiento individual", "🧾 Datos completos", "⚙️ Configuración"],
        label_visibility="collapsed",
        key="teacher_nav"
    )
    st.sidebar.markdown("<div class='saas-help-card'>Panel exclusivo para rol docente/admin. Exporta reportes por estudiante en CSV, Word y PDF.</div>", unsafe_allow_html=True)
    if st.sidebar.button("Cerrar sesión", key="teacher_logout"):
        st.session_state.clear()
        st.rerun()
    return page

# ---------------- Pages ----------------
def login_page():
    st.markdown("<div class='login-premium-page'>", unsafe_allow_html=True)
    left, center, right = st.columns([1.05, 0.82, 1.05], gap="large")

    with left:
        st.markdown("""
        <div class='ad-panel'>
            <div>
                <strong>Publicidad</strong>
                <span>Banner institucional<br>becas, cursos o aliados<br>320 x 360 px</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

    with center:
        st.markdown("""
        <div class='login-card-premium'>
            <div class='logo-orb'>BΣ</div>
            <div class='login-title'>BunsekiChat</div>
            <div class='login-subtitle'>Acceso académico inteligente</div>
        """, unsafe_allow_html=True)

        tab1, tab2 = st.tabs(["Ingresar", "Registrarse"])

        with tab1:
            u = st.text_input("Usuario", key='lu', placeholder="usuario")
            p = st.text_input("Contraseña", type='password', key='lp', placeholder="contraseña")
            if st.button("LOGIN", use_container_width=True):
                user = authenticate(u, p)
                if user:
                    st.session_state.user = user
                    st.session_state.last_activity = datetime.now()
                    st.rerun()
                else:
                    st.error("Credenciales inválidas o usuario inactivo.")

        with tab2:
            u = st.text_input("Nuevo usuario", key='ru', placeholder="mínimo 4 caracteres")
            p = st.text_input("Nueva contraseña", type='password', key='rp', placeholder="mínimo 6 caracteres")
            if st.button("Crear cuenta", use_container_width=True):
                if len(u) < 4 or len(p) < 6:
                    st.warning("Usuario mínimo 4 caracteres y contraseña mínimo 6.")
                elif user_by_username(u):
                    st.warning("Ese usuario ya existe.")
                else:
                    create_user(u, p)
                    st.success("Cuenta creada. Ahora ingresa.")

        st.markdown("""
            <div class='login-note'>BunsekiChat · Matemáticas con seguimiento académico</div>
            <div class='powered'>Powered by <strong>Christian Tomalá</strong></div>
        </div>
        """, unsafe_allow_html=True)

    with right:
        st.markdown("""
        <div class='ad-panel'>
            <div>
                <strong>Publicidad</strong>
                <span>Recursos educativos<br>eventos universitarios<br>320 x 360 px</span>
            </div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

def profile_page(user):
    prof=get_profile(user['id'])
    academic_defaults = academic_context_from_profile(prof)
    gps_actual = capture_browser_gps(user['id'], page='Ficha inicial', event_type='app_open', show_status=True)
    st.markdown("<div class='card'><h3>Ficha inicial del estudiante</h3><p class='small'>Registra datos oficiales. La app normaliza nombres, apellidos, materia, curso, paralelo y jornada para la investigación y los reportes.</p></div>", unsafe_allow_html=True)
    with st.form("perfil"):
        c1,c2,c3=st.columns(3)
        data={}
        with c1:
            data['first_names']=st.text_input("Nombres", prof.get('first_names') or '')
            data['last_names']=st.text_input("Apellidos", prof.get('last_names') or '')
            data['birth_date']=str(st.date_input("Fecha de nacimiento", value=birth_date_value(prof.get('birth_date'))))
            data['subject']=academic_selectbox("Materia", ACADEMIC_SUBJECTS, academic_defaults["subject"], "profile_subject")
        with c2:
            data['course_level']=academic_selectbox("Curso / nivel", ACADEMIC_COURSE_LEVELS, academic_defaults["course_level"], "profile_course_level", normalize_course_level)
            data['parallel']=academic_selectbox("Paralelo", ACADEMIC_PARALLELS, academic_defaults["parallel"], "profile_parallel", normalize_parallel)
            data['shift']=academic_selectbox("Jornada", ACADEMIC_SHIFTS, academic_defaults["shift"], "profile_shift", normalize_shift)
            data['teacher']=st.text_input("Docente", prof.get('teacher') or '')
            data['phone']=st.text_input("Número celular", prof.get('phone') or '')
        with c3:
            data['province']=st.text_input("Provincia", prof.get('province') or '')
            data['city']=st.text_input("Ciudad", prof.get('city') or '')
            data['canton']=st.text_input("Cantón", prof.get('canton') or '')
            data['address']=st.text_area("Dirección", prof.get('address') or '')
            auto_lat = (gps_actual or {}).get('lat') if gps_actual else None
            auto_lon = (gps_actual or {}).get('lon') if gps_actual else None
            data['gps_lat']=st.number_input("GPS latitud", value=float(auto_lat if auto_lat is not None else (prof.get('gps_lat') or 0.0)), format="%.6f")
            data['gps_lon']=st.number_input("GPS longitud", value=float(auto_lon if auto_lon is not None else (prof.get('gps_lon') or 0.0)), format="%.6f")
            data['gps_accuracy']=str((gps_actual or {}).get('accuracy') or prof.get('gps_accuracy') or '')
        submitted=st.form_submit_button("Guardar ficha")
        if submitted:
            data['first_names'] = normalize_person_name(data.get('first_names'))
            data['last_names'] = normalize_person_name(data.get('last_names'))
            data['subject'] = normalize_academic_text(data.get('subject'))
            data['course_level'] = normalize_course_level(data.get('course_level'))
            data['parallel'] = normalize_parallel(data.get('parallel'))
            data['shift'] = normalize_shift(data.get('shift'))
            data['cohort'] = build_course_label(data['course_level'], data['parallel'], data['shift'])
            data['course'] = data['cohort']
            data['full_name_normalized'] = normalize_spaces(f"{data['first_names']} {data['last_names']}")
            if not data['first_names'] or not data['last_names']:
                st.warning("Ingresa nombres y apellidos oficiales del estudiante.")
            elif not data['subject'] or not data['course_level'] or not data['parallel'] or not data['shift']:
                st.warning("Completa materia, curso, paralelo y jornada.")
            else:
                data['profile_completed']=1
                update_profile(user['id'], data)
                st.success(f"Ficha guardada: {data['full_name_normalized']} · {data['subject']} · {data['cohort']}")
                st.rerun()
    st.info("GPS automático activado con consentimiento del navegador. En producción debe ejecutarse en HTTPS; en desarrollo funciona en localhost.")

def student_page(user):
    prof=get_profile(user['id'])
    if not prof.get('profile_completed'): return profile_page(user)
    academic_context = academic_context_from_profile(prof)
    hero_html = (
        "<div class='hero'>"
        + bunseki_logo_html(APP_NAME, 'Bienvenido/a, ' + str(prof.get('first_names') or '') + ' · ' + academic_context.get('subject','Matemáticas universitarias'))
        + f"<div style='margin-top:12px'><span class='badge'>Nivel: {prof.get('level','Inicial')}</span> <span class='badge'>{academic_context.get('cohort','Sin curso')}</span></div>"
        + "</div>"
    )
    st.markdown(hero_html, unsafe_allow_html=True)
    page, topic, subtopic, level = render_student_sidebar(prof)
    gps_actual = capture_browser_gps(user['id'], page=page, topic=topic, subtopic=subtopic, event_type='app_open', show_status=False)
    rows=interactions(user['id']); user_q=[r for r in rows if r['role']=='user']
    m1,m2,m3,m4=st.columns(4)
    m1.markdown(f"<div class='metric'><b>{len(user_q)}</b><br><span>Consultas</span></div>", unsafe_allow_html=True)
    m2.markdown(f"<div class='metric'><b>{len(set([r['topic'] for r in user_q if r['topic']]))}</b><br><span>Áreas trabajadas</span></div>", unsafe_allow_html=True)
    m3.markdown(f"<div class='metric'><b>{sum(r['tokens_est'] for r in rows)}</b><br><span>Tokens estimados</span></div>", unsafe_allow_html=True)
    m4.markdown(f"<div class='metric'><b>{level}</b><br><span>Nivel actual</span></div>", unsafe_allow_html=True)

    if page == "🧠 Evaluación IA":
        render_student_adaptive_evaluation(user, prof, topic, subtopic)
        return

    if page == "📈 Mi progreso":
        st.markdown("<div class='card'><h3>📈 Mi progreso académico</h3><p class='small'>Resumen de tus consultas, áreas trabajadas y evolución.</p></div>", unsafe_allow_html=True)
        if user_q:
            df_prog = pd.DataFrame(user_q)
            c1, c2 = st.columns(2)
            c1.plotly_chart(px.histogram(df_prog, x='topic', title='Mis consultas por área'), use_container_width=True)
            c2.plotly_chart(px.histogram(df_prog, x='created_at', title='Mi actividad en el tiempo'), use_container_width=True)
            st.dataframe(df_prog, use_container_width=True, height=320)
        else:
            st.info("Aún no tienes consultas registradas.")
        return

    # Descarga individual para estudiantes: evidencia de lo investigado y aprendido
    st.markdown("""
    <div class='student-export-card'>
        <h3>📘 Descargar mi aprendizaje</h3>
        <p>Genera un reporte personal con tus consultas, respuestas recibidas, temas trabajados y trazabilidad académica.</p>
    </div>
    """, unsafe_allow_html=True)
    exp1, exp2, exp3 = st.columns(3)
    profile_export, report_df_export = build_student_summary(user['id'])
    exp1.download_button(
        "📊 Descargar CSV",
        report_df_export.to_csv(index=False).encode("utf-8-sig"),
        f"mi_aprendizaje_{profile_export.get('username','estudiante')}.csv",
        "text/csv",
        use_container_width=True,
    )
    try:
        exp2.download_button(
            "📝 Descargar Word",
            export_student_docx(user['id']),
            f"mi_aprendizaje_{profile_export.get('username','estudiante')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    except Exception as e:
        exp2.warning(f"Word no disponible: {e}")
    try:
        exp3.download_button(
            "📄 Descargar PDF",
            export_student_pdf(user['id']),
            f"mi_aprendizaje_{profile_export.get('username','estudiante')}.pdf",
            "application/pdf",
            use_container_width=True,
        )
    except Exception as e:
        exp3.warning(f"PDF no disponible: {e}")

    if page == "📘 Descargar aprendizaje":
        st.markdown("<div class='card'><h3>Vista previa de mi aprendizaje</h3></div>", unsafe_allow_html=True)
        st.dataframe(report_df_export, use_container_width=True, height=360)
        return

    left,right=st.columns([1.5,1])
    with left:
        st.markdown("<div class='card'><h3>💬 Tutor matemático</h3><p class='small'>Sugerencias para preguntar mejor:</p></div>", unsafe_allow_html=True)
        with st.expander("💡 Ver sugerencias de prompt", expanded=False):
            for sug in prompt_suggestions(topic, subtopic, level):
                st.markdown(f"- {sug}")
        for r in rows[-12:]: render_message(r['role'], r['clean_message'] or r['message'])
        q=st.chat_input("Escribe tu pregunta de matemáticas...")
        if q:
            if contains_bad_language(q):
                warning = "Tu consulta contiene lenguaje inapropiado. Reformúlala con respeto para poder ayudarte mejor."
                gps_actual = capture_browser_gps(user['id'], page=page, topic=topic, subtopic=subtopic, event_type='interaction', show_status=False)
                log_location_event(user['id'], 'chat_message_blocked', page=page, topic=topic, subtopic=subtopic, gps=gps_actual)
                log_interaction(user['id'],'user',q,topic,subtopic,level,gps=gps_actual)
                log_interaction(user['id'],'assistant',warning,topic,subtopic,level,"moderacion",0,gps=gps_actual)
                st.rerun()
            gps_actual = capture_browser_gps(user['id'], page=page, topic=topic, subtopic=subtopic, event_type='interaction', show_status=False)
            log_location_event(user['id'], 'chat_message', page=page, topic=topic, subtopic=subtopic, gps=gps_actual)
            log_interaction(user['id'],'user',q,topic,subtopic,level,gps=gps_actual)
            ans,model,ms=ask_ai(q,topic,subtopic,level,prof)
            log_interaction(user['id'],'assistant',ans,topic,subtopic,level,model,ms,gps=gps_actual)
            st.rerun()
    with right:
        quiz_items = adaptive_quiz_for_user(user['id'], topic)
        st.markdown(
            "<div class='quiz-card-premium'><h3>🎮 Prueba adaptativa para subir de nivel</h3>"
            "<p class='small'>10 preguntas según tu historial de consultas. Aprueba con 70% o más.</p>"
            f"<span class='quiz-topic-chip'>Base actual: {topic}</span></div>",
            unsafe_allow_html=True,
        )
        with st.form("quiz"):
            answers = {}
            for i, (q_topic, qdata) in enumerate(quiz_items):
                text, opts, correct = qdata
                answers[str(i)] = st.radio(f"{i+1}. [{q_topic}] {text}", opts, key=f"q{i}")
            if st.form_submit_button("Calificar prueba"):
                total = max(1, len(quiz_items))
                score = sum(
                    1 for i, (q_topic, qdata) in enumerate(quiz_items)
                    if answers[str(i)] == qdata[1][qdata[2]]
                ) / total * 100
                passed = score >= 70
                old = prof.get('level') or 'Inicial'
                new = old
                if passed:
                    idx = min(LEVELS.index(old) + 1, len(LEVELS) - 1)
                    new = LEVELS[idx]
                    update_profile(user['id'], {'level': new})
                gps_actual = capture_browser_gps(user['id'], page=page, topic=topic, subtopic=subtopic, event_type='quiz_submitted', show_status=False)
                log_location_event(user['id'], 'quiz_submitted', page=page, topic=topic, subtopic=subtopic, gps=gps_actual)
                save_quiz(user['id'], topic, old, new, score, passed, answers)
                st.success(f"Puntaje: {score:.0f}%. {'Subiste de nivel.' if passed else 'Sigue practicando.'}")
        df = pd.DataFrame(user_q)
        if not df.empty:
            fig = px.histogram(df, x='topic', title='Consultas por área')
            st.plotly_chart(fig, use_container_width=True)


def render_teacher_gps_heatmap(logs: pd.DataFrame, filtered_students: pd.DataFrame | None = None):
    """Mapa de calor docente basado en las coordenadas GPS guardadas en interactions.
    Usa únicamente registros con gps_lat/gps_lon válidos y permite respetar los filtros
    académicos activos del dashboard.
    """
    st.markdown("""
    <div class='teacher-card'>
        <h3>🗺️ Mapa de calor GPS por interacciones</h3>
        <p class='small'>Visualiza la concentración geográfica de uso de BunsekiChat según las coordenadas guardadas en cada interacción.</p>
    </div>
    """, unsafe_allow_html=True)

    if logs is None or logs.empty:
        st.info("Aún no existen interacciones para generar el mapa de calor GPS.")
        return

    gps_df = logs.copy()

    # Respeta los filtros de curso/docente/nivel aplicados al dashboard.
    if filtered_students is not None and not filtered_students.empty and "id" in filtered_students.columns:
        allowed_ids = set(filtered_students["id"].dropna().astype(int).tolist())
        gps_df = gps_df[gps_df["user_id"].isin(allowed_ids)]

    # El mapa se construye desde la tabla interactions, no desde el registro manual del perfil.
    if "role" in gps_df.columns:
        gps_df = gps_df[gps_df["role"].astype(str).str.lower().eq("user")]

    required_cols = {"gps_lat", "gps_lon"}
    if not required_cols.issubset(set(gps_df.columns)):
        st.warning("La tabla interactions todavía no tiene columnas gps_lat/gps_lon disponibles.")
        return

    gps_df["gps_lat"] = pd.to_numeric(gps_df["gps_lat"], errors="coerce")
    gps_df["gps_lon"] = pd.to_numeric(gps_df["gps_lon"], errors="coerce")
    gps_df = gps_df.dropna(subset=["gps_lat", "gps_lon"])
    gps_df = gps_df[gps_df["gps_lat"].between(-90, 90) & gps_df["gps_lon"].between(-180, 180)]

    if gps_df.empty:
        st.warning("No hay interacciones con GPS válido. Verifica que el navegador haya autorizado la ubicación y que se esté ejecutando en localhost o HTTPS.")
        return

    # Columnas auxiliares para un hover docente más claro.
    gps_df["estudiante"] = gps_df.get("username", "Sin usuario")
    if "first_names" in gps_df.columns or "last_names" in gps_df.columns:
        gps_df["nombre"] = (
            gps_df.get("first_names", "").fillna("").astype(str).str.strip()
            + " "
            + gps_df.get("last_names", "").fillna("").astype(str).str.strip()
        ).str.strip()
        gps_df.loc[gps_df["nombre"].eq(""), "nombre"] = gps_df["estudiante"]
    else:
        gps_df["nombre"] = gps_df["estudiante"]

    center_lat = float(gps_df["gps_lat"].mean())
    center_lon = float(gps_df["gps_lon"].mean())
    zoom = 11 if len(gps_df) <= 10 else 10

    h1, h2, h3 = st.columns(3)
    h1.markdown(f"<div class='metric'><b>{len(gps_df)}</b><br><span>Interacciones con GPS</span></div>", unsafe_allow_html=True)
    h2.markdown(f"<div class='metric'><b>{gps_df['user_id'].nunique()}</b><br><span>Estudiantes geolocalizados</span></div>", unsafe_allow_html=True)
    last_gps = gps_df["created_at"].max() if "created_at" in gps_df.columns else "N/D"
    h3.markdown(f"<div class='metric'><b>{last_gps}</b><br><span>Último GPS registrado</span></div>", unsafe_allow_html=True)

    fig_heat = px.density_mapbox(
        gps_df,
        lat="gps_lat",
        lon="gps_lon",
        radius=28,
        center={"lat": center_lat, "lon": center_lon},
        zoom=zoom,
        mapbox_style="open-street-map",
        hover_name="nombre",
        hover_data={
            "estudiante": True,
            "topic": True,
            "subtopic": True,
            "created_at": True,
            "gps_lat": ":.6f",
            "gps_lon": ":.6f",
            "gps_accuracy": True if "gps_accuracy" in gps_df.columns else False,
        },
        title="Mapa de calor de uso por coordenadas GPS de interacciones",
    )
    fig_heat.update_layout(margin={"r":0,"t":45,"l":0,"b":0}, height=560)
    st.plotly_chart(fig_heat, use_container_width=True)

    fig_points = px.scatter_mapbox(
        gps_df,
        lat="gps_lat",
        lon="gps_lon",
        color="topic" if "topic" in gps_df.columns else None,
        hover_name="nombre",
        hover_data={
            "estudiante": True,
            "topic": True if "topic" in gps_df.columns else False,
            "subtopic": True if "subtopic" in gps_df.columns else False,
            "created_at": True if "created_at" in gps_df.columns else False,
            "gps_lat": ":.6f",
            "gps_lon": ":.6f",
        },
        center={"lat": center_lat, "lon": center_lon},
        zoom=zoom,
        mapbox_style="open-street-map",
        title="Puntos GPS individuales por interacción",
    )
    fig_points.update_layout(margin={"r":0,"t":45,"l":0,"b":0}, height=500)
    with st.expander("Ver puntos GPS individuales"):
        st.plotly_chart(fig_points, use_container_width=True)
        cols = [c for c in ["id", "user_id", "username", "first_names", "last_names", "course", "teacher", "topic", "subtopic", "created_at", "gps_lat", "gps_lon", "gps_accuracy", "gps_source"] if c in gps_df.columns]
        st.dataframe(gps_df[cols].sort_values("created_at", ascending=False) if "created_at" in cols else gps_df[cols], use_container_width=True, height=280)
        st.download_button(
            "📍 Descargar interacciones con GPS",
        gps_df[cols].to_csv(index=False).encode("utf-8-sig"),
            "interacciones_gps_mapa_calor.csv",
            "text/csv",
            use_container_width=True,
        )


def get_research_results():
    return fetchall("""
        SELECT
            aq.id, aq.user_id, aq.quiz_type, aq.title, aq.subject,
            aq.course_level, aq.parallel, aq.shift, aq.cohort,
            aq.score, aq.passed, aq.question_count, aq.status,
            aq.created_at, aq.completed_at, aq.research_title, aq.random_seed,
            u.username,
            p.first_names, p.last_names, p.full_name_normalized,
            COALESCE(aq.subject, p.subject) AS profile_subject,
            COALESCE(aq.course_level, p.course_level) AS profile_course_level,
            COALESCE(aq.parallel, p.parallel) AS profile_parallel,
            COALESCE(aq.shift, p.shift) AS profile_shift,
            COALESCE(aq.cohort, p.cohort, p.course) AS profile_cohort
        FROM adaptive_quizzes aq
        JOIN users u ON u.id=aq.user_id
        LEFT JOIN profiles p ON p.user_id=aq.user_id
        WHERE COALESCE(aq.quiz_type,'adaptive') IN ('pretest','posttest')
        ORDER BY aq.id DESC
    """)


def render_teacher_research_dashboard(user):
    st.markdown("""
    <div class='teacher-card'>
        <h3>🧪 Investigación pretest/posttest</h3>
        <p class='small'>Resultados parametrizados por materia, curso, paralelo y jornada para análisis del efecto del tutor inteligente.</p>
    </div>
    """, unsafe_allow_html=True)

    df = pd.DataFrame(get_research_results())
    if df.empty:
        st.info("Aún no existen pretest o posttest registrados.")
        return

    df["score"] = pd.to_numeric(df.get("score"), errors="coerce")
    df["student"] = df.apply(
        lambda r: r.get("full_name_normalized") or f"{r.get('first_names') or ''} {r.get('last_names') or ''}".strip() or r.get("username"),
        axis=1,
    )
    for col in ["profile_subject", "profile_course_level", "profile_parallel", "profile_shift", "profile_cohort", "quiz_type", "status"]:
        if col not in df.columns:
            df[col] = ""

    def opts(col):
        values = sorted([str(x).strip() for x in df[col].dropna().unique() if str(x).strip()])
        return ["Todos"] + values

    f1, f2, f3, f4, f5 = st.columns(5)
    sel_subject = f1.selectbox("Materia", opts("profile_subject"), key="research_subject_filter")
    sel_course = f2.selectbox("Curso", opts("profile_course_level"), key="research_course_filter")
    sel_parallel = f3.selectbox("Paralelo", opts("profile_parallel"), key="research_parallel_filter")
    sel_shift = f4.selectbox("Jornada", opts("profile_shift"), key="research_shift_filter")
    sel_type = f5.selectbox("Instrumento", ["Todos", "pretest", "posttest"], key="research_type_filter")

    filtered = df.copy()
    if sel_subject != "Todos": filtered = filtered[filtered["profile_subject"].astype(str) == sel_subject]
    if sel_course != "Todos": filtered = filtered[filtered["profile_course_level"].astype(str) == sel_course]
    if sel_parallel != "Todos": filtered = filtered[filtered["profile_parallel"].astype(str) == sel_parallel]
    if sel_shift != "Todos": filtered = filtered[filtered["profile_shift"].astype(str) == sel_shift]
    if sel_type != "Todos": filtered = filtered[filtered["quiz_type"].astype(str) == sel_type]

    completed = filtered[filtered["status"].astype(str).eq("completed")].copy()
    pre_avg = completed.loc[completed["quiz_type"].eq("pretest"), "score"].mean()
    post_avg = completed.loc[completed["quiz_type"].eq("posttest"), "score"].mean()
    latest = completed.sort_values(["user_id", "quiz_type", "id"]).drop_duplicates(["user_id", "quiz_type"], keep="last")
    pivot = latest.pivot_table(index="user_id", columns="quiz_type", values="score", aggfunc="last")
    gain = (pivot["posttest"] - pivot["pretest"]).mean() if {"pretest", "posttest"}.issubset(set(pivot.columns)) else None

    m1, m2, m3, m4 = st.columns(4)
    m1.markdown(f"<div class='metric'><b>{filtered['user_id'].nunique()}</b><br><span>Estudiantes</span></div>", unsafe_allow_html=True)
    m2.markdown(f"<div class='metric'><b>{0 if pd.isna(pre_avg) else pre_avg:.1f}%</b><br><span>Promedio pretest</span></div>", unsafe_allow_html=True)
    m3.markdown(f"<div class='metric'><b>{0 if pd.isna(post_avg) else post_avg:.1f}%</b><br><span>Promedio posttest</span></div>", unsafe_allow_html=True)
    m4.markdown(f"<div class='metric'><b>{'N/D' if gain is None or pd.isna(gain) else f'{gain:.1f} pts'}</b><br><span>Ganancia media</span></div>", unsafe_allow_html=True)

    if not completed.empty:
        fig = px.box(completed, x="quiz_type", y="score", color="profile_cohort", points="all", title="Distribución de resultados pre/post")
        fig.update_layout(height=430, margin={"r":10,"t":60,"l":10,"b":80})
        st.plotly_chart(fig, use_container_width=True)

    view_cols = [c for c in [
        "id", "student", "username", "quiz_type", "score", "passed", "status",
        "profile_subject", "profile_cohort", "profile_course_level", "profile_parallel", "profile_shift",
        "question_count", "created_at", "completed_at", "random_seed"
    ] if c in filtered.columns]
    st.dataframe(filtered[view_cols], use_container_width=True, height=420)
    st.download_button(
        "Descargar CSV investigación",
        filtered[view_cols].to_csv(index=False).encode("utf-8-sig"),
        "investigacion_pretest_posttest_bunsekichat.csv",
        "text/csv",
        use_container_width=True,
    )

def admin_page(user):
    teacher_page = render_teacher_sidebar()

    # CSS específico para que el dashboard docente no se corte y ocupe todo el ancho.
    st.markdown("""
    <style>
    .block-container{
        max-width:100% !important;
        padding-left:2rem !important;
        padding-right:2rem !important;
        padding-bottom:5rem !important;
    }
    div[data-testid="stVerticalBlock"]{overflow:visible !important;}
    .stPlotlyChart{width:100% !important; overflow:visible !important;}
    div[data-testid="stDataFrame"]{width:100% !important;}
    </style>
    """, unsafe_allow_html=True)

    st.markdown("<div class='hero'>" + bunseki_logo_html("Dashboard docente PRO", "Analítica, trazabilidad, exportación y seguimiento académico") + "</div>", unsafe_allow_html=True)

    users, logs, quizzes = get_teacher_tables()
    students = users[users["role"].astype(str).str.lower().eq("student")].copy() if not users.empty and "role" in users.columns else pd.DataFrame()

    # Configuración solo cuando se elige el menú correspondiente
    if teacher_page == "⚙️ Configuración":
        st.markdown("<div class='teacher-card'><h3>⚙️ Configuración</h3><p class='small'>Ajustes generales de sesión y parámetros administrativos.</p></div>", unsafe_allow_html=True)
        timeout = st.number_input("Inactividad (minutos)", 5, 180, get_timeout())
        if st.button("Guardar configuración", use_container_width=True):
            set_timeout(int(timeout)); st.success("Configuración guardada.")
        return

    if teacher_page == "📚 Planes analíticos":
        render_teacher_plan_manager(user)
        return

    if teacher_page == "🧪 Investigación pre/post":
        render_teacher_research_dashboard(user)
        return

    total_students = len(students)
    user_logs_all = logs[logs["role"].astype(str).str.lower().eq("user")].copy() if not logs.empty and "role" in logs.columns else pd.DataFrame()
    total_questions = len(user_logs_all)
    active_students = user_logs_all["user_id"].nunique() if not user_logs_all.empty and "user_id" in user_logs_all.columns else 0

    if not quizzes.empty and "score" in quizzes.columns:
        quizzes["score"] = pd.to_numeric(quizzes["score"], errors="coerce")
        avg_score = quizzes["score"].dropna().mean()
        if pd.isna(avg_score):
            avg_score = 0
    else:
        avg_score = 0

    k1,k2,k3,k4 = st.columns(4)
    k1.markdown(f"<div class='metric'><b>{total_students}</b><br><span>Estudiantes</span></div>", unsafe_allow_html=True)
    k2.markdown(f"<div class='metric'><b>{total_questions}</b><br><span>Consultas realizadas</span></div>", unsafe_allow_html=True)
    k3.markdown(f"<div class='metric'><b>{active_students}</b><br><span>Estudiantes activos</span></div>", unsafe_allow_html=True)
    k4.markdown(f"<div class='metric'><b>{avg_score:.0f}%</b><br><span>Promedio pruebas</span></div>", unsafe_allow_html=True)

    st.markdown("<div class='teacher-card'><h3>Filtros académicos</h3></div>", unsafe_allow_html=True)
    f1,f2,f3,f4,f5,f6 = st.columns(6)

    def opts(df, col):
        if df.empty or col not in df.columns:
            return ["Todos"]
        vals = sorted([str(x).strip() for x in df[col].dropna().unique() if str(x).strip()])
        return ["Todos"] + vals

    sel_subject = f1.selectbox("Materia", opts(students, "subject"))
    sel_course = f2.selectbox("Curso", opts(students, "course"))
    sel_parallel = f3.selectbox("Paralelo", opts(students, "parallel"))
    sel_shift = f4.selectbox("Jornada", opts(students, "shift"))
    sel_teacher = f5.selectbox("Docente", opts(students, "teacher"))
    sel_level = f6.selectbox("Nivel", opts(students, "level"))

    filtered_students = students.copy()
    if not filtered_students.empty:
        if sel_subject != "Todos": filtered_students = filtered_students[filtered_students["subject"].astype(str) == sel_subject]
        if sel_course != "Todos": filtered_students = filtered_students[filtered_students["course"].astype(str) == sel_course]
        if sel_parallel != "Todos": filtered_students = filtered_students[filtered_students["parallel"].astype(str) == sel_parallel]
        if sel_shift != "Todos": filtered_students = filtered_students[filtered_students["shift"].astype(str) == sel_shift]
        if sel_teacher != "Todos": filtered_students = filtered_students[filtered_students["teacher"].astype(str) == sel_teacher]
        if sel_level != "Todos": filtered_students = filtered_students[filtered_students["level"].astype(str) == sel_level]

    allowed_ids = set(filtered_students["id"].dropna().astype(int).tolist()) if not filtered_students.empty and "id" in filtered_students.columns else set()

    filtered_logs = logs.copy()
    if allowed_ids and not filtered_logs.empty and "user_id" in filtered_logs.columns:
        filtered_logs = filtered_logs[filtered_logs["user_id"].isin(allowed_ids)]

    filtered_quizzes = quizzes.copy()
    if allowed_ids and not filtered_quizzes.empty and "user_id" in filtered_quizzes.columns:
        filtered_quizzes = filtered_quizzes[filtered_quizzes["user_id"].isin(allowed_ids)]

    st.markdown(f"<div class='teacher-card'><b>Vista seleccionada en menú:</b> {teacher_page}</div>", unsafe_allow_html=True)

    if teacher_page == "📊 Dashboard docente PRO":
        st.markdown("<div class='teacher-card'><h3>📊 Dashboard docente PRO</h3><p class='small'>Resumen visual del desempeño, participación, cursos y resultados de pruebas.</p></div>", unsafe_allow_html=True)

        user_logs = filtered_logs[filtered_logs["role"].astype(str).str.lower().eq("user")].copy() if not filtered_logs.empty and "role" in filtered_logs.columns else pd.DataFrame()

        if user_logs.empty:
            st.info("Aún no existen consultas registradas con los filtros seleccionados.")
        else:
            c1, c2 = st.columns(2, gap="large")
            fig1 = px.histogram(user_logs, x="topic", color="course" if "course" in user_logs.columns else None, title="Consultas por área y curso")
            fig1.update_layout(height=430, margin={"r":10,"t":60,"l":10,"b":80})
            c1.plotly_chart(fig1, use_container_width=True)

            fig2 = px.histogram(user_logs, x="username" if "username" in user_logs.columns else "user_id", title="Participación por estudiante")
            fig2.update_layout(height=430, margin={"r":10,"t":60,"l":10,"b":100})
            c2.plotly_chart(fig2, use_container_width=True)

        if not filtered_quizzes.empty and "score" in filtered_quizzes.columns:
            fig3 = px.box(filtered_quizzes, x="topic", y="score", color="course" if "course" in filtered_quizzes.columns else None, title="Distribución de puntajes por tema")
            fig3.update_layout(height=430, margin={"r":10,"t":60,"l":10,"b":80})
            st.plotly_chart(fig3, use_container_width=True)
        else:
            st.info("Aún no existen pruebas registradas con los filtros seleccionados.")

        render_teacher_gps_heatmap(filtered_logs, filtered_students)
        return

    if teacher_page == "📥 Exportación":
        st.markdown("""
        <div class='teacher-card'>
            <h3>📥 Exportación por estudiante</h3>
            <p class='small'>Botones independientes para exportar CSV, Word y PDF con las consultas reales del estudiante.</p>
        </div>
        """, unsafe_allow_html=True)
        if filtered_students.empty:
            st.warning("No hay estudiantes con los filtros seleccionados.")
            return

        filtered_students = filtered_students.copy()
        filtered_students["display"] = filtered_students.apply(lambda r: f"{r.get('username','')} · {r.get('first_names') or ''} {r.get('last_names') or ''} · {r.get('subject') or ''} · {r.get('course') or ''}", axis=1)
        selected_label = st.selectbox("Selecciona estudiante para exportar", filtered_students["display"].tolist(), key="export_student_select")
        selected_id = int(filtered_students.loc[filtered_students["display"] == selected_label, "id"].iloc[0])
        profile, report_df = build_student_summary(selected_id)

        st.markdown(
            f"<div class='teacher-card'><h3>{profile.get('first_names') or ''} {profile.get('last_names') or profile.get('username')}</h3>"
            f"<span class='teacher-chip'>Usuario: {profile.get('username') or ''}</span>"
            f"<span class='teacher-chip'>Materia: {profile.get('subject') or 'Sin materia'}</span>"
            f"<span class='teacher-chip'>Curso: {profile.get('course') or 'Sin curso'}</span>"
            f"<span class='teacher-chip'>Docente: {profile.get('teacher') or 'Sin docente'}</span>"
            f"<span class='teacher-chip'>Nivel: {profile.get('level') or 'Sin nivel'}</span>"
            f"<span class='teacher-chip'>Registros: {len(report_df)}</span></div>",
            unsafe_allow_html=True,
        )

        if report_df.empty:
            st.warning("Este estudiante no tiene consultas registradas. El archivo se generará con encabezados, pero sin filas.")
        b1, b2, b3 = st.columns(3)
        b1.download_button("📊 Exportar CSV", report_df.to_csv(index=False).encode("utf-8-sig"), f"consultas_{profile.get('username','estudiante')}_{selected_id}.csv", "text/csv", use_container_width=True)
        try:
            b2.download_button("📝 Exportar Word", export_student_docx(selected_id), f"reporte_{profile.get('username','estudiante')}_{selected_id}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        except Exception as e:
            b2.warning(f"Word no disponible: {e}")
        try:
            b3.download_button("📄 Exportar PDF", export_student_pdf(selected_id), f"reporte_{profile.get('username','estudiante')}_{selected_id}.pdf", "application/pdf", use_container_width=True)
        except Exception as e:
            b3.warning(f"PDF no disponible: {e}")

        st.markdown("<div class='teacher-card'><h3>Vista previa de consultas del estudiante</h3></div>", unsafe_allow_html=True)
        st.dataframe(report_df, use_container_width=True, height=420)
        return

    if teacher_page == "👤 Seguimiento individual":
        st.markdown("<div class='teacher-card'><h3>👤 Seguimiento individual</h3><p class='small'>Revisa el comportamiento académico de cada estudiante antes de exportar.</p></div>", unsafe_allow_html=True)
        if filtered_students.empty:
            st.warning("No hay estudiantes con los filtros seleccionados.")
            return

        filtered_students = filtered_students.copy()
        filtered_students["display"] = filtered_students.apply(lambda r: f"{r.get('username','')} · {r.get('first_names') or ''} {r.get('last_names') or ''} · {r.get('subject') or ''} · {r.get('course') or ''}", axis=1)
        selected_label = st.selectbox("Selecciona estudiante para seguimiento", filtered_students["display"].tolist(), key="follow_student_select")
        selected_id = int(filtered_students.loc[filtered_students["display"] == selected_label, "id"].iloc[0])
        profile, report_df = build_student_summary(selected_id)
        st.markdown(
            f"<div class='teacher-card'><h3>{profile.get('first_names') or ''} {profile.get('last_names') or profile.get('username')}</h3>"
            f"<span class='teacher-chip'>Usuario: {profile.get('username') or ''}</span>"
            f"<span class='teacher-chip'>Materia: {profile.get('subject') or 'Sin materia'}</span>"
            f"<span class='teacher-chip'>Curso: {profile.get('course') or 'Sin curso'}</span>"
            f"<span class='teacher-chip'>Nivel: {profile.get('level') or 'Sin nivel'}</span></div>",
            unsafe_allow_html=True,
        )
        if report_df.empty:
            st.info("Este estudiante aún no registra consultas.")
        else:
            g1, g2 = st.columns(2, gap="large")
            user_df = report_df[report_df['role'].astype(str).str.lower().eq('user')].copy()
            if not user_df.empty:
                fig_a = px.histogram(user_df, x='topic', title='Temas consultados')
                fig_a.update_layout(height=420, margin={"r":10,"t":60,"l":10,"b":80})
                g1.plotly_chart(fig_a, use_container_width=True)
            fig_b = px.histogram(report_df, x='created_at', title='Actividad en el tiempo')
            fig_b.update_layout(height=420, margin={"r":10,"t":60,"l":10,"b":80})
            g2.plotly_chart(fig_b, use_container_width=True)
            st.dataframe(report_df, use_container_width=True, height=380)
        return

    if teacher_page == "🧾 Datos completos":
        st.markdown("<div class='teacher-card'><h3>🧾 Datos completos</h3><p class='small'>Tablas completas y descargas generales.</p></div>", unsafe_allow_html=True)
        st.subheader("Estudiantes")
        st.dataframe(filtered_students.drop(columns=["display"], errors="ignore"), use_container_width=True, height=300)
        st.subheader("Trazabilidad de interacciones")
        st.dataframe(filtered_logs, use_container_width=True, height=360)
        loc_df = pd.DataFrame(get_location_events())
        st.subheader("Eventos GPS para mapas de calor")
        if not loc_df.empty:
            st.dataframe(loc_df, use_container_width=True, height=300)
            st.download_button("📍 CSV eventos GPS", loc_df.to_csv(index=False).encode('utf-8-sig'), "eventos_gps_bunsekichat.csv", "text/csv", use_container_width=True)
        else:
            st.info("Aún no hay eventos GPS registrados.")
        st.subheader("Pruebas y gamificación")
        st.dataframe(filtered_quizzes, use_container_width=True, height=300)
        x1, x2, x3 = st.columns(3)
        x1.download_button("📊 CSV interacciones", filtered_logs.to_csv(index=False).encode('utf-8-sig'), "interacciones_bunsekichat.csv", "text/csv", use_container_width=True)
        x2.download_button("👥 CSV estudiantes", filtered_students.drop(columns=["display"], errors="ignore").to_csv(index=False).encode('utf-8-sig'), "estudiantes_bunsekichat.csv", "text/csv", use_container_width=True)
        x3.download_button("🎮 CSV pruebas", filtered_quizzes.to_csv(index=False).encode('utf-8-sig'), "pruebas_bunsekichat.csv", "text/csv", use_container_width=True)
        return

setup_database_once()
session_guard()

if 'user' not in st.session_state:
    login_page()
else:
    u = st.session_state.user

    # =========================================
    # GPS GLOBAL AUTOMÁTICO
    # =========================================
    capture_browser_gps(
        uid=u['id'],
        page='global',
        event_type='app_open',
        show_status=False
    )

    if u.get('role') in ['admin', 'teacher', 'docente']:
        admin_page(u)
    else:
        student_page(u)
