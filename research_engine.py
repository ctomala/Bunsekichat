"""Scientific and predictive research package for BunsekiChat."""

from __future__ import annotations

import json
import math
import os
import re
import zipfile
from datetime import datetime
from io import BytesIO

os.environ["MPLCONFIGDIR"] = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".matplotlib-cache")
os.makedirs(os.environ["MPLCONFIGDIR"], exist_ok=True)

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sklearn.compose import TransformedTargetRegressor
from sklearn.ensemble import GradientBoostingRegressor, RandomForestClassifier, RandomForestRegressor
from sklearn.linear_model import LassoCV, LinearRegression, LogisticRegression, RidgeCV
from sklearn.metrics import (
    accuracy_score, auc, confusion_matrix, f1_score, mean_absolute_error,
    mean_squared_error, precision_score, r2_score, recall_score, roc_curve,
)
from sklearn.model_selection import KFold, StratifiedKFold, cross_val_score, train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.tree import DecisionTreeClassifier

try:
    import statsmodels.api as sm
    import statsmodels.formula.api as smf
    from statsmodels.stats.anova import anova_lm
except Exception:
    sm = smf = anova_lm = None

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from research_analytics import (
    ALPHA, COLORS, _add_figure, _add_note, _add_table, _bool_series,
    _cronbach_alpha, _descriptive, _latest_completed, _number, _paired_scores,
    _paired_statistics, _pvalue, _safe_json, _shapiro,
    build_research_word_report,
)


RANDOM_STATE = 20260626
LEVEL_MAP = {"novato": 1, "inicial": 1, "intermedio": 2, "avanzado": 3}
PREDICTOR_COLUMNS = [
    "pretest_total", "tiempo_uso_total", "numero_sesiones", "numero_ejercicios",
    "retroalimentaciones_recibidas", "clases_vistas", "nivel_adaptativo",
    "motivacion", "autoeficacia", "usabilidad", "grupo_experimental",
]


def _numeric(series):
    return pd.to_numeric(series, errors="coerce")


def _json_dict(value):
    return _safe_json(value)


def _session_metrics(timestamps, gap_minutes=30):
    times = pd.to_datetime(pd.Series(timestamps), errors="coerce").dropna().sort_values()
    if times.empty:
        return 0, 0.0
    sessions, active_seconds = 1, 60.0
    previous = times.iloc[0]
    for current in times.iloc[1:]:
        gap = (current - previous).total_seconds()
        if gap > gap_minutes * 60:
            sessions += 1
            active_seconds += 60
        elif gap > 0:
            active_seconds += min(gap, 10 * 60)
        previous = current
    return sessions, active_seconds


def _safe_divide(numerator, denominator):
    try:
        numerator, denominator = float(numerator), float(denominator)
    except (TypeError, ValueError):
        return np.nan
    return numerator / denominator if np.isfinite(denominator) and denominator != 0 else np.nan


def build_scientific_dataset(research_df, exercise_df=None, survey_df=None, interaction_df=None, event_df=None, item_df=None):
    research_df = pd.DataFrame(research_df).copy()
    exercise_df = pd.DataFrame() if exercise_df is None else pd.DataFrame(exercise_df).copy()
    survey_df = pd.DataFrame() if survey_df is None else pd.DataFrame(survey_df).copy()
    interaction_df = pd.DataFrame() if interaction_df is None else pd.DataFrame(interaction_df).copy()
    event_df = pd.DataFrame() if event_df is None else pd.DataFrame(event_df).copy()
    item_df = pd.DataFrame() if item_df is None else pd.DataFrame(item_df).copy()
    latest = _latest_completed(research_df)
    paired = _paired_scores(latest)
    if latest.empty:
        return pd.DataFrame(), []

    users = sorted(set(latest["user_id"].dropna().tolist()))
    rows = []
    for uid in users:
        user_tests = latest[latest["user_id"].eq(uid)]
        pre = user_tests[user_tests["quiz_type"].eq("pretest")]
        post = user_tests[user_tests["quiz_type"].eq("posttest")]
        meta = post.iloc[-1] if not post.empty else (pre.iloc[-1] if not pre.empty else user_tests.iloc[-1])
        pre_score = float(pre.iloc[-1]["score"]) if not pre.empty and pd.notna(pre.iloc[-1]["score"]) else np.nan
        post_score = float(post.iloc[-1]["score"]) if not post.empty and pd.notna(post.iloc[-1]["score"]) else np.nan
        gain = post_score - pre_score if np.isfinite(pre_score) and np.isfinite(post_score) else np.nan
        pre_dims = _json_dict(pre.iloc[-1].get("dimension_scores_json")) if not pre.empty else {}
        post_dims = _json_dict(post.iloc[-1].get("dimension_scores_json")) if not post.empty else {}

        user_ex = exercise_df[exercise_df.get("user_id", pd.Series(dtype=float)).eq(uid)] if not exercise_df.empty else pd.DataFrame()
        user_survey = survey_df[survey_df.get("user_id", pd.Series(dtype=float)).eq(uid)] if not survey_df.empty else pd.DataFrame()
        user_interactions = interaction_df[interaction_df.get("user_id", pd.Series(dtype=float)).eq(uid)] if not interaction_df.empty else pd.DataFrame()
        user_events = event_df[event_df.get("user_id", pd.Series(dtype=float)).eq(uid)] if not event_df.empty else pd.DataFrame()
        user_items = item_df[item_df.get("user_id", pd.Series(dtype=float)).eq(uid)] if not item_df.empty else pd.DataFrame()

        timestamps = []
        for frame in (user_interactions, user_events, user_ex):
            if not frame.empty and "created_at" in frame.columns:
                timestamps.extend(frame["created_at"].tolist())
        sessions, active_seconds = _session_metrics(timestamps)
        quiz_time = _numeric(user_tests.get("total_time_seconds", pd.Series(dtype=float))).fillna(0).sum()
        event_time = _numeric(user_events.get("duration_seconds", pd.Series(dtype=float))).fillna(0).sum() if not user_events.empty else 0
        total_minutes = (active_seconds + quiz_time + event_time) / 60

        survey_means = {}
        if not user_survey.empty and {"dimension", "score"}.issubset(user_survey.columns):
            temp = user_survey.assign(score_num=_numeric(user_survey["score"]))
            survey_means = temp.groupby("dimension")["score_num"].mean().to_dict()
        survey_lookup = {re.sub(r"[^a-z]", "", str(key).lower()): value for key, value in survey_means.items()}
        motivation = np.nanmean([value for key, value in survey_lookup.items() if "motiv" in key]) if any("motiv" in key for key in survey_lookup) else np.nan
        self_efficacy = np.nanmean([value for key, value in survey_lookup.items() if "autoef" in key]) if any("autoef" in key for key in survey_lookup) else np.nan
        satisfaction = np.nanmean([value for key, value in survey_lookup.items() if "satisf" in key]) if any("satisf" in key for key in survey_lookup) else np.nan
        usability_values = [value for key, value in survey_lookup.items() if "usabil" in key or "facilidad" in key]
        usability = np.nanmean(usability_values) if usability_values else np.nan

        event_types = user_events.get("event_type", pd.Series(dtype=str)).astype(str).str.lower() if not user_events.empty else pd.Series(dtype=str)
        classes = int(event_types.str.contains("class").sum())
        feedback_count = len(user_ex)
        errors_pre = int((user_items.get("quiz_type", pd.Series(dtype=str)).eq("pretest") & ~_bool_series(user_items.get("is_correct", pd.Series(dtype=object))).fillna(False)).sum()) if not user_items.empty else np.nan
        errors_post = int((user_items.get("quiz_type", pd.Series(dtype=str)).eq("posttest") & ~_bool_series(user_items.get("is_correct", pd.Series(dtype=object))).fillna(False)).sum()) if not user_items.empty else np.nan
        cognitive = str(meta.get("cognitive_profile") or "").lower()
        group = str(meta.get("research_group") or "Sin asignar").strip() or "Sin asignar"

        row = {
            "estudiante_id": uid,
            "grupo": group,
            "grupo_experimental": 1 if group.lower() == "experimental" else 0 if group.lower() == "control" else np.nan,
            "materia": meta.get("profile_subject") or meta.get("subject"),
            "curso": meta.get("profile_course_level") or meta.get("course_level"),
            "paralelo": meta.get("profile_parallel") or meta.get("parallel"),
            "jornada": meta.get("profile_shift") or meta.get("shift"),
            "cohorte": meta.get("profile_cohort") or meta.get("cohort"),
            "pretest_total": pre_score,
            "postest_total": post_score,
            "ganancia_aprendizaje": gain,
            "ganancia_normalizada": _safe_divide(gain, 100 - pre_score) if np.isfinite(pre_score) and pre_score < 100 else np.nan,
            "tiempo_uso_total": total_minutes,
            "numero_sesiones": sessions,
            "numero_ejercicios": len(user_ex),
            "retroalimentaciones_recibidas": feedback_count,
            "clases_vistas": classes,
            "nivel_adaptativo": LEVEL_MAP.get(cognitive, np.nan),
            "perfil_cognitivo": cognitive or "Sin dato",
            "motivacion": motivation,
            "autoeficacia": self_efficacy,
            "satisfaccion": satisfaction,
            "usabilidad": usability,
            "errores_conceptuales_pre": errors_pre,
            "errores_conceptuales_post": errors_post,
        }
        dimensions = {
            "Comprensión conceptual": "comprension_conceptual",
            "Procedimientos": "procedimientos",
            "Aplicaciones": "aplicaciones",
            "Resolución de problemas": "resolucion_problemas",
        }
        post_dimension_values = []
        for label, code in dimensions.items():
            pre_value = pd.to_numeric(pre_dims.get(label), errors="coerce")
            post_value = pd.to_numeric(post_dims.get(label), errors="coerce")
            row[f"{code}_pre"] = pre_value
            row[f"{code}_post"] = post_value
            row[f"mejora_{code}"] = post_value - pre_value if pd.notna(pre_value) and pd.notna(post_value) else np.nan
            if pd.notna(post_value):
                post_dimension_values.append(float(post_value))
        row["indice_dominio"] = np.mean(post_dimension_values) if post_dimension_values else np.nan
        row["nivel_dominio"] = "Bajo" if row["indice_dominio"] < 60 else "Medio" if row["indice_dominio"] < 80 else "Alto" if pd.notna(row["indice_dominio"]) else "Sin dato"
        row["indice_reduccion_error"] = _safe_divide(errors_pre - errors_post, errors_pre) if pd.notna(errors_pre) and pd.notna(errors_post) else np.nan
        row["eficiencia_aprendizaje"] = _safe_divide(gain, total_minutes)
        row["indice_perseverancia"] = _safe_divide(len(user_ex), sessions)
        row["aprovechamiento_feedback_proxy"] = _safe_divide(gain, feedback_count)
        rows.append(row)

    dataset = pd.DataFrame(rows)
    required = [
        "estudiante_id", "grupo", "pretest_total", "postest_total", "ganancia_aprendizaje",
        "tiempo_uso_total", "numero_sesiones", "numero_ejercicios", "retroalimentaciones_recibidas",
        "clases_vistas", "nivel_adaptativo", "comprension_conceptual_pre",
        "comprension_conceptual_post", "procedimientos_pre", "procedimientos_post",
        "aplicaciones_pre", "aplicaciones_post", "resolucion_problemas_pre",
        "resolucion_problemas_post", "motivacion", "autoeficacia", "satisfaccion", "usabilidad",
    ]
    missing_columns = [column for column in required if column not in dataset.columns or dataset[column].isna().all()]
    return dataset, missing_columns


def clean_dataset(dataset):
    cleaned = dataset.copy()
    log_rows = []
    protected = {"estudiante_id", "pretest_total", "postest_total", "ganancia_aprendizaje", "grupo"}
    for column in cleaned.columns:
        missing_pct = float(cleaned[column].isna().mean() * 100) if len(cleaned) else 0
        decision = "Sin cambios"
        if missing_pct == 0:
            pass
        elif missing_pct <= 5 and column not in protected:
            if pd.api.types.is_numeric_dtype(cleaned[column]):
                values = _numeric(cleaned[column]).dropna()
                normal = _shapiro(values)
                fill = values.mean() if np.isfinite(normal.get("p", np.nan)) and normal["p"] > ALPHA else values.median()
                cleaned[column] = _numeric(cleaned[column]).fillna(fill)
                decision = f"Imputación simple ({'media' if fill == values.mean() else 'mediana'})"
            else:
                mode = cleaned[column].mode(dropna=True)
                if not mode.empty:
                    cleaned[column] = cleaned[column].fillna(mode.iloc[0])
                decision = "Imputación por moda"
        elif missing_pct <= 20:
            decision = "No imputada automáticamente; usar casos completos o imputación múltiple externa"
        else:
            decision = "Excluir del modelo principal por >20% de datos perdidos"
        z_outliers = iqr_outliers = 0
        if pd.api.types.is_numeric_dtype(cleaned[column]):
            values = _numeric(cleaned[column]).dropna()
            if len(values) >= 3 and values.std(ddof=1) > 0:
                z_outliers = int((np.abs(stats.zscore(values)) > 3).sum())
                q1, q3 = values.quantile([0.25, 0.75])
                iqr = q3 - q1
                iqr_outliers = int(((values < q1 - 1.5 * iqr) | (values > q3 + 1.5 * iqr)).sum()) if iqr > 0 else 0
        log_rows.append({"variable": column, "tipo": str(cleaned[column].dtype), "perdidos_pct": missing_pct, "decision": decision, "atipicos_z": z_outliers, "atipicos_iqr": iqr_outliers})
    numeric_columns = [column for column in cleaned.select_dtypes(include=[np.number]).columns if column != "estudiante_id" and cleaned[column].nunique(dropna=True) > 1]
    multivariate = cleaned[numeric_columns].dropna() if numeric_columns else pd.DataFrame()
    if len(numeric_columns) >= 2 and len(multivariate) > len(numeric_columns) + 5:
        centered = multivariate - multivariate.mean()
        inverse_covariance = np.linalg.pinv(np.cov(multivariate.to_numpy(), rowvar=False))
        distances = np.einsum("ij,jk,ik->i", centered.to_numpy(), inverse_covariance, centered.to_numpy())
        threshold = stats.chi2.ppf(0.999, len(numeric_columns))
        log_rows.append({"variable": "Análisis multivariado", "tipo": "Distancia de Mahalanobis", "perdidos_pct": 0.0, "decision": f"Umbral χ²(0,999; gl={len(numeric_columns)})={threshold:.2f}; revisar, no eliminar automáticamente", "atipicos_z": "N/A", "atipicos_iqr": int((distances > threshold).sum())})
    return cleaned, pd.DataFrame(log_rows)


def normality_analysis(dataset):
    rows = []
    for column in dataset.select_dtypes(include=[np.number]).columns:
        if column == "estudiante_id":
            continue
        values = _numeric(dataset[column]).dropna()
        description = _descriptive(values)
        normality = _shapiro(values)
        rows.append({
            "variable": column, "n": len(values), "prueba": normality.get("name"),
            "estadistico": normality.get("statistic"), "p": normality.get("p"),
            "decision": normality.get("status"), "asimetria": description.get("skew"),
            "curtosis": description.get("kurtosis"),
        })
    return pd.DataFrame(rows)


def _omega_total(matrix):
    complete = matrix.apply(pd.to_numeric, errors="coerce").dropna()
    if complete.shape[0] < 3 or complete.shape[1] < 3:
        return np.nan
    corr = complete.corr().to_numpy()
    if not np.isfinite(corr).all():
        return np.nan
    values, vectors = np.linalg.eigh(corr)
    first = np.argmax(values)
    loadings = vectors[:, first] * math.sqrt(max(values[first], 0))
    if loadings.sum() < 0:
        loadings *= -1
    uniqueness = np.clip(1 - loadings ** 2, 0.001, None)
    return float(loadings.sum() ** 2 / (loadings.sum() ** 2 + uniqueness.sum()))


def psychometric_analysis(item_df):
    item_df = pd.DataFrame(item_df).copy()
    reliability_rows, item_rows, matrices = [], [], {}
    if item_df.empty or not {"user_id", "quiz_type", "item_code", "is_correct"}.issubset(item_df.columns):
        return pd.DataFrame(), pd.DataFrame(), matrices
    item_df["correct"] = _bool_series(item_df["is_correct"]).astype(float)
    item_df["version_code"] = item_df.get("version_code", "Sin versión").fillna("Sin versión").astype(str)
    for (quiz_type, version), form in item_df.groupby(["quiz_type", "version_code"], dropna=False):
        matrix = form.pivot_table(index="user_id", columns="item_code", values="correct", aggfunc="last")
        complete = matrix.dropna(axis=0)
        alpha, n_complete, k = _cronbach_alpha(matrix)
        omega = _omega_total(matrix)
        reliability_rows.append({"instrumento": quiz_type, "version": version, "participantes": form["user_id"].nunique(), "casos_completos": n_complete, "items": k, "alfa_cronbach": alpha, "omega_mcdonald_aprox": omega, "cobertura_completa_pct": n_complete / max(1, form["user_id"].nunique()) * 100})
        if n_complete >= 2 and k >= 2:
            matrices[f"{quiz_type}_{version}"] = complete.corr()
        totals = form.groupby("user_id")["correct"].sum()
        for item_code, responses in form.groupby("item_code"):
            y = responses.set_index("user_id")["correct"].dropna()
            total_other = totals.reindex(y.index) - y
            if len(y) >= 3 and y.nunique() > 1 and total_other.nunique() > 1:
                discrimination = stats.pointbiserialr(y, total_other).statistic
            else:
                discrimination = np.nan
            difficulty = float(y.mean()) if len(y) else np.nan
            question = str(responses["question"].iloc[0]) if "question" in responses else item_code
            item_rows.append({
                "instrumento": quiz_type, "version": version, "item_code": item_code,
                "pregunta": question, "n": len(y), "dificultad_p": difficulty,
                "clasificacion_dificultad": "Difícil" if difficulty < 0.30 else "Adecuado" if difficulty <= 0.70 else "Fácil",
                "discriminacion_item_total": discrimination,
                "calidad_item": "Débil: revisar" if not np.isfinite(discrimination) or discrimination < 0.20 else "Aceptable",
            })
    return pd.DataFrame(reliability_rows), pd.DataFrame(item_rows), matrices


def _between_group_test(dataset, variable):
    data = dataset[["grupo", variable]].dropna()
    data = data[data["grupo"].isin(["Experimental", "Control"])]
    groups = {name: group[variable].astype(float).to_numpy() for name, group in data.groupby("grupo")}
    if set(groups) != {"Experimental", "Control"} or min(map(len, groups.values())) < 2:
        return None
    exp, control = groups["Experimental"], groups["Control"]
    normal_exp, normal_control = _shapiro(exp), _shapiro(control)
    normal = all(np.isfinite(x.get("p", np.nan)) and x["p"] >= ALPHA for x in [normal_exp, normal_control])
    levene = stats.levene(exp, control, center="median")
    if normal:
        equal = levene.pvalue >= ALPHA
        test = stats.ttest_ind(exp, control, equal_var=equal)
        name = "t independiente" if equal else "t de Welch"
        pooled = math.sqrt(((len(exp) - 1) * np.var(exp, ddof=1) + (len(control) - 1) * np.var(control, ddof=1)) / (len(exp) + len(control) - 2))
        effect = (np.mean(exp) - np.mean(control)) / pooled if pooled else np.nan
        effect_name = "Cohen d"
    else:
        test = stats.mannwhitneyu(exp, control, alternative="two-sided")
        name = "U de Mann-Whitney"
        effect = 1 - (2 * test.statistic) / (len(exp) * len(control))
        effect_name = "Correlación biserial de rangos"
    return {"comparacion": "Experimental vs Control", "variable": variable, "n_experimental": len(exp), "n_control": len(control), "media_experimental": np.mean(exp), "media_control": np.mean(control), "normalidad": normal, "levene_p": float(levene.pvalue), "prueba": name, "estadistico": float(test.statistic), "p": float(test.pvalue), "efecto": effect, "tipo_efecto": effect_name}


def group_comparisons(dataset):
    rows = []
    for variable in ["pretest_total", "postest_total", "ganancia_aprendizaje", "comprension_conceptual_post", "procedimientos_post", "aplicaciones_post", "resolucion_problemas_post"]:
        if variable in dataset.columns:
            result = _between_group_test(dataset, variable)
            if result:
                rows.append(result)
    for group_name in ["Experimental", "Control"]:
        group = dataset[dataset.get("grupo", pd.Series(dtype=str)).eq(group_name)][["pretest_total", "postest_total"]].dropna()
        if len(group) < 2:
            continue
        paired = group.assign(gain=group["postest_total"] - group["pretest_total"])
        result = _paired_statistics(paired)
        primary_name = result.get("primary")
        primary = next((test for test in result.get("tests", []) if test["name"] == primary_name), None)
        rows.append({
            "comparacion": f"Pretest vs Postest - {group_name}", "variable": "cambio_intragrupo",
            "n_experimental": len(group) if group_name == "Experimental" else np.nan,
            "n_control": len(group) if group_name == "Control" else np.nan,
            "media_experimental": group["postest_total"].mean() - group["pretest_total"].mean() if group_name == "Experimental" else np.nan,
            "media_control": group["postest_total"].mean() - group["pretest_total"].mean() if group_name == "Control" else np.nan,
            "normalidad": result.get("normality", {}).get("p", np.nan) >= ALPHA,
            "levene_p": np.nan, "prueba": primary_name or "Solo descriptiva",
            "estadistico": primary.get("statistic") if primary else np.nan,
            "p": primary.get("p_greater") if primary else np.nan,
            "efecto": result.get("effect", {}).get("cohen_dz", np.nan), "tipo_efecto": "Cohen d_z",
        })
    return pd.DataFrame(rows)


def ancova_analysis(dataset):
    result = {"available": False, "reason": ""}
    if smf is None or anova_lm is None:
        result["reason"] = "statsmodels no disponible."
        return result
    data = dataset[["grupo", "pretest_total", "postest_total"]].dropna()
    data = data[data["grupo"].isin(["Experimental", "Control"])].copy()
    counts = data["grupo"].value_counts()
    if len(data) < 20 or not {"Experimental", "Control"}.issubset(counts.index) or counts.min() < 5:
        result["reason"] = "ANCOVA requiere ambos grupos, al menos 5 casos por grupo y 20 casos completos como mínimo operativo."
        return result
    data["grupo_codigo"] = (data["grupo"] == "Experimental").astype(int)
    interaction_model = smf.ols("postest_total ~ pretest_total * grupo_codigo", data=data).fit()
    interaction_p = float(interaction_model.pvalues.get("pretest_total:grupo_codigo", np.nan))
    model = smf.ols("postest_total ~ pretest_total + grupo_codigo", data=data).fit()
    anova = anova_lm(model, typ=2)
    group_ss = float(anova.loc["grupo_codigo", "sum_sq"])
    residual_ss = float(anova.loc["Residual", "sum_sq"])
    eta_partial = group_ss / (group_ss + residual_ss) if group_ss + residual_ss else np.nan
    mean_pre = data["pretest_total"].mean()
    adjusted = {
        "Control": float(model.predict(pd.DataFrame({"pretest_total": [mean_pre], "grupo_codigo": [0]})).iloc[0]),
        "Experimental": float(model.predict(pd.DataFrame({"pretest_total": [mean_pre], "grupo_codigo": [1]})).iloc[0]),
    }
    residual_normality = _shapiro(model.resid)
    levene = stats.levene(model.resid[data["grupo"].eq("Experimental")], model.resid[data["grupo"].eq("Control")], center="median")
    conf = model.conf_int().loc["grupo_codigo"].tolist()
    result.update({
        "available": True, "n": len(data), "model": model, "anova": anova.reset_index().rename(columns={"index": "fuente"}),
        "group_beta": float(model.params["grupo_codigo"]), "group_p": float(model.pvalues["grupo_codigo"]),
        "group_ci": conf, "eta_partial": eta_partial, "adjusted_means": adjusted,
        "r2": float(model.rsquared), "r2_adjusted": float(model.rsquared_adj),
        "residual_shapiro_p": residual_normality.get("p"), "levene_residual_p": float(levene.pvalue),
        "slope_interaction_p": interaction_p,
    })
    return result


def correlation_analysis(dataset):
    outcomes = [column for column in ["postest_total", "ganancia_aprendizaje", "comprension_conceptual_post", "procedimientos_post", "aplicaciones_post", "resolucion_problemas_post"] if column in dataset]
    predictors = [column for column in PREDICTOR_COLUMNS if column in dataset]
    rows = []
    for predictor in predictors:
        for outcome in outcomes:
            pair = dataset[[predictor, outcome]].apply(pd.to_numeric, errors="coerce").dropna()
            if len(pair) < 5 or pair[predictor].nunique() < 2 or pair[outcome].nunique() < 2:
                continue
            ordinal = predictor in {"nivel_adaptativo", "motivacion", "autoeficacia", "usabilidad", "grupo_experimental"}
            normal = _shapiro(pair[predictor]).get("p", 0) >= ALPHA and _shapiro(pair[outcome]).get("p", 0) >= ALPHA
            if normal and not ordinal:
                test, method = stats.pearsonr(pair[predictor], pair[outcome]), "Pearson"
            else:
                test, method = stats.spearmanr(pair[predictor], pair[outcome]), "Spearman"
            rows.append({"predictor": predictor, "resultado": outcome, "n": len(pair), "metodo": method, "correlacion": float(test.statistic), "p": float(test.pvalue)})
    return pd.DataFrame(rows)


def _prepare_model_data(dataset, outcome):
    candidates = [column for column in PREDICTOR_COLUMNS if column in dataset and column != outcome]
    usable = []
    for column in candidates:
        values = _numeric(dataset[column])
        if values.isna().mean() <= 0.20 and values.nunique(dropna=True) > 1:
            usable.append(column)
    complete = dataset[[outcome] + usable].copy()
    complete[outcome] = _numeric(complete[outcome])
    complete = complete.dropna(subset=[outcome])
    for column in usable:
        complete[column] = _numeric(complete[column])
        complete[column] = complete[column].fillna(complete[column].median())
    max_features = max(1, len(complete) // 10)
    if len(usable) > max_features:
        correlations = complete[usable].corrwith(complete[outcome]).abs().sort_values(ascending=False)
        usable = correlations.head(max_features).index.tolist()
    return complete, usable


def predictive_models(dataset, outcome="postest_total"):
    result = {"available": False, "reason": "", "metrics": pd.DataFrame(), "importance": pd.DataFrame()}
    data, features = _prepare_model_data(dataset, outcome)
    if len(data) < 30 or not features:
        result["reason"] = "Se requieren al menos 30 casos completos y predictores con variabilidad; además se limita el modelo a aproximadamente un predictor por cada 10 casos."
        return result
    X, y = data[features], data[outcome]
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.20, random_state=RANDOM_STATE)
    folds = min(5, max(3, len(X_train) // 8))
    cv = KFold(n_splits=folds, shuffle=True, random_state=RANDOM_STATE)
    models = {
        "Regresión lineal": Pipeline([("scale", StandardScaler()), ("model", LinearRegression())]),
        "Ridge": Pipeline([("scale", StandardScaler()), ("model", RidgeCV(alphas=np.logspace(-3, 3, 25)))]),
        "Lasso": Pipeline([("scale", StandardScaler()), ("model", LassoCV(cv=min(5, folds), random_state=RANDOM_STATE, max_iter=20000))]),
        "Random Forest": RandomForestRegressor(n_estimators=300, min_samples_leaf=2, random_state=RANDOM_STATE),
    }
    if len(data) >= 80:
        models["Gradient Boosting"] = GradientBoostingRegressor(random_state=RANDOM_STATE)
    metrics, fitted = [], {}
    for name, model in models.items():
        model.fit(X_train, y_train)
        pred = model.predict(X_test)
        cv_rmse = -cross_val_score(model, X_train, y_train, cv=cv, scoring="neg_root_mean_squared_error").mean()
        denominator = np.maximum(np.abs(y_test.to_numpy()), 1)
        r2 = r2_score(y_test, pred) if len(y_test) >= 2 else np.nan
        adjusted_r2 = 1 - (1 - r2) * (len(y_test) - 1) / (len(y_test) - len(features) - 1) if np.isfinite(r2) and len(y_test) > len(features) + 1 else np.nan
        metrics.append({"modelo": name, "R2_prueba": r2, "R2_ajustado": adjusted_r2, "MAE": mean_absolute_error(y_test, pred), "RMSE": math.sqrt(mean_squared_error(y_test, pred)), "MAPE_pct": float(np.mean(np.abs((y_test.to_numpy() - pred) / denominator)) * 100), "RMSE_CV": cv_rmse})
        fitted[name] = model
    metrics_df = pd.DataFrame(metrics).sort_values("RMSE_CV")
    best_name = metrics_df.iloc[0]["modelo"]
    best = fitted[best_name]
    if hasattr(best, "feature_importances_"):
        values = best.feature_importances_
    elif hasattr(best, "named_steps"):
        estimator = best.named_steps["model"]
        values = np.abs(getattr(estimator, "coef_", np.zeros(len(features))))
    else:
        values = np.zeros(len(features))
    importance = pd.DataFrame({"variable": features, "importancia": values}).sort_values("importancia", ascending=False)
    result.update({"available": True, "n": len(data), "features": features, "metrics": metrics_df, "best_name": best_name, "importance": importance, "y_test": y_test.to_numpy(), "predictions": best.predict(X_test), "residuals": y_test.to_numpy() - best.predict(X_test)})
    return result


def risk_models(dataset):
    result = {"available": False, "reason": "", "metrics": pd.DataFrame()}
    data, features = _prepare_model_data(dataset, "postest_total")
    if len(data) < 40 or not features:
        result["reason"] = "La clasificación de riesgo requiere al menos 40 casos completos y predictores con variabilidad."
        return result
    y = (data["postest_total"] < 70).astype(int)
    if y.value_counts().min() < 5:
        result["reason"] = "Se necesitan al menos 5 estudiantes en cada clase de riesgo para una validación mínima."
        return result
    X = data[features]
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.25, stratify=y, random_state=RANDOM_STATE)
    cv = StratifiedKFold(n_splits=min(5, int(y.value_counts().min())), shuffle=True, random_state=RANDOM_STATE)
    models = {
        "Regresión logística": Pipeline([("scale", StandardScaler()), ("model", LogisticRegression(max_iter=3000, class_weight="balanced", random_state=RANDOM_STATE))]),
        "Árbol de decisión": DecisionTreeClassifier(max_depth=4, min_samples_leaf=3, class_weight="balanced", random_state=RANDOM_STATE),
        "Random Forest": RandomForestClassifier(n_estimators=300, min_samples_leaf=2, class_weight="balanced", random_state=RANDOM_STATE),
    }
    metrics, fitted = [], {}
    for name, model in models.items():
        model.fit(X_train, y_train)
        pred = model.predict(X_test)
        prob = model.predict_proba(X_test)[:, 1]
        fpr, tpr, _ = roc_curve(y_test, prob)
        auc_cv = cross_val_score(model, X, y, cv=cv, scoring="roc_auc").mean()
        metrics.append({"modelo": name, "accuracy": accuracy_score(y_test, pred), "precision": precision_score(y_test, pred, zero_division=0), "recall": recall_score(y_test, pred, zero_division=0), "F1": f1_score(y_test, pred, zero_division=0), "AUC": auc(fpr, tpr), "AUC_CV": auc_cv})
        fitted[name] = (model, pred, prob)
    metrics_df = pd.DataFrame(metrics).sort_values(["AUC", "recall"], ascending=False)
    best_name = metrics_df.iloc[0]["modelo"]
    _, pred, prob = fitted[best_name]
    fpr, tpr, _ = roc_curve(y_test, prob)
    result.update({"available": True, "n": len(data), "features": features, "metrics": metrics_df, "best_name": best_name, "confusion": confusion_matrix(y_test, pred), "fpr": fpr, "tpr": tpr, "auc": auc(fpr, tpr)})
    return result


def _fig_bytes(fig):
    stream = BytesIO()
    fig.savefig(stream, format="png", dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    stream.seek(0)
    return stream.getvalue()


def _slug(value):
    return re.sub(r"[^a-z0-9]+", "_", str(value).lower()).strip("_")


def _eda_numeric_plot(series, label):
    values = _numeric(series).dropna().to_numpy(dtype=float)
    fig, axes = plt.subplots(2, 2, figsize=(10, 7.2))
    bins = min(15, max(5, int(math.sqrt(len(values))) + 1))
    axes[0, 0].hist(values, bins=bins, color=COLORS["pre"], edgecolor="white", alpha=0.82)
    axes[0, 0].set_title("Histograma")
    if len(values) >= 3 and np.ptp(values) > 0:
        grid = np.linspace(values.min(), values.max(), 200)
        try:
            density = stats.gaussian_kde(values)
            axes[0, 1].plot(grid, density(grid), color=COLORS["gain"], linewidth=2)
            axes[0, 1].fill_between(grid, density(grid), color=COLORS["gain"], alpha=0.18)
        except np.linalg.LinAlgError:
            axes[0, 1].text(0.5, 0.5, "Densidad no estimable", ha="center", va="center", transform=axes[0, 1].transAxes)
    axes[0, 1].set_title("Densidad")
    axes[1, 0].boxplot(values, vert=False, patch_artist=True, boxprops={"facecolor": COLORS["post"], "alpha": 0.65})
    axes[1, 0].set_title("Boxplot")
    if len(values) >= 3 and np.ptp(values) > 0:
        stats.probplot(values, dist="norm", plot=axes[1, 1])
    else:
        axes[1, 1].text(0.5, 0.5, "Q-Q no estimable", ha="center", va="center", transform=axes[1, 1].transAxes)
    axes[1, 1].set_title("Gráfico Q-Q")
    for ax in axes.ravel():
        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(axis="y", color="#E2E8F0", alpha=0.5)
    fig.suptitle(f"EDA: {label}", fontsize=14, fontweight="bold")
    fig.tight_layout()
    return fig


def _eda_categorical_plot(series, label):
    counts = series.fillna("Sin dato").astype(str).value_counts().head(15)
    fig, axes = plt.subplots(1, 2 if len(counts) <= 6 else 1, figsize=(10, 4.8))
    axes = np.atleast_1d(axes)
    axes[0].barh(counts.index[::-1], counts.values[::-1], color=COLORS["gain"])
    axes[0].set_title("Frecuencia absoluta")
    axes[0].set_xlabel("n")
    if len(axes) > 1:
        axes[1].pie(counts.values, labels=counts.index, autopct="%1.1f%%", colors=[COLORS["pre"], COLORS["post"], COLORS["gain"], COLORS["accent"], "#7C3AED", "#0891B2"][:len(counts)])
        axes[1].set_title("Distribución porcentual")
    fig.suptitle(f"EDA: {label}", fontsize=14, fontweight="bold")
    fig.tight_layout()
    return fig


def _heatmap(correlations):
    pivot = correlations.pivot_table(index="predictor", columns="resultado", values="correlacion")
    fig, ax = plt.subplots(figsize=(9, max(4, 0.45 * len(pivot) + 2)))
    image = ax.imshow(pivot.fillna(0), cmap="RdBu_r", vmin=-1, vmax=1, aspect="auto")
    ax.set_xticks(range(len(pivot.columns)), pivot.columns, rotation=25, ha="right")
    ax.set_yticks(range(len(pivot.index)), pivot.index)
    fig.colorbar(image, ax=ax, label="Correlación")
    ax.set_title("Matriz de correlaciones predictoras-resultados")
    return fig


def _item_matrix_heatmap(matrix, label):
    matrix = pd.DataFrame(matrix)
    fig, ax = plt.subplots(figsize=(8.5, 7.0))
    image = ax.imshow(matrix.fillna(0), cmap="RdBu_r", vmin=-1, vmax=1, aspect="auto")
    short_labels = [str(value)[:12] for value in matrix.columns]
    ax.set_xticks(range(len(short_labels)), short_labels, rotation=90, fontsize=6)
    ax.set_yticks(range(len(short_labels)), short_labels, fontsize=6)
    fig.colorbar(image, ax=ax, label="Correlación entre ítems")
    ax.set_title(f"Matriz de correlaciones: {label}")
    fig.tight_layout()
    return fig


def _importance_plot(importance):
    ordered = importance.sort_values("importancia")
    fig, ax = plt.subplots(figsize=(8, max(3.5, 0.45 * len(ordered) + 1.5)))
    ax.barh(ordered["variable"], ordered["importancia"], color=COLORS["gain"])
    ax.set_title("Importancia de variables del mejor modelo")
    ax.set_xlabel("Importancia relativa")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _residual_plot(model_result):
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.scatter(model_result["predictions"], model_result["residuals"], color=COLORS["pre"], alpha=0.8)
    ax.axhline(0, color=COLORS["accent"], linestyle="--")
    ax.set_xlabel("Valor predicho")
    ax.set_ylabel("Residuo")
    ax.set_title("Residuos del modelo predictivo")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _risk_plots(risk):
    fig1, ax1 = plt.subplots(figsize=(5.5, 4.5))
    ax1.imshow(risk["confusion"], cmap="Blues")
    for (i, j), value in np.ndenumerate(risk["confusion"]):
        ax1.text(j, i, str(value), ha="center", va="center", fontsize=14)
    ax1.set_xticks([0, 1], ["Sin riesgo", "Riesgo"])
    ax1.set_yticks([0, 1], ["Sin riesgo", "Riesgo"])
    ax1.set_xlabel("Predicción")
    ax1.set_ylabel("Real")
    ax1.set_title("Matriz de confusión")
    fig2, ax2 = plt.subplots(figsize=(5.5, 4.5))
    ax2.plot(risk["fpr"], risk["tpr"], color=COLORS["post"], label=f"AUC={risk['auc']:.3f}")
    ax2.plot([0, 1], [0, 1], color=COLORS["muted"], linestyle="--")
    ax2.set_xlabel("Tasa de falsos positivos")
    ax2.set_ylabel("Sensibilidad")
    ax2.set_title("Curva ROC")
    ax2.legend(frameon=False)
    return fig1, fig2


def _group_boxplot(dataset):
    data = dataset[dataset["grupo"].isin(["Experimental", "Control"])]
    fig, ax = plt.subplots(figsize=(8, 4.8))
    labels, values = [], []
    for group in ["Experimental", "Control"]:
        for moment, column in [("Pretest", "pretest_total"), ("Postest", "postest_total")]:
            current = _numeric(data.loc[data["grupo"].eq(group), column]).dropna().to_numpy()
            if len(current):
                labels.append(f"{group}\n{moment}")
                values.append(current)
    boxes = ax.boxplot(values, tick_labels=labels, patch_artist=True)
    for box, color in zip(boxes["boxes"], [COLORS["pre"], COLORS["post"]] * 2):
        box.set_facecolor(color)
        box.set_alpha(0.7)
    ax.set_ylabel("Puntaje")
    ax.set_ylim(0, 105)
    ax.set_title("Puntajes por grupo y momento")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _radar_dimensions(dataset):
    codes = ["comprension_conceptual", "procedimientos", "aplicaciones", "resolucion_problemas"]
    labels = ["Comprensión", "Procedimientos", "Aplicaciones", "Resolución"]
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]
    fig, ax = plt.subplots(figsize=(6.5, 6.0), subplot_kw={"polar": True})
    for suffix, label, color in [("pre", "Pretest", COLORS["pre"]), ("post", "Postest", COLORS["post"])]:
        values = [float(_numeric(dataset.get(f"{code}_{suffix}", pd.Series(dtype=float))).mean()) for code in codes]
        values += values[:1]
        ax.plot(angles, values, label=label, color=color, linewidth=2)
        ax.fill(angles, values, color=color, alpha=0.12)
    ax.set_xticks(angles[:-1], labels)
    ax.set_ylim(0, 100)
    ax.set_title("Radar de competencias")
    ax.legend(loc="upper right", bbox_to_anchor=(1.2, 1.1), frameon=False)
    return fig


def _learning_curve(exercise_df):
    data = pd.DataFrame(exercise_df).copy()
    data["correct"] = _bool_series(data["is_correct"]).astype(float)
    if "created_at" in data:
        data = data.sort_values("created_at")
    data["orden_intento"] = data.groupby("user_id").cumcount() + 1
    curve = data.groupby("orden_intento")["correct"].agg(["mean", "count"]).reset_index()
    fig, ax = plt.subplots(figsize=(8, 4.8))
    ax.plot(curve["orden_intento"], curve["mean"] * 100, marker="o", color=COLORS["gain"], linewidth=2)
    ax.set_xlabel("Orden del ejercicio")
    ax.set_ylabel("Aciertos (%)")
    ax.set_ylim(0, 105)
    ax.set_title("Curva de aprendizaje durante ejercicios guiados")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _ancova_adjusted_plot(ancova):
    labels = ["Control", "Experimental"]
    values = [ancova["adjusted_means"][label] for label in labels]
    fig, ax = plt.subplots(figsize=(6.5, 4.5))
    bars = ax.bar(labels, values, color=[COLORS["pre"], COLORS["post"]])
    ax.bar_label(bars, labels=[f"{value:.1f}" for value in values], padding=3)
    ax.set_ylim(0, 105)
    ax.set_ylabel("Media postest ajustada")
    ax.set_title("Medias ajustadas por el pretest (ANCOVA)")
    ax.spines[["top", "right"]].set_visible(False)
    return fig


def _df_rows(df, columns, formatters=None, max_rows=100):
    formatters = formatters or {}
    rows = []
    for _, row in df.head(max_rows).iterrows():
        values = []
        for column in columns:
            value = row.get(column, "")
            values.append(formatters[column](value) if column in formatters else str(value))
        rows.append(values)
    return rows


def _append_advanced_word(base_word, dataset, cleaning, normality, reliability, items, groups, ancova, correlations, prediction, risk, figures, title):
    doc = Document(BytesIO(base_word))
    doc.add_page_break()
    doc.add_heading("10. Motor estadístico-predictivo", level=1)
    doc.add_paragraph("Esta sección amplía el análisis confirmatorio, psicométrico y predictivo. Cada procedimiento se ejecuta únicamente cuando existen variables, grupos, tamaño muestral y variabilidad suficientes.")
    doc.add_heading("10.1 Dataset científico y limpieza", level=2)
    _add_table(doc, ["Variable", "% perdidos", "Decisión", "Atípicos RIC"], _df_rows(cleaning, ["variable", "perdidos_pct", "decision", "atipicos_iqr"], {"perdidos_pct": lambda x: _number(x, 1) + "%"}), font_size=7.5)
    _add_note(doc, "Imputación", "Los resultados principales y la asignación de grupo no se imputan. Las variables auxiliares con hasta 5% de ausencia pueden imputarse de forma simple; entre 5% y 20% se advierte la necesidad de imputación múltiple; por encima de 20% se excluyen de modelos principales.")
    _add_note(doc, "Definiciones operativas", "El tiempo de uso es una estimación formada por duración de evaluaciones, eventos declarados e intervalos activos entre registros, limitados a 10 minutos por intervalo. Las clases vistas se aproximan mediante clases generadas y la retroalimentación mediante ejercicios revisados; estas métricas deben reportarse como aproximaciones de trazas digitales.")

    doc.add_heading("10.2 Normalidad por variable", level=2)
    _add_table(doc, ["Variable", "n", "Prueba", "Estadístico", "p", "Decisión", "Asimetría", "Curtosis"], _df_rows(normality, ["variable", "n", "prueba", "estadistico", "p", "decision", "asimetria", "curtosis"], {"estadistico": lambda x: _number(x, 3), "p": _pvalue, "asimetria": lambda x: _number(x, 3), "curtosis": lambda x: _number(x, 3)}), font_size=6.8)

    doc.add_heading("10.3 Análisis psicométrico", level=2)
    if reliability.empty:
        doc.add_paragraph("No hay respuestas por ítem suficientes para estimar confiabilidad.")
    else:
        _add_table(doc, ["Instrumento", "Versión", "n", "Ítems", "Alfa", "Omega", "Cobertura"], _df_rows(reliability, ["instrumento", "version", "casos_completos", "items", "alfa_cronbach", "omega_mcdonald_aprox", "cobertura_completa_pct"], {"alfa_cronbach": lambda x: _number(x, 3), "omega_mcdonald_aprox": lambda x: _number(x, 3), "cobertura_completa_pct": lambda x: _number(x, 1) + "%"}), font_size=8)
        _add_note(doc, "Interpretación", "Valores de alfa u omega desde 0,70 se consideran aceptables como referencia general. El omega se estima mediante una solución unifactorial aproximada y requiere revisión de dimensionalidad.")
    if not items.empty:
        _add_table(doc, ["Instrumento", "Versión", "Ítem", "n", "p", "Dificultad", "Discriminación", "Calidad"], _df_rows(items, ["instrumento", "version", "item_code", "n", "dificultad_p", "clasificacion_dificultad", "discriminacion_item_total", "calidad_item"], {"dificultad_p": lambda x: _number(x, 3), "discriminacion_item_total": lambda x: _number(x, 3)}, max_rows=80), font_size=6.8)

    doc.add_heading("10.4 Comparación experimental-control", level=2)
    if groups.empty:
        doc.add_paragraph("No disponible: asigne estudiantes a los grupos Experimental y Control y complete ambas mediciones.")
    else:
        _add_table(doc, ["Comparación", "Variable", "n Exp.", "n Ctrl.", "M/dif. Exp.", "M/dif. Ctrl.", "Prueba", "p", "Efecto"], _df_rows(groups, ["comparacion", "variable", "n_experimental", "n_control", "media_experimental", "media_control", "prueba", "p", "efecto"], {"media_experimental": lambda x: _number(x), "media_control": lambda x: _number(x), "p": _pvalue, "efecto": lambda x: _number(x, 3)}), font_size=6.8)
        _add_note(doc, "Selección", "Se utiliza t independiente con normalidad y varianzas homogéneas, Welch ante heterogeneidad y Mann-Whitney cuando la normalidad no es defendible.")

    doc.add_heading("10.5 ANCOVA", level=2)
    if not ancova.get("available"):
        doc.add_paragraph("No ejecutada: " + ancova.get("reason", "datos insuficientes"))
    else:
        _add_table(doc, ["n", "β grupo", "IC 95%", "p grupo", "η² parcial", "R² ajustado", "p interacción pendientes"], [[ancova["n"], _number(ancova["group_beta"]), f"{_number(ancova['group_ci'][0])} a {_number(ancova['group_ci'][1])}", _pvalue(ancova["group_p"]), _number(ancova["eta_partial"], 3), _number(ancova["r2_adjusted"], 3), _pvalue(ancova["slope_interaction_p"])]] )
        _add_note(doc, "Interpretación", f"Al controlar el pretest, la diferencia ajustada Experimental-Control fue {_number(ancova['group_beta'])} puntos. El efecto del grupo {'fue significativo' if ancova['group_p'] < ALPHA else 'no fue significativo'} (p {_pvalue(ancova['group_p'])}). La homogeneidad de pendientes {'es compatible con el modelo' if ancova['slope_interaction_p'] >= ALPHA else 'presenta evidencia de incumplimiento'}.")

    doc.add_heading("10.6 Correlaciones y analítica de aprendizaje", level=2)
    if correlations.empty:
        doc.add_paragraph("No existen pares de variables con tamaño y variabilidad suficientes.")
    else:
        strongest = correlations.assign(abs_r=correlations["correlacion"].abs()).sort_values("abs_r", ascending=False).head(25)
        _add_table(doc, ["Predictor", "Resultado", "n", "Método", "r/ρ", "p"], _df_rows(strongest, ["predictor", "resultado", "n", "metodo", "correlacion", "p"], {"correlacion": lambda x: _number(x, 3), "p": _pvalue}), font_size=7.5)
        if "correlaciones.png" in figures:
            doc.add_picture(BytesIO(figures["correlaciones.png"]), width=Inches(6.25))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("Figura 11. Heatmap de correlaciones")
            caption_run.bold = True
            caption_run.font.size = Pt(9)
            _add_note(doc, "Interpretación", "Los colores muestran dirección y magnitud. Las correlaciones son asociaciones y no demuestran causalidad.")

    doc.add_heading("10.7 Modelo predictivo de rendimiento", level=2)
    if not prediction.get("available"):
        doc.add_paragraph("No entrenado: " + prediction.get("reason", "datos insuficientes"))
    else:
        _add_table(doc, ["Modelo", "R² prueba", "R² ajustado", "MAE", "RMSE", "MAPE %", "RMSE CV"], _df_rows(prediction["metrics"], ["modelo", "R2_prueba", "R2_ajustado", "MAE", "RMSE", "MAPE_pct", "RMSE_CV"], {column: lambda x: _number(x, 3) for column in ["R2_prueba", "R2_ajustado", "MAE", "RMSE", "MAPE_pct", "RMSE_CV"]}), font_size=7.5)
        _add_note(doc, "Modelo seleccionado", f"{prediction['best_name']} por menor RMSE de validación cruzada. El resultado es predictivo, no causal, y requiere validación externa antes de decisiones individuales.")

    doc.add_heading("10.8 Clasificación de riesgo académico", level=2)
    if not risk.get("available"):
        doc.add_paragraph("No entrenada: " + risk.get("reason", "datos insuficientes"))
    else:
        _add_table(doc, ["Modelo", "Accuracy", "Precisión", "Recall", "F1", "AUC", "AUC CV"], _df_rows(risk["metrics"], ["modelo", "accuracy", "precision", "recall", "F1", "AUC", "AUC_CV"], {column: lambda x: _number(x, 3) for column in ["accuracy", "precision", "recall", "F1", "AUC", "AUC_CV"]}), font_size=7.5)
        _add_note(doc, "Uso responsable", "La clasificación sirve para priorizar apoyo pedagógico y nunca para sancionar, excluir o etiquetar de forma permanente. Debe revisarse sesgo, calibración y rendimiento por subgrupo.")

    doc.add_heading("Anexo C. EDA gráfico por variable", level=1)
    figure_number = 12
    for column in dataset.columns:
        prefix = "eda_numerica" if pd.api.types.is_numeric_dtype(dataset[column]) else "eda_categorica"
        key = f"{prefix}_{_slug(column)}.png"
        if key not in figures:
            continue
        doc.add_picture(BytesIO(figures[key]), width=Inches(6.25))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(f"Figura {figure_number}. EDA de {column}")
        caption_run.bold = True
        caption_run.font.size = Pt(9)
        if pd.api.types.is_numeric_dtype(dataset[column]):
            values = _numeric(dataset[column]).dropna()
            interpretation = f"n={len(values)}, media={_number(values.mean())}, mediana={_number(values.median())} y DE={_number(values.std(ddof=1) if len(values)>1 else np.nan)}. El histograma, densidad, boxplot y Q-Q deben leerse conjuntamente para valorar forma, dispersión, atípicos y normalidad."
        else:
            counts = dataset[column].fillna("Sin dato").astype(str).value_counts()
            interpretation = f"La categoría modal es {counts.index[0]} con {int(counts.iloc[0])} casos ({counts.iloc[0] / max(1, counts.sum()) * 100:.1f}%). Las frecuencias describen la composición y no implican diferencias estadísticamente significativas."
        _add_note(doc, "Interpretación", interpretation)
        figure_number += 1

    doc.add_heading("Anexo D. Guía metodológica de interpretación", level=1)
    method_rows = [
        ["Shapiro-Wilk / Lilliefors", "Evaluar normalidad", "Datos continuos e independencia", "La distribución es normal", "Sensible al tamaño muestral", "Shapiro y Wilk (1965)"],
        ["Levene", "Comparar varianzas", "Grupos independientes", "Las varianzas son iguales", "Baja potencia con n pequeño", "Field (2024)"],
        ["t pareada", "Comparar pre-post", "Diferencias aproximadamente normales", "La diferencia media es cero", "No controla historia o maduración", "Cohen (1988)"],
        ["Wilcoxon", "Comparar rangos pre-post", "Pares y diferencias simétricas", "La distribución de diferencias se centra en cero", "No estima directamente diferencia de medias", "Wilcoxon (1945)"],
        ["Mann-Whitney", "Comparar dos grupos", "Independencia y forma comparable", "Las distribuciones son equivalentes", "No siempre prueba medianas", "Mann y Whitney (1947)"],
        ["ANCOVA", "Comparar postest controlando pretest", "Linealidad, pendientes homogéneas, residuos adecuados", "El grupo no aporta efecto ajustado", "No corrige sesgo de asignación", "Tabachnick y Fidell (2021)"],
        ["Pearson / Spearman", "Medir asociación", "Linealidad para Pearson; monotonía para Spearman", "La asociación es cero", "Correlación no implica causalidad", "Pearson (1896); Spearman (1904)"],
        ["Alfa / Omega", "Estimar consistencia interna", "Ítems relacionados; unidimensionalidad para lectura simple", "No se usa como contraste causal", "No demuestra validez", "McDonald (1999)"],
        ["Regresión", "Predecir resultado continuo", "Linealidad/residuos para inferencia OLS", "Los coeficientes son cero", "Sobreajuste y confusión", "Hair et al. (2022)"],
        ["Clasificación", "Identificar riesgo <70", "Clases observadas y validación independiente", "Rendimiento no superior al azar", "Puede amplificar sesgos", "Hair et al. (2022)"],
    ]
    _add_table(doc, ["Método", "Objetivo", "Supuestos", "H0", "Limitación", "Referencia"], method_rows, font_size=6.8)

    doc.add_heading("11. Referencias", level=1)
    references = [
        "American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.).",
        "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum.",
        "Field, A. (2024). Discovering statistics using IBM SPSS Statistics. SAGE.",
        "Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2022). Multivariate data analysis. Cengage.",
        "Hake, R. R. (1998). Interactive-engagement versus traditional methods. American Journal of Physics, 66(1), 64-74.",
        "Mann, H. B., & Whitney, D. R. (1947). On a test of whether one of two random variables is stochastically larger than the other. The Annals of Mathematical Statistics, 18(1), 50-60.",
        "McDonald, R. P. (1999). Test theory: A unified treatment. Lawrence Erlbaum.",
        "Pallant, J. (2020). SPSS survival manual. McGraw-Hill Education.",
        "Pearson, K. (1896). Mathematical contributions to the theory of evolution. Philosophical Transactions of the Royal Society of London.",
        "Shapiro, S. S., & Wilk, M. B. (1965). An analysis of variance test for normality. Biometrika, 52(3/4), 591-611.",
        "Spearman, C. (1904). The proof and measurement of association between two things. The American Journal of Psychology, 15(1), 72-101.",
        "Wilcoxon, F. (1945). Individual comparisons by ranking methods. Biometrics Bulletin, 1(6), 80-83.",
        "Tabachnick, B. G., & Fidell, L. S. (2021). Using multivariate statistics. Pearson.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(reference)
        paragraph.paragraph_format.left_indent = 18
        paragraph.paragraph_format.first_line_indent = -18
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _excel_bytes(sheets):
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        for name, frame in sheets.items():
            pd.DataFrame(frame).to_excel(writer, sheet_name=name[:31], index=False)
            sheet = writer.sheets[name[:31]]
            sheet.freeze_panes = "A2"
            sheet.auto_filter.ref = sheet.dimensions
            for column_cells in sheet.columns:
                width = min(45, max(10, max(len(str(cell.value or "")) for cell in column_cells) + 2))
                sheet.column_dimensions[column_cells[0].column_letter].width = width
    return output.getvalue()


def _pdf_bytes(title, dataset, reliability, groups, ancova, prediction, risk, figures):
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=1.7 * cm, leftMargin=1.7 * cm, topMargin=1.7 * cm, bottomMargin=1.7 * cm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#6F0F49")))
    story = [Paragraph("Informe estadístico-predictivo", styles["CenterTitle"]), Paragraph(title, styles["Heading2"]), Spacer(1, 10)]
    story.append(Paragraph(f"Muestra procesada: {len(dataset)} estudiantes. Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["BodyText"]))
    sections = [
        ("Confiabilidad", reliability, ["instrumento", "version", "casos_completos", "items", "alfa_cronbach", "omega_mcdonald_aprox"]),
        ("Comparación entre grupos", groups, ["variable", "prueba", "p", "efecto"]),
        ("Modelos predictivos", prediction.get("metrics", pd.DataFrame()), ["modelo", "R2_prueba", "R2_ajustado", "MAE", "RMSE", "RMSE_CV"]),
        ("Clasificación de riesgo", risk.get("metrics", pd.DataFrame()), ["modelo", "accuracy", "precision", "recall", "F1", "AUC", "AUC_CV"]),
    ]
    for heading, frame, columns_to_show in sections:
        story.extend([Spacer(1, 10), Paragraph(heading, styles["Heading2"])])
        frame = pd.DataFrame(frame)
        if frame.empty:
            story.append(Paragraph("No disponible por datos insuficientes o variables ausentes.", styles["BodyText"]))
            continue
        columns_valid = [column for column in columns_to_show if column in frame.columns]
        data = [columns_valid] + [[str(value)[:34] for value in row] for row in frame[columns_valid].round(3).head(25).itertuples(index=False, name=None)]
        table = Table(data, repeatRows=1, hAlign="LEFT")
        table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#6F0F49")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.35, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 7), ("VALIGN", (0, 0), (-1, -1), "TOP")]))
        story.append(table)
    story.extend([PageBreak(), Paragraph("Gráficos", styles["Heading1"])])
    for name, raw in figures.items():
        explanation = "El panel combina distribución, dispersión, valores atípicos y ajuste aproximado a normalidad." if name.startswith("eda_numerica") else "El gráfico resume frecuencias y porcentajes de la variable." if name.startswith("eda_categorica") else "Gráfico complementario para interpretar asociación, predicción o clasificación; no constituye por sí solo evidencia causal."
        story.extend([Paragraph(name.replace("_", " ").replace(".png", "").title(), styles["Heading3"]), Image(BytesIO(raw), width=16 * cm, height=9 * cm), Paragraph(explanation, styles["BodyText"]), Spacer(1, 8)])
    doc.build(story)
    return output.getvalue()


def _python_script():
    return """# Análisis reproducible BunsekiChat\nimport pandas as pd\nfrom scipy import stats\n\ndf = pd.read_csv('dataset_procesado.csv')\npaired = df[['pretest_total','postest_total']].dropna()\ndiff = paired.postest_total - paired.pretest_total\nprint(df.describe(include='all'))\nif len(diff) >= 3:\n    print('Shapiro:', stats.shapiro(diff))\n    print('t pareada:', stats.ttest_rel(paired.postest_total, paired.pretest_total))\n    print('Wilcoxon:', stats.wilcoxon(paired.postest_total, paired.pretest_total))\n"""


def _r_script():
    return """# Análisis reproducible BunsekiChat\ndatos <- read.csv('dataset_procesado.csv')\npares <- na.omit(datos[, c('pretest_total','postest_total')])\nprint(summary(datos))\nif (nrow(pares) >= 3) {\n  diferencia <- pares$postest_total - pares$pretest_total\n  print(shapiro.test(diferencia))\n  print(t.test(pares$postest_total, pares$pretest_total, paired=TRUE))\n  print(wilcox.test(pares$postest_total, pares$pretest_total, paired=TRUE))\n}\n"""


def _json_records(frame):
    frame = pd.DataFrame(frame)
    if frame.empty:
        return []
    return json.loads(frame.to_json(orient="records", force_ascii=False, date_format="iso"))


def build_scientific_package(research_df, exercise_df=None, survey_df=None, interaction_df=None, event_df=None, item_df=None, filters=None, research_title="BunsekiChat"):
    research_df = pd.DataFrame(research_df).copy()
    exercise_df = pd.DataFrame() if exercise_df is None else pd.DataFrame(exercise_df).copy()
    survey_df = pd.DataFrame() if survey_df is None else pd.DataFrame(survey_df).copy()
    interaction_df = pd.DataFrame() if interaction_df is None else pd.DataFrame(interaction_df).copy()
    event_df = pd.DataFrame() if event_df is None else pd.DataFrame(event_df).copy()
    item_df = pd.DataFrame() if item_df is None else pd.DataFrame(item_df).copy()
    dataset, missing_columns = build_scientific_dataset(research_df, exercise_df, survey_df, interaction_df, event_df, item_df)
    cleaned, cleaning_log = clean_dataset(dataset)
    normality = normality_analysis(cleaned)
    reliability, item_stats, item_matrices = psychometric_analysis(item_df)
    comparisons = group_comparisons(cleaned)
    ancova = ancova_analysis(cleaned)
    correlations = correlation_analysis(cleaned)
    prediction = predictive_models(cleaned, "postest_total")
    risk = risk_models(cleaned)
    figures = {}
    for form_name, matrix in item_matrices.items():
        figures[f"matriz_items_{_slug(form_name)}.png"] = _fig_bytes(_item_matrix_heatmap(matrix, form_name))
    for column in cleaned.select_dtypes(include=[np.number]).columns:
        if column in {"estudiante_id", "grupo_experimental"}:
            continue
        values = _numeric(cleaned[column]).dropna()
        if len(values) >= 2 and values.nunique() > 1:
            figures[f"eda_numerica_{_slug(column)}.png"] = _fig_bytes(_eda_numeric_plot(values, column))
    for column in ["grupo", "materia", "curso", "paralelo", "jornada", "cohorte", "perfil_cognitivo", "nivel_dominio"]:
        if column in cleaned.columns and cleaned[column].notna().any():
            figures[f"eda_categorica_{_slug(column)}.png"] = _fig_bytes(_eda_categorical_plot(cleaned[column], column))
    if {"Experimental", "Control"}.issubset(set(cleaned.get("grupo", pd.Series(dtype=str)).dropna())):
        figures["boxplot_por_grupo.png"] = _fig_bytes(_group_boxplot(cleaned))
    dimension_columns = ["comprension_conceptual_pre", "procedimientos_pre", "aplicaciones_pre", "resolucion_problemas_pre", "comprension_conceptual_post", "procedimientos_post", "aplicaciones_post", "resolucion_problemas_post"]
    if all(column in cleaned.columns for column in dimension_columns) and cleaned[dimension_columns].notna().any().all():
        figures["radar_competencias.png"] = _fig_bytes(_radar_dimensions(cleaned))
    if not exercise_df.empty and {"user_id", "is_correct"}.issubset(exercise_df.columns):
        figures["curva_aprendizaje.png"] = _fig_bytes(_learning_curve(exercise_df))
    if ancova.get("available"):
        figures["medias_ajustadas_ancova.png"] = _fig_bytes(_ancova_adjusted_plot(ancova))
    if not correlations.empty:
        figures["correlaciones.png"] = _fig_bytes(_heatmap(correlations))
    if prediction.get("available"):
        figures["importancia_variables.png"] = _fig_bytes(_importance_plot(prediction["importance"]))
        figures["residuos_modelo.png"] = _fig_bytes(_residual_plot(prediction))
    if risk.get("available"):
        confusion_fig, roc_fig = _risk_plots(risk)
        figures["matriz_confusion.png"] = _fig_bytes(confusion_fig)
        figures["curva_roc.png"] = _fig_bytes(roc_fig)

    base_word = build_research_word_report(research_df, exercise_df, survey_df, filters, research_title)
    word = _append_advanced_word(base_word, cleaned, cleaning_log, normality, reliability, item_stats, comparisons, ancova, correlations, prediction, risk, figures, research_title)
    sheets = {
        "dataset_procesado": cleaned,
        "limpieza": cleaning_log,
        "normalidad": normality,
        "confiabilidad": reliability,
        "psicometria_items": item_stats,
        "comparacion_grupos": comparisons,
        "correlaciones": correlations,
        "modelos_predictivos": prediction.get("metrics", pd.DataFrame()),
        "riesgo_academico": risk.get("metrics", pd.DataFrame()),
        "respuestas_items": item_df,
        "encuesta": survey_df,
        "ejercicios": exercise_df,
    }
    if ancova.get("available"):
        sheets["ANCOVA"] = ancova["anova"]
    excel = _excel_bytes(sheets)
    csv_bytes = cleaned.to_csv(index=False).encode("utf-8-sig")
    json_payload = {
        "metadata": {"title": research_title, "generated_at": datetime.now().isoformat(), "filters": filters or {}, "missing_required_columns": missing_columns},
        "dataset": _json_records(cleaned),
        "cleaning": _json_records(cleaning_log),
        "normality": _json_records(normality),
        "reliability": _json_records(reliability),
        "item_statistics": _json_records(item_stats),
        "group_comparisons": _json_records(comparisons),
        "correlations": _json_records(correlations),
        "predictive_metrics": _json_records(prediction.get("metrics", pd.DataFrame())),
        "risk_metrics": _json_records(risk.get("metrics", pd.DataFrame())),
    }
    json_bytes = json.dumps(json_payload, ensure_ascii=False, indent=2, default=str, allow_nan=False).encode("utf-8")
    pdf = _pdf_bytes(research_title, cleaned, reliability, comparisons, ancova, prediction, risk, figures)
    files = {
        "informe_estadistico_predictivo.docx": word,
        "informe_estadistico_predictivo.pdf": pdf,
        "tablas_y_dataset.xlsx": excel,
        "dataset_procesado.csv": csv_bytes,
        "reporte_estadistico.json": json_bytes,
        "analisis_reproducible.py": _python_script().encode("utf-8"),
        "analisis_reproducible.R": _r_script().encode("utf-8"),
        **{f"graficos/{name}": raw for name, raw in figures.items()},
    }
    archive = BytesIO()
    with zipfile.ZipFile(archive, "w", zipfile.ZIP_DEFLATED) as zipped:
        for name, raw in files.items():
            zipped.writestr(name, raw)
    return {
        "word": word, "pdf": pdf, "excel": excel, "csv": csv_bytes, "json": json_bytes,
        "zip": archive.getvalue(), "dataset": cleaned, "summary": {
            "students": len(cleaned), "paired": int(cleaned[["pretest_total", "postest_total"]].dropna().shape[0]) if not cleaned.empty else 0,
            "psychometric_forms": len(reliability), "group_comparisons": len(comparisons),
            "predictive_available": prediction.get("available", False), "risk_available": risk.get("available", False),
            "warnings": ([f"Variables obligatorias sin datos disponibles: {', '.join(missing_columns)}"] if missing_columns else [])
            + [message for message in [prediction.get("reason"), risk.get("reason"), ancova.get("reason")] if message],
        },
    }
